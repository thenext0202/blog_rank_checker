"""docx 원본 줄 추출 — 검토용 임시 스크립트"""
import sys
from pathlib import Path
from docx import Document


def run_to_md(run):
    """run의 굵게 표시(bold) 마크업을 ** ** 로 변환"""
    text = run.text
    if not text:
        return ""
    if run.bold:
        # 양 끝 공백은 마크업 밖으로
        stripped = text.strip()
        if not stripped:
            return text
        lead = text[: len(text) - len(text.lstrip())]
        trail = text[len(text.rstrip()):]
        return f"{lead}**{stripped}**{trail}"
    return text


def para_to_md(para):
    return "".join(run_to_md(r) for r in para.runs)


def extract_lines(docx_path: Path):
    doc = Document(str(docx_path))
    lines = []
    # 본문 단락
    for para in doc.paragraphs:
        text = para_to_md(para)
        # 단락 안의 줄바꿈은 보존 (Word의 soft line break)
        for line in text.split("\n"):
            lines.append(line)
    # 표는 별도로 — 표 셀 텍스트를 파이프로 둘러싼 한 줄로 출력 (메타헤더 감지용)
    for table in doc.tables:
        for row in table.rows:
            cells_md = []
            for cell in row.cells:
                cell_text = " ".join(para_to_md(p) for p in cell.paragraphs).strip()
                cells_md.append(cell_text)
            lines.append("| " + " | ".join(cells_md) + " |")
    return lines


if __name__ == "__main__":
    path = Path(sys.argv[1])
    out = Path(sys.argv[2]) if len(sys.argv) > 2 else Path("output/_raw_lines.txt")
    out.parent.mkdir(parents=True, exist_ok=True)
    lines = extract_lines(path)
    with out.open("w", encoding="utf-8") as f:
        for i, line in enumerate(lines, 1):
            f.write(f"{i:3d}| {line}\n")
    print(f"wrote {len(lines)} lines -> {out}")
