"""
STEP 1: 본문 식별기
입력 docx에서 디자인 안내·광고·해시태그·목차·구분선·기존 마커 등을 제거하고
순수 본문만 output/body_text.txt 로 추출한다.
제거된 줄은 사유 코드와 함께 output/excluded_log.txt 에 기록한다.
"""
import io
import re
import sys
from pathlib import Path
from collections import Counter
from docx import Document

from utils import count_body_chars

# Windows 콘솔에서도 한글·유니코드 출력이 깨지지 않게 stdout을 UTF-8로 고정
try:
    sys.stdout.reconfigure(encoding="utf-8")
except Exception:
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")


# 사유 코드 (중복 없이 한 번씩만 — dict.fromkeys 로 순서 유지 + 자동 중복 제거)
REASONS = list(dict.fromkeys([
    "메타헤더", "디자인지시", "광고안내", "광고링크",
    "해시태그", "제목반복", "제목라인", "목차", "구분선", "기존마커",
]))


# ---------- docx → 줄 리스트 ----------

def _run_to_md(run):
    """run의 bold(굵게)를 ** ** 마크업으로 변환. 양 끝 공백은 마크업 밖으로 뺀다."""
    text = run.text
    if not text:
        return ""
    if not run.bold:
        return text
    stripped = text.strip()
    if not stripped:
        return text
    lead = text[: len(text) - len(text.lstrip())]
    trail = text[len(text.rstrip()):]
    return f"{lead}**{stripped}**{trail}"


def _para_to_md(para):
    return "".join(_run_to_md(r) for r in para.runs)


def extract_lines(docx_path: Path):
    """docx에서 줄 단위 텍스트 리스트 반환. 표는 파이프로 둘러싼 한 줄로 변환."""
    doc = Document(str(docx_path))
    lines = []
    for para in doc.paragraphs:
        text = _para_to_md(para)
        for line in text.split("\n"):
            lines.append(line)
    for table in doc.tables:
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_text = " ".join(_para_to_md(p) for p in cell.paragraphs).strip()
                cells.append(cell_text)
            lines.append("| " + " | ".join(cells) + " |")
    return lines


# ---------- 줄 분류 규칙 ----------

# 디자인지시: ** 로 감싸졌든 안 감싸졌든 'ㄴ' 으로 시작
RE_DESIGN = re.compile(r"^(\*\*)?ㄴ(\*\*)?(\s|$)")

# 기존 마커: 단순 숫자(0~999), bold 마크업 허용 (예: 0, **2**, **11**, 01)
RE_LEFT_MARKER = re.compile(r"^(\*\*)?\d{1,3}(\*\*)?$")

# 구분선: --- 또는 *** (3개 이상)
RE_DIVIDER = re.compile(r"^[-*]{3,}$")

# 해시태그: #한글/영문 으로 시작
RE_HASHTAG = re.compile(r"^#\S")

# 광고링크: 네이버 mkt 쇼핑 링크
RE_AD_LINK = re.compile(r"^https?://mkt\.shopping\.naver\.com/", re.I)

# 광고안내
RE_AD_NOTICE = re.compile(r"^(광고\s*이미지\s*번호\s*[:：]|광고\s*링크\s*[:：])")

# 메타헤더: 블로거 요청사항 또는 표 셀 안에 그것이 들어간 줄
RE_BLOGGER_REQ = re.compile(r"★\s*블로거\s*요청사항\s*★")

# 목차 블록 시작
RE_TOC_HEADER = re.compile(r"^<\s*목\s*차\s*>")
# 목차 항목: "1- ...", "1. ..." 등 - 블록 안에서만 적용
RE_TOC_ITEM = re.compile(r"^\d+[\-\.\)]\s")

# 제목 라인: "제목 : ..." 또는 "제목: ..." (docx 메타 정보, 본문 흐름 아님)
RE_TITLE_LINE = re.compile(r"^제목\s*[:：]")


def classify(lines):
    """
    각 줄을 (kept: bool, reason: str | None) 로 분류.
    상태가 필요한 패턴(블로거 요청사항 블록, 목차 블록, 제목반복)은 순회하면서 처리.
    """
    n = len(lines)
    result = [None] * n  # (kept, reason)

    in_blogger_meta = False  # 블로거 요청사항 블록 안인지
    in_toc = False           # <목차> 블록 안인지

    for i, raw in enumerate(lines):
        s = raw.strip()

        # 1) 빈 줄: 본문 (줄바꿈 보존)
        if s == "":
            in_blogger_meta = False  # 블록은 빈 줄에서 종료
            in_toc = False
            result[i] = (True, None)
            continue

        # 2) 블로거 요청사항 블록 진입/유지
        if RE_BLOGGER_REQ.search(s):
            in_blogger_meta = True
            result[i] = (False, "메타헤더")
            continue
        if in_blogger_meta:
            # 블록 안에서 "- ..." 안내 줄이면 메타헤더, 아니면 블록 종료 후 재분류
            if s.startswith("- "):
                result[i] = (False, "메타헤더")
                continue
            else:
                in_blogger_meta = False
                # 아래 일반 분류로 흘려보냄

        # 3) 목차 블록
        if RE_TOC_HEADER.match(s):
            in_toc = True
            result[i] = (False, "목차")
            continue
        if in_toc:
            if RE_TOC_ITEM.match(s):
                result[i] = (False, "목차")
                continue
            else:
                in_toc = False
                # 일반 분류로 진행

        # 4) 구분선
        if RE_DIVIDER.match(s):
            result[i] = (False, "구분선")
            continue

        # 4-1) 제목 라인 (docx 메타 — "제목 : ..." / "제목: ...")
        if RE_TITLE_LINE.match(s):
            result[i] = (False, "제목라인")
            continue

        # 5) 기존 마커 (단순 숫자만)
        if RE_LEFT_MARKER.match(s):
            result[i] = (False, "기존마커")
            continue

        # 6) 광고링크
        if RE_AD_LINK.match(s):
            result[i] = (False, "광고링크")
            continue

        # 7) 광고안내
        if RE_AD_NOTICE.match(s):
            result[i] = (False, "광고안내")
            continue

        # 8) 해시태그
        if RE_HASHTAG.match(s):
            result[i] = (False, "해시태그")
            continue

        # 9) 디자인지시
        if RE_DESIGN.match(s):
            result[i] = (False, "디자인지시")
            continue

        # 10) 표 셀 안의 메타헤더 (한 줄 전체가 |...| 이고 안에 ★블로거 요청사항★)
        if s.startswith("|") and s.endswith("|") and RE_BLOGGER_REQ.search(s):
            result[i] = (False, "메타헤더")
            continue

        # 그 외: 일단 본문으로 둠 - 제목반복은 다음 패스에서 처리
        result[i] = (True, None)

    # 11) 제목반복 - 같은 비-빈 줄이 연속해서 2회 이상 나오면 두 번째부터 제거
    prev_text = None
    for i, raw in enumerate(lines):
        kept, reason = result[i]
        if not kept:
            prev_text = None
            continue
        s = raw.strip()
        if s == "":
            prev_text = None
            continue
        if prev_text is not None and s == prev_text:
            result[i] = (False, "제목반복")
        else:
            prev_text = s

    return result


# ---------- 의심 케이스 탐지 ----------

# 본문에 보존했지만 사람 검토가 필요한 케이스를 감지
def detect_suspects(lines, classified):
    suspects = []
    for i, raw in enumerate(lines):
        kept, reason = classified[i]
        s = raw.strip()
        if not kept or s == "":
            continue
        # 본문 안의 일반 https 링크 (mkt 쇼핑 외) - 광고링크 패턴엔 없음
        if re.search(r"https?://", s) and not RE_AD_LINK.match(s):
            suspects.append((i + 1, raw, "본문 보존",
                             "mkt.shopping.naver.com 외 링크 - 광고링크 패턴 밖이라 보존"))
    return suspects


# ---------- 보고용 ----------

def main(docx_path: Path, out_dir: Path):
    out_dir.mkdir(parents=True, exist_ok=True)
    lines = extract_lines(docx_path)
    classified = classify(lines)

    body_lines = []
    excluded_records = []
    reason_counter = Counter({r: 0 for r in REASONS})

    for i, raw in enumerate(lines):
        kept, reason = classified[i]
        if kept:
            body_lines.append(raw)
        else:
            excluded_records.append((i + 1, reason, raw))
            reason_counter[reason] += 1

    # body_text.txt: 본문 줄바꿈 보존
    (out_dir / "body_text.txt").write_text(
        "\n".join(body_lines) + "\n", encoding="utf-8")

    # excluded_log.txt
    log_lines = ["[줄번호] [사유코드] 원본내용", "-" * 60]
    for ln, reason, raw in excluded_records:
        log_lines.append(f"[{ln:3d}] [{reason}] {raw}")
    (out_dir / "excluded_log.txt").write_text(
        "\n".join(log_lines) + "\n", encoding="utf-8")

    # 보고
    body_text = "\n".join(body_lines)
    body_chars = count_body_chars(body_text)
    suspects = detect_suspects(lines, classified)

    print(f"# STEP 1 보고")
    print(f"샘플: {docx_path.name}")
    print()
    print(f"## 1. 줄 수 통계")
    print(f"- 원본 줄 수: {len(lines)}")
    print(f"- 본문 줄 수: {len(body_lines)}")
    print(f"- 제외된 줄 수: {len(excluded_records)}")
    print()
    print(f"## 2. 제외 사유별 카운트")
    seen = set()
    for r in REASONS:
        if r in seen:
            continue
        seen.add(r)
        print(f"- {r}: {reason_counter[r]}")
    print()
    print(f"## 3. 본문 추정 글자수 ([가-힣A-Za-z0-9] 카운트)")
    print(f"- {body_chars}")
    print()
    print(f"## 4. 의심 케이스 ({len(suspects)}건)")
    if not suspects:
        print("- 없음")
    else:
        for ln, raw, decision, why in suspects:
            print(f"- [줄 {ln}] {raw!r} → {decision}: {why}")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        # 샘플 미지정: samples/ 첫 파일
        samples_dir = Path("samples")
        candidates = sorted(samples_dir.glob("*.docx"))
        if not candidates:
            print("samples/ 폴더에 docx 파일이 없습니다.", file=sys.stderr)
            sys.exit(1)
        docx_path = candidates[0]
        print(f"(샘플 미지정 - samples/ 첫 파일 사용: {docx_path.name})\n")
    else:
        docx_path = Path(sys.argv[1])
    main(docx_path, Path("output"))
