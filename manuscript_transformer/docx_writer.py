"""
원고 치환기 — Word 출력 모듈
원고제작기의 save_as_docx 로직을 모듈화하여 재사용
"""
import re

# ── 색상명 → RGBColor 매핑 ──
def _get_color_name_to_rgb():
    from docx.shared import RGBColor
    return {
        '빨간색': RGBColor(0xFF, 0x00, 0x00),
        '파란색': RGBColor(0x00, 0x70, 0xC0),
        '청록색': RGBColor(0x00, 0x80, 0x80),
        '초록색': RGBColor(0x00, 0x80, 0x00),
        '보라색': RGBColor(0x70, 0x30, 0xA0),
        '주황색': RGBColor(0xED, 0x7D, 0x31),
        '회색': RGBColor(0x80, 0x80, 0x80),
        '많이옅은회색': RGBColor(0xC0, 0xC0, 0xC0),
        '옅은회색': RGBColor(0xA0, 0xA0, 0xA0),
        '진한회색': RGBColor(0x50, 0x50, 0x50),
    }


# ── ㄴ 서식 지시 파싱 ──
def parse_annotation(annotation_text):
    """ㄴ 서식 지시 줄 → 서식 딕셔너리"""
    from docx.enum.text import WD_COLOR_INDEX
    text = annotation_text.lstrip('ㄴ').strip()
    fmt = {
        'font_size': None, 'bold': False, 'italic': False,
        'colored_words': [], 'full_text_color': None,
        'highlight': None, 'quote': None, 'link': False,
        'multi_line': 1, 'is_image_desc': False,
    }

    if text.startswith('(') and text.endswith(')'):
        fmt['is_image_desc'] = True
        return fmt

    m = re.search(r'인용구\s*(\d+)\s*번', text)
    if m:
        fmt['quote'] = int(m.group(1))

    VALID_FONT_SIZES = [11, 13, 15, 16, 19, 24, 28]
    m = re.search(r'글자\s*크기\s*(\d+)', text)
    if m:
        requested = int(m.group(1))
        fmt['font_size'] = min(VALID_FONT_SIZES, key=lambda x: abs(x - requested))

    if re.search(r'글꼴\s*두껍게|두껍게', text):
        fmt['bold'] = True

    if re.search(r'이탤릭|기울임|글꼴\s*기울임', text):
        fmt['italic'] = True

    full_color_map = {
        '옅은 회색': '옅은회색', '많이 옅은 회색': '많이옅은회색',
        '회색': '회색', '진한 회색': '진한회색',
        '빨간색': '빨간색', '파란색': '파란색', '청록색': '청록색',
        '초록색': '초록색', '보라색': '보라색', '주황색': '주황색',
    }
    for pattern, color_key in full_color_map.items():
        if re.search(rf'글자\s*색\s*{re.escape(pattern)}', text):
            fmt['full_text_color'] = color_key
            break

    color_names = ['빨간색', '파란색', '청록색', '초록색', '보라색', '주황색', '회색']
    for color_name in color_names:
        m = re.search(rf"((?:'[^']+'\s*,?\s*)+)\s*{color_name}", text)
        if m:
            words = re.findall(r"'([^']+)'", m.group(1))
            for w in words:
                fmt['colored_words'].append((w, color_name))

    highlight_map = {
        '노란|노랑': WD_COLOR_INDEX.YELLOW,
        '검정|검은': WD_COLOR_INDEX.BLACK,
        '파란|파랑': WD_COLOR_INDEX.BLUE,
        '빨간|빨강': WD_COLOR_INDEX.RED,
        '초록': WD_COLOR_INDEX.GREEN,
        '청록': WD_COLOR_INDEX.TEAL,
    }
    for hl_pattern, hl_val in highlight_map.items():
        if re.search(rf'(?:{hl_pattern})색?\s*형광펜', text):
            fmt['highlight'] = hl_val
            break

    if '링크도구로연결' in text:
        fmt['link'] = True

    num_map = {'두': 2, '세': 3, '네': 4, '다섯': 5}
    m = re.search(r'(두|세|네|다섯)\s*줄\s*모두', text)
    if m:
        fmt['multi_line'] = num_map.get(m.group(1), 1)

    return fmt


def _is_format_annotation(text):
    """ㄴ로 시작하는 줄이 서식 지시인지 판별"""
    stripped = text.lstrip('ㄴ').strip()
    if stripped.startswith('(') and stripped.endswith(')'):
        return True
    if re.search(r'글자\s*크기|글꼴\s*두껍게|두껍게|형광펜|인용구|이탤릭|기울임|링크도구|줄\s*모두|글자\s*색', stripped):
        return True
    if re.search(r"'[^']+'\s*(빨간색|파란색|청록색|초록색|보라색|주황색|회색)", stripped):
        return True
    return False


# ── 마크다운 + 색상 세그먼트 빌드 ──
def _build_styled_segments(original_text, colored_words):
    md_re = re.compile(r'\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*')
    chunks = []
    last = 0
    for m in md_re.finditer(original_text):
        if m.start() > last:
            chunks.append((original_text[last:m.start()], False, False))
        if m.group(1):
            chunks.append((m.group(1), True, True))
        elif m.group(2):
            chunks.append((m.group(2), True, False))
        elif m.group(3):
            chunks.append((m.group(3), False, True))
        last = m.end()
    if last < len(original_text):
        chunks.append((original_text[last:], False, False))

    if not colored_words:
        return [(t, {'bold': b, 'italic': it, 'color': None}) for t, b, it in chunks]

    visible_text = ''.join(c[0] for c in chunks)
    char_colors = [None] * len(visible_text)
    for word, color_name in colored_words:
        for m in re.finditer(re.escape(word), visible_text):
            for j in range(m.start(), m.end()):
                char_colors[j] = color_name

    segments = []
    pos = 0
    for chunk_text, is_bold, is_italic in chunks:
        i = 0
        while i < len(chunk_text):
            current_color = char_colors[pos + i]
            j = i
            while j < len(chunk_text) and char_colors[pos + j] == current_color:
                j += 1
            segments.append((chunk_text[i:j], {'bold': is_bold, 'italic': is_italic, 'color': current_color}))
            i = j
        pos += len(chunk_text)
    return segments


def _clear_paragraph_runs(para):
    from docx.oxml.ns import qn
    p_elem = para._element
    for r in list(p_elem.findall(qn('w:r'))):
        p_elem.remove(r)


def _apply_quote_border(paragraph, quote_num):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    color_map = {
        1: '4472C4', 2: '70AD47', 3: 'ED7D31',
        4: 'FFC000', 5: '5B9BD5', 6: '7030A0',
    }
    border_color = color_map.get(quote_num, '4472C4')
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '18')
    left.set(qn('w:space'), '8')
    left.set(qn('w:color'), border_color)
    pBdr.append(left)
    pPr.append(pBdr)


def _apply_formatting_to_para(para, original_text, fmt):
    from docx.shared import Pt
    BLUE_C = _get_color_name_to_rgb()['파란색']
    _clear_paragraph_runs(para)
    segments = _build_styled_segments(original_text, fmt.get('colored_words', []))
    is_quote = bool(fmt.get('quote'))

    for seg_text, seg_props in segments:
        run = para.add_run(seg_text)
        if fmt.get('font_size'):
            run.font.size = Pt(fmt['font_size'])
        if is_quote:
            run.bold = True
        elif seg_props.get('bold'):
            run.bold = True
        elif fmt.get('bold') and seg_props.get('color'):
            run.bold = True
        elif fmt.get('bold') and not fmt.get('colored_words'):
            run.bold = True
        if fmt.get('italic') or seg_props.get('italic'):
            run.italic = True
        if not is_quote:
            _color_map = _get_color_name_to_rgb()
            color_name = seg_props.get('color')
            if color_name and color_name in _color_map:
                run.font.color.rgb = _color_map[color_name]
            elif fmt.get('full_text_color'):
                ftc = fmt['full_text_color']
                if ftc in _color_map:
                    run.font.color.rgb = _color_map[ftc]
            if fmt.get('highlight'):
                run.font.highlight_color = fmt['highlight']

    if fmt.get('link'):
        for run in para.runs:
            run.font.color.rgb = BLUE_C
            run.underline = True


def _add_text_runs(para, text):
    md_re = re.compile(r'\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*')
    last = 0
    for m in md_re.finditer(text):
        if m.start() > last:
            para.add_run(text[last:m.start()])
        if m.group(1):
            run = para.add_run(m.group(1))
            run.bold = True
            run.italic = True
        elif m.group(2):
            run = para.add_run(m.group(2))
            run.bold = True
        elif m.group(3):
            run = para.add_run(m.group(3))
            run.italic = True
        last = m.end()
    if last < len(text):
        para.add_run(text[last:])


def _split_colored_words_across_targets(targets, colored_words):
    if not colored_words or len(targets) <= 1:
        return None
    spanning = [w for w, _ in colored_words if not any(w in t for _, t in targets)]
    if not spanning:
        return None

    all_texts = [t for _, t in targets]
    joined = ' '.join(all_texts)
    char_colors = [None] * len(joined)
    for word, color_name in colored_words:
        for m in re.finditer(re.escape(word), joined):
            for j in range(m.start(), m.end()):
                char_colors[j] = color_name
    result = {}
    pos = 0
    for idx, (_, para_text) in enumerate(targets):
        para_cw = []
        pc = char_colors[pos:pos + len(para_text)]
        i = 0
        while i < len(para_text):
            if pc[i] is not None:
                color = pc[i]
                j = i
                while j < len(para_text) and pc[j] == color:
                    j += 1
                para_cw.append((para_text[i:j], color))
                i = j
            else:
                i += 1
        result[idx] = para_cw
        pos += len(para_text) + 1
    return result


def _add_blogger_request_box(doc, lines):
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)

    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'FFF8E1')
    shading.set(qn('w:val'), 'clear')
    cell._element.get_or_add_tcPr().append(shading)

    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '12')
        border.set(qn('w:color'), 'FF0000')
        border.set(qn('w:space'), '0')
        tcBorders.append(border)
    cell._element.get_or_add_tcPr().append(tcBorders)

    first_para = cell.paragraphs[0]
    for i, line in enumerate(lines):
        para = cell.add_paragraph() if i > 0 else first_para
        run = para.add_run(line)
        run.bold = True
        run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
        run.font.size = Pt(11)
    doc.add_paragraph('')


# ── 메인: 텍스트 → Word 파일 ──
def save_as_docx(text, filepath):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '맑은 고딕'
    style.font.size = Pt(11)

    annotation_re = re.compile(r'^ㄴ\s*')
    image_num_re = re.compile(r'^\d{1,2}$')
    GREEN = RGBColor(0x00, 0x80, 0x00)
    BLUE = RGBColor(0x00, 0x70, 0xC0)

    lines = text.split('\n')
    recent = []
    pending_fmts = []
    blogger_req_lines = []
    in_blogger_req = False

    for line in lines:
        stripped = line.strip()

        # ★ 블로거 요청사항
        if '★' in stripped:
            in_blogger_req = True
            blogger_req_lines.append(stripped)
            continue
        if in_blogger_req:
            if not stripped:
                _add_blogger_request_box(doc, blogger_req_lines)
                blogger_req_lines = []
                in_blogger_req = False
                recent.append((doc.paragraphs[-1] if doc.paragraphs else None, ''))
                continue
            else:
                blogger_req_lines.append(stripped)
                continue

        # 빈 줄
        if not stripped:
            p = doc.add_paragraph('')
            recent.append((p, ''))
            continue

        # 본문 중간 ㄴ 서식
        mid_ann = re.match(r'^(.+?)\s+ㄴ\s+(.+)$', stripped)
        if mid_ann and _is_format_annotation('ㄴ ' + mid_ann.group(2)):
            content_part = mid_ann.group(1).strip()
            ann_part = 'ㄴ ' + mid_ann.group(2).strip()
            p = doc.add_paragraph()
            _add_text_runs(p, content_part)
            recent.append((p, content_part))
            mid_fmt = parse_annotation(ann_part)
            if mid_fmt.get('colored_words'):
                if all(w in content_part for w, _ in mid_fmt['colored_words']):
                    _apply_formatting_to_para(p, content_part, mid_fmt)
                else:
                    pending_fmts.append((mid_fmt, []))
            elif not mid_fmt['is_image_desc']:
                _apply_formatting_to_para(p, content_part, mid_fmt)
            ap = doc.add_paragraph()
            run = ap.add_run(ann_part)
            run.bold = True
            run.font.color.rgb = GREEN
            run.font.size = Pt(24)
            run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
            recent.append((ap, ann_part))
            continue

        # ㄴ 서식 지시 줄
        if annotation_re.match(stripped) and _is_format_annotation(stripped):
            fmt = parse_annotation(stripped)
            if fmt['is_image_desc']:
                continue

            content_paras = [(p, t) for p, t in recent
                             if t.strip() and not (re.match(r'^ㄴ\s*', t.strip()) and _is_format_annotation(t.strip()))]
            target_count = fmt['multi_line']
            targets = content_paras[-target_count:] if content_paras else []

            applied = False
            if fmt.get('colored_words'):
                if targets:
                    all_target_text = ' '.join(t for _, t in targets)
                    missing = any(w not in all_target_text for w, _ in fmt['colored_words'])
                    if missing and len(content_paras) > target_count:
                        found = False
                        for ext in range(target_count + 1, min(target_count + 5, len(content_paras) + 1)):
                            targets = content_paras[-ext:]
                            all_target_text = ' '.join(t for _, t in targets)
                            if all(w in all_target_text for w, _ in fmt['colored_words']):
                                found = True
                                break
                        if found:
                            applied = True
                        else:
                            pending_fmts.append((fmt, []))
                    elif not missing:
                        applied = True
                    else:
                        pending_fmts.append((fmt, []))
                else:
                    pending_fmts.append((fmt, []))
            else:
                applied = True

            if applied:
                per_para_cw = _split_colored_words_across_targets(targets, fmt.get('colored_words', []))
                for idx, (para, para_text) in enumerate(targets):
                    if per_para_cw and idx in per_para_cw:
                        para_fmt = dict(fmt)
                        para_fmt['colored_words'] = per_para_cw[idx]
                        _apply_formatting_to_para(para, para_text, para_fmt)
                    else:
                        _apply_formatting_to_para(para, para_text, fmt)

            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.bold = True
            run.font.color.rgb = GREEN
            run.font.size = Pt(24)
            run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
            recent.append((p, stripped))
            continue

        # 이미지 번호
        if image_num_re.match(stripped):
            display_num = str(int(stripped))
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(display_num)
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = BLUE
            recent.append((p, display_num))
            continue

        # 마크다운 헤딩
        if stripped.startswith('### '):
            p = doc.add_heading(stripped.lstrip('# ').strip(), level=3)
            recent.append((p, stripped))
            continue
        if stripped.startswith('## '):
            p = doc.add_heading(stripped.lstrip('# ').strip(), level=2)
            recent.append((p, stripped))
            continue
        if stripped.startswith('# '):
            p = doc.add_heading(stripped.lstrip('# ').strip(), level=1)
            recent.append((p, stripped))
            continue

        # 일반 텍스트
        p = doc.add_paragraph()
        _add_text_runs(p, stripped)
        recent.append((p, stripped))

        # 대기 중인 서식 적용
        if pending_fmts:
            new_pending = []
            for pfmt, collected in pending_fmts:
                collected.append((p, stripped))
                if pfmt.get('colored_words'):
                    all_text = ' '.join(t for _, t in collected)
                    if all(w in all_text for w, _ in pfmt['colored_words']):
                        per_para_cw = _split_colored_words_across_targets(collected, pfmt.get('colored_words', []))
                        for cidx, (cp, ct) in enumerate(collected):
                            if per_para_cw and cidx in per_para_cw:
                                p_fmt = dict(pfmt)
                                p_fmt['colored_words'] = per_para_cw[cidx]
                                _apply_formatting_to_para(cp, ct, p_fmt)
                            else:
                                _apply_formatting_to_para(cp, ct, pfmt)
                    elif len(collected) < 5:
                        new_pending.append((pfmt, collected))
                else:
                    _apply_formatting_to_para(p, stripped, pfmt)
            pending_fmts = new_pending

        if len(recent) > 15:
            recent = recent[-15:]

    if blogger_req_lines:
        _add_blogger_request_box(doc, blogger_req_lines)

    doc.save(filepath)
