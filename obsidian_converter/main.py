"""
옵시디언 MD 변환기 v2.1
- 로컬 파일/폴더를 선택하여 Obsidian 볼트에 마크다운으로 저장
- 미리보기 + 속성 편집 + [[]] 자동 링크
- AI 모드 (Claude API) / 일반 모드 (키워드 매칭) 선택 가능
- 지원: .txt, .py, .js, .json, .yaml, .toml, .xml, .html, .css,
         .bat, .sh, .sql, .csv, .xlsx, .docx, .pdf, .md, .log, .ini, .cfg
"""

import os
import re
import csv
import json
import threading
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from pathlib import Path
from datetime import datetime
from dataclasses import dataclass, field
from collections import Counter

# ── 기본 설정 ──────────────────────────────────────────────
DEFAULT_VAULT = Path.home() / "Desktop" / "옵시디언" / "효진의 창고"
WINDOW_TITLE = "옵시디언 MD 변환기 v2.1"
WINDOW_SIZE = "900x700"
BASE_DIR = Path(__file__).parent
API_KEY_FILE = BASE_DIR / ".api_key"

# 지원하는 파일 확장자
CODE_EXTENSIONS = {
    ".py": "python", ".js": "javascript", ".ts": "typescript",
    ".jsx": "jsx", ".tsx": "tsx", ".java": "java", ".c": "c",
    ".cpp": "cpp", ".cs": "csharp", ".go": "go", ".rs": "rust",
    ".rb": "ruby", ".php": "php", ".swift": "swift",
    ".kt": "kotlin", ".r": "r", ".m": "matlab",
    ".sh": "bash", ".bat": "batch", ".ps1": "powershell",
    ".sql": "sql", ".html": "html", ".css": "css",
    ".xml": "xml", ".yaml": "yaml", ".yml": "yaml",
    ".toml": "toml", ".ini": "ini", ".cfg": "ini",
}

TEXT_EXTENSIONS = {".txt", ".md", ".log", ".rst"}
DOC_EXTENSIONS = {".docx", ".xlsx", ".xls", ".csv", ".pdf", ".json"}
ALL_SUPPORTED = set(CODE_EXTENSIONS.keys()) | TEXT_EXTENSIONS | DOC_EXTENSIONS

FRONTMATTER_FIELDS = [
    ("tags", "태그 (쉼표 구분)", "list"),
    ("aliases", "별칭 (쉼표 구분)", "list"),
    ("date", "날짜", "text"),
    ("version", "버전", "text"),
    ("type", "유형", "text"),
    ("description", "설명", "text"),
    ("status", "상태", "text"),
]


# ── API 키 관리 ───────────────────────────────────────────
def load_api_key() -> str:
    if API_KEY_FILE.exists():
        return API_KEY_FILE.read_text(encoding="utf-8").strip()
    return ""


def save_api_key(key: str):
    API_KEY_FILE.write_text(key.strip(), encoding="utf-8")


# ── 데이터 클래스 ─────────────────────────────────────────
@dataclass
class ConvertedFile:
    source_path: Path
    filename: str
    frontmatter: dict = field(default_factory=dict)
    body: str = ""


@dataclass
class VaultContext:
    """볼트에서 수집한 컨텍스트 정보"""
    notes: list[str] = field(default_factory=list)           # 노트 이름 목록
    all_tags: list[str] = field(default_factory=list)        # 사용 중인 태그
    all_types: list[str] = field(default_factory=list)       # 사용 중인 type 값
    all_statuses: list[str] = field(default_factory=list)    # 사용 중인 status 값
    folders: list[str] = field(default_factory=list)         # 폴더 목록
    folder_descriptions: dict = field(default_factory=dict)  # 폴더별 노트 이름 샘플


# ── 볼트 스캔 ─────────────────────────────────────────────
def scan_vault(vault_path: Path) -> VaultContext:
    """볼트 전체 스캔: 노트 이름, 태그, 폴더 구조 수집"""
    ctx = VaultContext()
    tags_counter = Counter()
    types_set = set()
    statuses_set = set()

    # 폴더 목록
    try:
        for p in sorted(vault_path.iterdir()):
            if p.is_dir() and not p.name.startswith("."):
                ctx.folders.append(p.name)
                # 폴더별 노트 샘플 (최대 5개)
                samples = []
                for f in p.rglob("*.md"):
                    samples.append(f.stem)
                    if len(samples) >= 5:
                        break
                ctx.folder_descriptions[p.name] = samples
    except Exception:
        pass

    # 노트 + 프론트매터 수집
    try:
        for f in vault_path.rglob("*.md"):
            name = f.stem
            if len(name) >= 2:
                ctx.notes.append(name)

            # 프론트매터 파싱 (간단 YAML)
            try:
                text = f.read_text(encoding="utf-8", errors="replace")
                if text.startswith("---"):
                    end = text.find("---", 3)
                    if end > 0:
                        fm_text = text[3:end]
                        for line in fm_text.split("\n"):
                            line = line.strip()
                            if line.startswith("- "):
                                # 리스트 아이템 (태그 등)
                                tag = line[2:].strip().strip('"').strip("'")
                                if tag:
                                    tags_counter[tag] += 1
                            elif line.startswith("type:"):
                                val = line.split(":", 1)[1].strip().strip('"').strip("'")
                                if val:
                                    types_set.add(val)
                            elif line.startswith("status:"):
                                val = line.split(":", 1)[1].strip().strip('"').strip("'")
                                if val:
                                    statuses_set.add(val)
            except Exception:
                pass
    except Exception:
        pass

    ctx.notes = sorted(set(ctx.notes), key=len, reverse=True)
    ctx.all_tags = [t for t, _ in tags_counter.most_common(100)]
    ctx.all_types = sorted(types_set)
    ctx.all_statuses = sorted(statuses_set)
    return ctx


# ── 키워드 기반 분석 (일반 모드) ──────────────────────────
def analyze_keyword(content: str, filename: str, ext: str,
                    vault_ctx: VaultContext) -> dict:
    """키워드 매칭으로 태그, 폴더, 설명 추천"""
    result = {
        "tags": "",
        "aliases": "",
        "date": datetime.now().strftime("%Y-%m-%d"),
        "version": "",
        "type": _guess_type(ext),
        "description": "",
        "status": "draft",
        "suggested_folder": "",
    }

    content_lower = content.lower()

    # 태그 매칭: 볼트에 이미 쓰이는 태그 중 본문에 등장하는 것
    matched_tags = []
    for tag in vault_ctx.all_tags:
        if tag.lower() in content_lower or tag.lower() in filename.lower():
            matched_tags.append(tag)
        if len(matched_tags) >= 5:
            break

    # 확장자 기반 태그 추가
    if ext in CODE_EXTENSIONS:
        lang = CODE_EXTENSIONS[ext]
        if lang not in [t.lower() for t in matched_tags]:
            matched_tags.append(lang.capitalize())

    result["tags"] = ", ".join(matched_tags)

    # 설명: 첫 번째 의미있는 줄
    for line in content.split("\n"):
        line = line.strip().strip("#").strip("*").strip()
        if len(line) >= 5 and not line.startswith("```") and not line.startswith("---"):
            result["description"] = line[:80]
            break

    # 폴더 추천: 키워드 기반
    best_folder = ""
    best_score = 0
    for folder in vault_ctx.folders:
        score = 0
        folder_lower = folder.lower()
        # 폴더 이름이 본문/파일명에 있으면 점수
        if folder_lower in content_lower or folder_lower in filename.lower():
            score += 3
        # 폴더 내 노트 이름이 본문에 있으면 점수
        for sample in vault_ctx.folder_descriptions.get(folder, []):
            if sample.lower() in content_lower:
                score += 1
        if score > best_score:
            best_score = score
            best_folder = folder

    if best_folder:
        result["suggested_folder"] = best_folder

    return result


# ── Claude API 분석 (AI 모드) ─────────────────────────────
def analyze_with_ai(content: str, filename: str, ext: str,
                    vault_ctx: VaultContext, api_key: str) -> dict:
    """Claude API로 태그, 폴더, 설명 등 자동 추천"""
    try:
        import anthropic
    except ImportError:
        raise ImportError("anthropic 패키지 미설치. `pip install anthropic` 실행 필요")

    client = anthropic.Anthropic(api_key=api_key)

    # 본문이 너무 길면 앞부분만
    content_preview = content[:3000] if len(content) > 3000 else content

    # 프롬프트 구성
    vault_info = f"""
## 볼트 기존 태그 (사용 빈도순)
{', '.join(vault_ctx.all_tags[:30]) if vault_ctx.all_tags else '(아직 없음)'}

## 볼트 기존 유형(type)
{', '.join(vault_ctx.all_types) if vault_ctx.all_types else '(아직 없음)'}

## 볼트 기존 상태(status)
{', '.join(vault_ctx.all_statuses) if vault_ctx.all_statuses else 'draft, active, archived'}

## 볼트 폴더 구조
"""
    for folder in vault_ctx.folders:
        samples = vault_ctx.folder_descriptions.get(folder, [])
        sample_str = ", ".join(samples[:3]) if samples else "(비어있음)"
        vault_info += f"- {folder}: {sample_str}\n"

    prompt = f"""다음 파일을 옵시디언 볼트에 저장하려 합니다. 속성과 저장 폴더를 추천해주세요.

## 파일 정보
- 파일명: {filename}
- 확장자: {ext}

## 파일 내용 (앞부분)
{content_preview}

## 볼트 현황
{vault_info}

## 요청사항
아래 JSON 형식으로만 답해주세요. 다른 설명 없이 JSON만:
{{
    "tags": "태그1, 태그2, 태그3",
    "aliases": "별칭1, 별칭2",
    "description": "한줄 요약 설명",
    "type": "유형 (기존 유형 우선 사용)",
    "status": "상태",
    "version": "버전 (해당시)",
    "suggested_folder": "추천 폴더명 (기존 폴더 중 선택, 없으면 새 폴더명)"
}}"""

    response = client.messages.create(
        model="claude-haiku-4-5-20251001",
        max_tokens=500,
        messages=[{"role": "user", "content": prompt}],
    )

    # JSON 파싱
    text = response.content[0].text.strip()
    # JSON 블록 추출
    if "```" in text:
        text = text.split("```")[1]
        if text.startswith("json"):
            text = text[4:]
        text = text.strip()

    data = json.loads(text)
    data["date"] = datetime.now().strftime("%Y-%m-%d")
    return data


# ── 링크 감지/적용 ────────────────────────────────────────
def find_linkable_terms(content: str, vault_notes: list[str]) -> list[tuple[str, int]]:
    results = []
    cleaned = re.sub(r'\[\[[^\]]*\]\]', lambda m: ' ' * len(m.group()), content)
    for note in vault_notes:
        if len(note) < 3:
            continue
        pattern = re.escape(note)
        matches = list(re.finditer(pattern, cleaned, re.IGNORECASE))
        if matches:
            results.append((note, len(matches)))
    return results


def apply_auto_links(content: str, note_names: list[str]) -> str:
    sorted_names = sorted(note_names, key=len, reverse=True)
    for name in sorted_names:
        pattern = re.escape(name)

        def replacer(m, _content=content):
            start = m.start()
            before = _content[:start]
            open_count = before.count('[[') - before.count(']]')
            if open_count > 0:
                return m.group()
            return f'[[{m.group()}]]'

        content = re.sub(pattern, replacer, content, flags=re.IGNORECASE)
    return content


# ── 프론트매터 직렬화 ─────────────────────────────────────
def serialize_frontmatter(fm: dict) -> str:
    lines = ["---"]
    for key, _, ftype in FRONTMATTER_FIELDS:
        val = fm.get(key, "")
        if not val:
            continue
        if ftype == "list":
            items = [v.strip() for v in val.split(",") if v.strip()] if isinstance(val, str) else val
            if items:
                lines.append(f"{key}:")
                for item in items:
                    lines.append(f"  - {item}")
        else:
            lines.append(f'{key}: "{val}"')
    lines.append("---")
    return "\n".join(lines) + "\n\n"


# ── 변환 함수들 ────────────────────────────────────────────
def convert_text_file(filepath: Path) -> str:
    text = filepath.read_text(encoding="utf-8", errors="replace").strip()
    if filepath.suffix.lower() == ".md":
        return text
    return text


def convert_code_file(filepath: Path) -> str:
    lang = CODE_EXTENSIONS.get(filepath.suffix.lower(), "")
    code = filepath.read_text(encoding="utf-8", errors="replace").strip()
    result = f"**파일:** `{filepath.name}`\n"
    result += f"**언어:** {lang.upper() if lang else '기타'}\n\n"
    result += f"```{lang}\n{code}\n```"
    return result


def convert_json_file(filepath: Path) -> str:
    raw = filepath.read_text(encoding="utf-8", errors="replace")
    try:
        data = json.loads(raw)
        pretty = json.dumps(data, indent=2, ensure_ascii=False)
        result = f"**파일:** `{filepath.name}`\n\n"
        if isinstance(data, dict) and len(data) <= 30:
            result += "## 구조 요약\n\n"
            for key in data:
                val = data[key]
                if isinstance(val, (list, dict)):
                    val_preview = f"{type(val).__name__} ({len(val)}개)"
                else:
                    val_preview = str(val)[:80]
                result += f"- **{key}**: {val_preview}\n"
            result += "\n"
        result += "## 전체 내용\n\n"
        result += f"```json\n{pretty}\n```"
        return result
    except json.JSONDecodeError:
        return f"```json\n{raw}\n```"


def convert_csv_file(filepath: Path) -> str:
    rows = []
    with open(filepath, "r", encoding="utf-8-sig", errors="replace") as f:
        reader = csv.reader(f)
        for row in reader:
            rows.append(row)
    if not rows:
        return "*빈 CSV 파일*"
    result = f"**파일:** `{filepath.name}`  \n"
    result += f"**행 수:** {len(rows) - 1}개 (헤더 제외)\n\n"
    header = rows[0]
    result += "| " + " | ".join(header) + " |\n"
    result += "| " + " | ".join(["---"] * len(header)) + " |\n"
    for row in rows[1:101]:
        padded = row + [""] * (len(header) - len(row))
        cells = [c.replace("|", "\\|").replace("\n", " ") for c in padded[:len(header)]]
        result += "| " + " | ".join(cells) + " |\n"
    if len(rows) - 1 > 100:
        result += f"\n> *... 외 {len(rows) - 101}행 생략*\n"
    return result


def convert_xlsx_file(filepath: Path) -> str:
    try:
        import openpyxl
    except ImportError:
        return "openpyxl 미설치. `pip install openpyxl` 실행 후 다시 시도하세요."
    wb = openpyxl.load_workbook(filepath, read_only=True, data_only=True)
    result = f"**파일:** `{filepath.name}`  \n"
    result += f"**시트 수:** {len(wb.sheetnames)}개\n\n"
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            result += f"## {sheet_name}\n\n*빈 시트*\n\n"
            continue
        result += f"## {sheet_name}\n\n"
        header = [str(c) if c is not None else "" for c in rows[0]]
        col_count = len(header)
        result += "| " + " | ".join(header) + " |\n"
        result += "| " + " | ".join(["---"] * col_count) + " |\n"
        for row in rows[1:101]:
            cells = [str(c).replace("|", "\\|").replace("\n", " ") if c is not None else ""
                     for c in row[:col_count]]
            cells += [""] * (col_count - len(cells))
            result += "| " + " | ".join(cells) + " |\n"
        if len(rows) - 1 > 100:
            result += f"\n> *... 외 {len(rows) - 101}행 생략*\n"
        result += "\n"
    wb.close()
    return result


def convert_docx_file(filepath: Path) -> str:
    try:
        from docx import Document
    except ImportError:
        return "python-docx 미설치. `pip install python-docx` 실행 후 다시 시도하세요."
    doc = Document(str(filepath))
    result = f"**파일:** `{filepath.name}`\n\n---\n\n"
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            result += "\n"
            continue
        style = para.style.name.lower() if para.style else ""
        if "heading 1" in style:
            result += f"# {text}\n\n"
        elif "heading 2" in style:
            result += f"## {text}\n\n"
        elif "heading 3" in style:
            result += f"### {text}\n\n"
        elif "heading" in style:
            result += f"#### {text}\n\n"
        elif "list" in style or "bullet" in style:
            result += f"- {text}\n"
        else:
            md_line = ""
            for run in para.runs:
                t = run.text
                if not t:
                    continue
                if run.bold and run.italic:
                    md_line += f"***{t}***"
                elif run.bold:
                    md_line += f"**{t}**"
                elif run.italic:
                    md_line += f"*{t}*"
                else:
                    md_line += t
            result += (md_line or text) + "\n\n"
    for table in doc.tables:
        result += "\n"
        for i, row in enumerate(table.rows):
            cells = [cell.text.replace("|", "\\|").replace("\n", " ") for cell in row.cells]
            result += "| " + " | ".join(cells) + " |\n"
            if i == 0:
                result += "| " + " | ".join(["---"] * len(cells)) + " |\n"
        result += "\n"
    return result.strip()


def convert_pdf_file(filepath: Path) -> str:
    try:
        import PyPDF2
    except ImportError:
        try:
            import pypdf as PyPDF2
        except ImportError:
            return "PyPDF2 미설치. `pip install PyPDF2` 실행 후 다시 시도하세요."
    result = f"**파일:** `{filepath.name}`\n\n"
    with open(filepath, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        total_pages = len(reader.pages)
        result += f"**페이지 수:** {total_pages}p\n\n---\n\n"
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text and text.strip():
                result += f"### {i + 1}페이지\n\n"
                result += text.strip() + "\n\n"
    return result


def convert_file(filepath: Path) -> str:
    ext = filepath.suffix.lower()
    if ext in CODE_EXTENSIONS:
        return convert_code_file(filepath)
    elif ext == ".json":
        return convert_json_file(filepath)
    elif ext == ".csv":
        return convert_csv_file(filepath)
    elif ext in {".xlsx", ".xls"}:
        return convert_xlsx_file(filepath)
    elif ext == ".docx":
        return convert_docx_file(filepath)
    elif ext == ".pdf":
        return convert_pdf_file(filepath)
    elif ext in TEXT_EXTENSIONS:
        return convert_text_file(filepath)
    else:
        try:
            return filepath.read_text(encoding="utf-8", errors="strict")
        except (UnicodeDecodeError, Exception):
            return f"지원하지 않는 파일 형식: `{ext}`"


def _guess_type(ext: str) -> str:
    if ext in CODE_EXTENSIONS:
        return "code"
    elif ext in {".docx", ".pdf", ".txt", ".md"}:
        return "document"
    elif ext in {".xlsx", ".xls", ".csv"}:
        return "data"
    elif ext == ".json":
        return "config"
    return "file"


# ── 미리보기 창 ───────────────────────────────────────────
class PreviewWindow:
    def __init__(self, parent: tk.Tk, files: list[ConvertedFile],
                 vault_path: Path, subfolder: str, vault_ctx: VaultContext,
                 api_key_var: tk.StringVar = None):
        self.parent = parent
        self.files = files
        self.vault_path = vault_path
        self.vault_ctx = vault_ctx
        self.current_index = 0

        self.win = tk.Toplevel(parent)
        self.win.title("미리보기 및 편집")
        self.win.geometry("1200x850")
        self.win.minsize(900, 600)
        self.win.grab_set()
        self.win.configure(bg="#F7F8FA")

        self.subfolder_var = tk.StringVar(value=subfolder)
        self.fm_vars: dict[str, tk.StringVar] = {}
        self.use_ai = tk.BooleanVar(value=False)
        # 메인 화면의 api_key_var를 공유
        self.api_key_var = api_key_var if api_key_var else tk.StringVar(value=load_api_key())

        self._build_ui()
        self._show_file(0)

    def _build_ui(self):
        # ── 상단: 네비게이션 + 모드 선택 ──
        nav_frame = ttk.Frame(self.win, padding=(16, 8))
        nav_frame.pack(fill="x")

        self.prev_btn = ttk.Button(nav_frame, text="< 이전", command=self._go_prev)
        self.prev_btn.pack(side="left")
        self.nav_label = ttk.Label(nav_frame, text="", style="Title.TLabel")
        self.nav_label.pack(side="left", padx=16)
        self.next_btn = ttk.Button(nav_frame, text="다음 >", command=self._go_next)
        self.next_btn.pack(side="left")

        self.filename_var = tk.StringVar()
        ttk.Label(nav_frame, text="파일명:").pack(side="left", padx=(32, 4))
        ttk.Entry(nav_frame, textvariable=self.filename_var, width=30,
                  font=("맑은 고딕", 10)).pack(side="left")

        # AI / 일반 모드 토글
        mode_frame = ttk.Frame(nav_frame)
        mode_frame.pack(side="right")
        ttk.Checkbutton(mode_frame, text="AI 모드 (Claude API)",
                         variable=self.use_ai,
                         command=self._on_mode_change).pack(side="left")
        self.ai_status = ttk.Label(mode_frame, text="", style="Subtitle.TLabel")
        self.ai_status.pack(side="left", padx=(8, 0))

        # ── 중앙: 좌(속성+링크) / 우(본문) ──
        body = ttk.Frame(self.win, padding=(16, 0))
        body.pack(fill="both", expand=True)

        paned = ttk.PanedWindow(body, orient="horizontal")
        paned.pack(fill="both", expand=True)

        # ── 좌측 패널 ──
        left = ttk.Frame(paned, padding=8)
        paned.add(left, weight=35)

        # 속성 섹션 헤더 + 자동 채우기 버튼
        prop_header = ttk.Frame(left)
        prop_header.pack(fill="x", pady=(0, 8))
        ttk.Label(prop_header, text="속성 (프론트매터)",
                  font=("맑은 고딕", 11, "bold")).pack(side="left")
        self.autofill_btn = ttk.Button(prop_header, text="자동 채우기",
                                        command=self._autofill_current)
        self.autofill_btn.pack(side="right")

        self.fm_vars = {}
        for key, label, ftype in FRONTMATTER_FIELDS:
            row = ttk.Frame(left)
            row.pack(fill="x", pady=2)
            ttk.Label(row, text=f"{label}:", width=14, anchor="e").pack(side="left")
            var = tk.StringVar()
            self.fm_vars[key] = var
            entry = ttk.Entry(row, textvariable=var, font=("맑은 고딕", 9))
            entry.pack(side="left", fill="x", expand=True, padx=(4, 0))

        # API 키 (AI 모드용, 접히는 프레임)
        self.api_frame = ttk.Frame(left)
        api_row = ttk.Frame(self.api_frame)
        api_row.pack(fill="x", pady=4)
        ttk.Label(api_row, text="API 키:", width=14, anchor="e").pack(side="left")
        self.api_entry = ttk.Entry(api_row, textvariable=self.api_key_var,
                                    font=("맑은 고딕", 9), show="*")
        self.api_entry.pack(side="left", fill="x", expand=True, padx=(4, 4))
        ttk.Button(api_row, text="저장", command=self._save_api_key).pack(side="left")
        # 처음에는 숨김 — _on_mode_change에서 pack/pack_forget

        # 구분선 (api_frame 바로 뒤에 올 수 있도록 참조 보관)
        self.left_separator = ttk.Separator(left, orient="horizontal")
        self.left_separator.pack(fill="x", pady=12)

        # 링크 섹션
        ttk.Label(left, text="[[ ]] 링크",
                  font=("맑은 고딕", 11, "bold")).pack(anchor="w", pady=(0, 8))

        link_btn_frame = ttk.Frame(left)
        link_btn_frame.pack(fill="x", pady=(0, 4))
        ttk.Button(link_btn_frame, text="자동 감지",
                   command=self._detect_links).pack(side="left", padx=(0, 4))
        ttk.Button(link_btn_frame, text="선택 적용",
                   command=self._apply_selected_links).pack(side="left", padx=(0, 4))
        ttk.Button(link_btn_frame, text="전체 적용",
                   command=self._apply_all_links).pack(side="left")

        self.link_status = ttk.Label(left,
                                      text=f"볼트 노트 {len(self.vault_ctx.notes)}개",
                                      style="Subtitle.TLabel")
        self.link_status.pack(anchor="w", pady=(4, 4))

        link_list_frame = ttk.Frame(left)
        link_list_frame.pack(fill="both", expand=True, pady=(0, 8))

        self.link_listbox = tk.Listbox(link_list_frame, selectmode="extended",
                                        font=("맑은 고딕", 9), bg="white",
                                        highlightthickness=1, highlightbackground="#E2E8F0")
        link_scroll = ttk.Scrollbar(link_list_frame, orient="vertical",
                                     command=self.link_listbox.yview)
        self.link_listbox.configure(yscrollcommand=link_scroll.set)
        self.link_listbox.pack(side="left", fill="both", expand=True)
        link_scroll.pack(side="right", fill="y")
        self.detected_links: list[tuple[str, int]] = []

        # 수동 링크 삽입
        ttk.Separator(left, orient="horizontal").pack(fill="x", pady=8)
        ttk.Label(left, text="수동 링크 삽입",
                  font=("맑은 고딕", 10, "bold")).pack(anchor="w", pady=(0, 4))

        search_frame = ttk.Frame(left)
        search_frame.pack(fill="x", pady=(0, 4))
        self.search_var = tk.StringVar()
        self.search_var.trace_add("write", self._filter_notes)
        ttk.Entry(search_frame, textvariable=self.search_var,
                  font=("맑은 고딕", 9)).pack(side="left", fill="x", expand=True, padx=(0, 4))
        ttk.Button(search_frame, text="삽입",
                   command=self._insert_manual_link).pack(side="left")

        self.suggest_listbox = tk.Listbox(left, height=5, font=("맑은 고딕", 9),
                                           bg="white", highlightthickness=1,
                                           highlightbackground="#E2E8F0")
        self.suggest_listbox.pack(fill="x", pady=(0, 4))

        # ── 우측 패널: 본문 편집 ──
        right = ttk.Frame(paned, padding=8)
        paned.add(right, weight=65)

        ttk.Label(right, text="본문 (마크다운)",
                  font=("맑은 고딕", 11, "bold")).pack(anchor="w", pady=(0, 4))

        text_frame = ttk.Frame(right)
        text_frame.pack(fill="both", expand=True)

        self.text_widget = tk.Text(text_frame, wrap="word", font=("Consolas", 10),
                                    bg="white", fg="#1E293B", insertbackground="#4A6CF7",
                                    highlightthickness=1, highlightbackground="#E2E8F0",
                                    undo=True, padx=8, pady=8)
        text_scroll = ttk.Scrollbar(text_frame, orient="vertical",
                                     command=self.text_widget.yview)
        self.text_widget.configure(yscrollcommand=text_scroll.set)
        self.text_widget.pack(side="left", fill="both", expand=True)
        text_scroll.pack(side="right", fill="y")

        # ── 하단: 폴더 확인 + 저장 ──
        bottom = ttk.Frame(self.win, padding=(16, 8))
        bottom.pack(fill="x")

        ttk.Label(bottom, text="저장 폴더:").pack(side="left")
        vault_folders = ["(볼트 루트)"] + self.vault_ctx.folders
        self.folder_combo = ttk.Combobox(bottom, textvariable=self.subfolder_var,
                                          values=vault_folders, width=40,
                                          font=("맑은 고딕", 9))
        self.folder_combo.pack(side="left", padx=(4, 4))
        ttk.Button(bottom, text="찾아보기",
                   command=self._browse_subfolder).pack(side="left")
        ttk.Label(bottom, text=f"(볼트: {self.vault_path.name})",
                  style="Subtitle.TLabel").pack(side="left", padx=(8, 0))

        ttk.Button(bottom, text="취소",
                   command=self._cancel).pack(side="right", padx=(8, 0))
        ttk.Button(bottom, text="전체 저장", style="Primary.TButton",
                   command=self._save_all).pack(side="right")
        ttk.Button(bottom, text="건너뛰기",
                   command=self._skip_current).pack(side="right", padx=(0, 8))

    # ── 모드 전환 ──
    def _on_mode_change(self):
        if self.use_ai.get():
            # 구분선 바로 앞에 삽입
            self.api_frame.pack(fill="x", before=self.left_separator)
            key = self.api_key_var.get().strip()
            if key:
                self.ai_status.config(text="AI 활성")
            else:
                self.ai_status.config(text="API 키 필요")
        else:
            self.api_frame.pack_forget()
            self.ai_status.config(text="")

    def _save_api_key(self):
        key = self.api_key_var.get().strip()
        if key:
            save_api_key(key)
            self.ai_status.config(text="API 키 저장됨")
        else:
            self.ai_status.config(text="키가 비어있습니다")

    # ── 자동 채우기 ──
    def _autofill_current(self):
        """현재 파일의 속성을 자동 분석하여 채움"""
        cf = self.files[self.current_index]
        content = self.text_widget.get("1.0", "end-1c")
        ext = cf.source_path.suffix.lower()

        if self.use_ai.get():
            key = self.api_key_var.get().strip()
            if not key:
                messagebox.showwarning("알림", "AI 모드를 사용하려면 API 키를 입력하세요.")
                return
            self.autofill_btn.config(state="disabled", text="분석 중...")

            def run_ai():
                try:
                    result = analyze_with_ai(
                        content, cf.filename, ext, self.vault_ctx, key
                    )

                    def apply():
                        self.autofill_btn.config(state="normal", text="자동 채우기")
                        for fkey, var in self.fm_vars.items():
                            val = result.get(fkey, "")
                            if val:
                                var.set(val)
                        # 폴더 추천
                        suggested = result.get("suggested_folder", "")
                        if suggested:
                            self.subfolder_var.set(suggested)
                        self.ai_status.config(text="AI 분석 완료")

                    self.win.after(0, apply)
                except Exception as e:
                    def on_error():
                        self.autofill_btn.config(state="normal", text="자동 채우기")
                        self.ai_status.config(text="AI 오류")
                        messagebox.showerror("AI 오류", str(e))
                    self.win.after(0, on_error)

            threading.Thread(target=run_ai, daemon=True).start()
        else:
            # 키워드 매칭 모드
            result = analyze_keyword(content, cf.filename, ext, self.vault_ctx)
            for fkey, var in self.fm_vars.items():
                val = result.get(fkey, "")
                if val:
                    var.set(val)
            suggested = result.get("suggested_folder", "")
            if suggested:
                self.subfolder_var.set(suggested)

    def _browse_subfolder(self):
        path = filedialog.askdirectory(title="저장 폴더 선택",
                                        initialdir=str(self.vault_path))
        if path:
            p = Path(path)
            try:
                rel = p.relative_to(self.vault_path)
                self.subfolder_var.set(str(rel))
            except ValueError:
                messagebox.showwarning("알림", "볼트 경로 내의 폴더를 선택해주세요.")

    # ── 파일 네비게이션 ──
    def _save_current_edits(self):
        cf = self.files[self.current_index]
        cf.filename = self.filename_var.get().strip()
        if not cf.filename.endswith(".md"):
            cf.filename += ".md"
        for key, var in self.fm_vars.items():
            cf.frontmatter[key] = var.get().strip()
        cf.body = self.text_widget.get("1.0", "end-1c")

    def _show_file(self, index: int):
        self.current_index = index
        cf = self.files[index]

        self.nav_label.config(text=f"파일 {index + 1} / {len(self.files)}")
        self.prev_btn.config(state="normal" if index > 0 else "disabled")
        self.next_btn.config(state="normal" if index < len(self.files) - 1 else "disabled")

        self.filename_var.set(cf.filename)
        for key, var in self.fm_vars.items():
            var.set(cf.frontmatter.get(key, ""))

        self.text_widget.delete("1.0", "end")
        self.text_widget.insert("1.0", cf.body)

        self.link_listbox.delete(0, "end")
        self.detected_links = []

    def _go_prev(self):
        if self.current_index > 0:
            self._save_current_edits()
            self._show_file(self.current_index - 1)

    def _go_next(self):
        if self.current_index < len(self.files) - 1:
            self._save_current_edits()
            self._show_file(self.current_index + 1)

    def _skip_current(self):
        if len(self.files) <= 1:
            self.files.pop(self.current_index)
            messagebox.showinfo("알림", "모든 파일이 건너뛰어졌습니다.")
            self.win.destroy()
            return
        self.files.pop(self.current_index)
        new_index = min(self.current_index, len(self.files) - 1)
        self._show_file(new_index)

    # ── 자동 링크 ──
    def _detect_links(self):
        content = self.text_widget.get("1.0", "end-1c")
        self.detected_links = find_linkable_terms(content, self.vault_ctx.notes)
        self.link_listbox.delete(0, "end")
        if not self.detected_links:
            self.link_status.config(text="연결 가능한 노트가 없습니다")
            return
        for name, count in self.detected_links:
            self.link_listbox.insert("end", f"{name}  ({count}회)")
        self.link_status.config(text=f"{len(self.detected_links)}개 노트 감지됨")

    def _apply_selected_links(self):
        selection = self.link_listbox.curselection()
        if not selection:
            messagebox.showinfo("알림", "적용할 항목을 선택하세요.")
            return
        names = [self.detected_links[i][0] for i in selection]
        content = self.text_widget.get("1.0", "end-1c")
        content = apply_auto_links(content, names)
        self.text_widget.delete("1.0", "end")
        self.text_widget.insert("1.0", content)

    def _apply_all_links(self):
        if not self.detected_links:
            return
        names = [name for name, _ in self.detected_links]
        content = self.text_widget.get("1.0", "end-1c")
        content = apply_auto_links(content, names)
        self.text_widget.delete("1.0", "end")
        self.text_widget.insert("1.0", content)

    # ── 수동 링크 삽입 ──
    def _filter_notes(self, *args):
        query = self.search_var.get().strip().lower()
        self.suggest_listbox.delete(0, "end")
        if not query:
            return
        matches = [n for n in self.vault_ctx.notes if query in n.lower()][:20]
        for m in matches:
            self.suggest_listbox.insert("end", m)

    def _insert_manual_link(self):
        sel = self.suggest_listbox.curselection()
        if not sel:
            text = self.search_var.get().strip()
            if not text:
                return
        else:
            text = self.suggest_listbox.get(sel[0])
        link = f"[[{text}]]"
        try:
            self.text_widget.delete("sel.first", "sel.last")
            self.text_widget.insert("insert", link)
        except tk.TclError:
            self.text_widget.insert("insert", link)

    # ── 저장 ──
    def _save_all(self):
        self._save_current_edits()

        sub = self.subfolder_var.get().strip()
        if sub == "(볼트 루트)" or not sub:
            output_dir = self.vault_path
        else:
            output_dir = self.vault_path / sub

        file_list = "\n".join(f"  - {cf.filename}" for cf in self.files)
        confirm = messagebox.askyesno(
            "저장 확인",
            f"다음 {len(self.files)}개 파일을 저장하시겠습니까?\n\n"
            f"{file_list}\n\n"
            f"저장 위치: {output_dir}"
        )
        if not confirm:
            return

        output_dir.mkdir(parents=True, exist_ok=True)
        success = 0
        errors = []

        for cf in self.files:
            try:
                has_fm = any(v for v in cf.frontmatter.values())
                if has_fm:
                    content = serialize_frontmatter(cf.frontmatter) + cf.body
                else:
                    content = cf.body

                out_path = output_dir / cf.filename
                counter = 1
                stem = Path(cf.filename).stem
                while out_path.exists():
                    out_path = output_dir / f"{stem}_{counter}.md"
                    counter += 1

                out_path.write_text(content, encoding="utf-8")
                success += 1
            except Exception as e:
                errors.append(f"{cf.filename}: {e}")

        msg = f"{success}개 파일 저장 완료"
        if errors:
            msg += f"\n\n실패 ({len(errors)}개):\n" + "\n".join(errors)
        msg += f"\n\n저장 위치: {output_dir}"
        messagebox.showinfo("완료", msg)
        self.win.destroy()

    def _cancel(self):
        if messagebox.askyesno("취소", "편집 내용을 버리고 닫으시겠습니까?"):
            self.win.destroy()


# ── 메인 GUI ──────────────────────────────────────────────
class ObsidianConverterApp:
    def __init__(self, root: tk.Tk):
        self.root = root
        self.root.title(WINDOW_TITLE)
        self.root.geometry(WINDOW_SIZE)
        self.root.minsize(700, 500)

        self.file_vars: dict[Path, tk.BooleanVar] = {}
        self.vault_path = tk.StringVar(value=str(DEFAULT_VAULT))
        self.subfolder_name = tk.StringVar(value="")
        self.api_key_var = tk.StringVar(value=load_api_key())

        # 볼트 스캔 (시작 시 백그라운드)
        self.vault_ctx = VaultContext()
        self._vault_scan_done = False
        threading.Thread(target=self._scan_vault_bg, daemon=True).start()

        self._setup_style()
        self._build_ui()

    def _scan_vault_bg(self):
        vault = Path(self.vault_path.get())
        if vault.exists():
            self.vault_ctx = scan_vault(vault)
        self._vault_scan_done = True
        self.root.after(0, lambda: self.status_label.config(
            text=f"볼트 스캔 완료 (노트 {len(self.vault_ctx.notes)}개, "
                 f"태그 {len(self.vault_ctx.all_tags)}개)"))

    def _setup_style(self):
        style = ttk.Style()
        style.theme_use("clam")
        BG = "#F7F8FA"
        PRIMARY = "#4A6CF7"
        TEXT = "#1E293B"
        SUBTLE = "#64748B"

        self.root.configure(bg=BG)
        style.configure("TFrame", background=BG)
        style.configure("Card.TFrame", background="#FFFFFF", relief="solid", borderwidth=1)
        style.configure("TLabel", background=BG, foreground=TEXT, font=("맑은 고딕", 10))
        style.configure("Title.TLabel", font=("맑은 고딕", 14, "bold"), foreground=TEXT)
        style.configure("Subtitle.TLabel", font=("맑은 고딕", 9), foreground=SUBTLE)
        style.configure("TButton", font=("맑은 고딕", 10), padding=(12, 6))
        style.configure("Primary.TButton", font=("맑은 고딕", 11, "bold"),
                         foreground="white", background=PRIMARY, padding=(16, 8))
        style.map("Primary.TButton",
                   background=[("active", "#3B5DE7"), ("disabled", "#94A3B8")])
        style.configure("TCheckbutton", background=BG, font=("맑은 고딕", 9))
        style.configure("TEntry", font=("맑은 고딕", 10), padding=4)

    def _build_ui(self):
        main = ttk.Frame(self.root, padding=16)
        main.pack(fill="both", expand=True)

        ttk.Label(main, text="옵시디언 MD 변환기", style="Title.TLabel").pack(anchor="w")
        ttk.Label(main, text="파일을 선택하면 미리보기에서 속성을 편집한 후 저장할 수 있습니다",
                  style="Subtitle.TLabel").pack(anchor="w", pady=(0, 12))

        # 볼트 경로
        vault_frame = ttk.Frame(main)
        vault_frame.pack(fill="x", pady=(0, 8))
        ttk.Label(vault_frame, text="볼트 경로:").pack(side="left")
        ttk.Entry(vault_frame, textvariable=self.vault_path, width=60).pack(
            side="left", padx=(8, 4), fill="x", expand=True)
        ttk.Button(vault_frame, text="변경", command=self._browse_vault).pack(side="left")

        # 저장 폴더
        sub_frame = ttk.Frame(main)
        sub_frame.pack(fill="x", pady=(0, 8))
        ttk.Label(sub_frame, text="저장 폴더:").pack(side="left")
        ttk.Entry(sub_frame, textvariable=self.subfolder_name, width=30).pack(
            side="left", padx=(8, 4))
        ttk.Label(sub_frame, text="(미리보기에서 변경 가능)",
                  style="Subtitle.TLabel").pack(side="left")

        # API 키 설정
        api_frame = ttk.Frame(main)
        api_frame.pack(fill="x", pady=(0, 12))
        ttk.Label(api_frame, text="API 키:").pack(side="left")
        api_entry = ttk.Entry(api_frame, textvariable=self.api_key_var,
                               width=50, show="*")
        api_entry.pack(side="left", padx=(8, 4))
        ttk.Button(api_frame, text="저장",
                   command=self._save_api_key_main).pack(side="left", padx=(0, 4))
        ttk.Button(api_frame, text="보기/숨기기",
                   command=lambda: api_entry.config(
                       show="" if api_entry.cget("show") == "*" else "*"
                   )).pack(side="left")
        self.api_status_label = ttk.Label(
            api_frame,
            text="저장됨" if self.api_key_var.get() else "미설정 (AI 모드용)",
            style="Subtitle.TLabel")
        self.api_status_label.pack(side="left", padx=(8, 0))

        # 파일 선택 버튼
        btn_frame = ttk.Frame(main)
        btn_frame.pack(fill="x", pady=(0, 8))
        ttk.Button(btn_frame, text="폴더 추가",
                   command=self._add_folder).pack(side="left", padx=(0, 8))
        ttk.Button(btn_frame, text="파일 추가",
                   command=self._add_files).pack(side="left", padx=(0, 8))
        ttk.Button(btn_frame, text="선택 항목 제거",
                   command=self._remove_selected).pack(side="left", padx=(0, 8))
        ttk.Button(btn_frame, text="전체 해제",
                   command=self._clear_all).pack(side="left")

        toggle_frame = ttk.Frame(main)
        toggle_frame.pack(fill="x", pady=(0, 4))
        ttk.Button(toggle_frame, text="전체 선택",
                   command=self._select_all).pack(side="left", padx=(0, 8))
        ttk.Button(toggle_frame, text="전체 해제 (체크)",
                   command=self._deselect_all).pack(side="left")

        # 파일 목록
        list_frame = ttk.Frame(main)
        list_frame.pack(fill="both", expand=True, pady=(0, 8))

        self.canvas = tk.Canvas(list_frame, bg="white", highlightthickness=1,
                                 highlightbackground="#E2E8F0")
        scrollbar = ttk.Scrollbar(list_frame, orient="vertical", command=self.canvas.yview)
        self.scroll_frame = ttk.Frame(self.canvas)
        self.scroll_frame.bind("<Configure>",
                                lambda e: self.canvas.configure(scrollregion=self.canvas.bbox("all")))
        self.canvas.create_window((0, 0), window=self.scroll_frame, anchor="nw")
        self.canvas.configure(yscrollcommand=scrollbar.set)
        self.canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        self.canvas.bind("<MouseWheel>",
                          lambda e: self.canvas.yview_scroll(-1 * (e.delta // 120), "units"))

        # 상태 + 변환 버튼
        bottom_frame = ttk.Frame(main)
        bottom_frame.pack(fill="x")

        self.status_label = ttk.Label(bottom_frame, text="볼트 스캔 중...",
                                       style="Subtitle.TLabel")
        self.status_label.pack(side="left")

        self.convert_btn = ttk.Button(bottom_frame, text="미리보기",
                                       style="Primary.TButton",
                                       command=self._start_convert)
        self.convert_btn.pack(side="right")

        self.file_count_label = ttk.Label(bottom_frame, text="0개 파일",
                                           style="Subtitle.TLabel")
        self.file_count_label.pack(side="right", padx=(0, 16))

    def _save_api_key_main(self):
        key = self.api_key_var.get().strip()
        if key:
            save_api_key(key)
            self.api_status_label.config(text="저장 완료")
        else:
            self.api_status_label.config(text="키가 비어있습니다")

    def _browse_vault(self):
        path = filedialog.askdirectory(title="옵시디언 볼트 선택")
        if path:
            self.vault_path.set(path)
            # 볼트 변경 시 재스캔
            self._vault_scan_done = False
            self.status_label.config(text="볼트 재스캔 중...")
            threading.Thread(target=self._scan_vault_bg, daemon=True).start()

    def _add_folder(self):
        folder = filedialog.askdirectory(title="변환할 폴더 선택")
        if not folder:
            return
        folder_path = Path(folder)
        count = 0
        for f in sorted(folder_path.rglob("*")):
            if f.is_file() and f.suffix.lower() in ALL_SUPPORTED:
                if f not in self.file_vars:
                    self.file_vars[f] = tk.BooleanVar(value=True)
                    count += 1
        self._refresh_file_list()
        self.status_label.config(text=f"'{folder_path.name}' 폴더에서 {count}개 파일 추가됨")

    def _add_files(self):
        filetypes = [
            ("지원 파일", " ".join(f"*{ext}" for ext in sorted(ALL_SUPPORTED))),
            ("문서", "*.docx *.pdf *.txt *.md"),
            ("코드", "*.py *.js *.ts *.json *.yaml *.html *.css *.sql"),
            ("데이터", "*.csv *.xlsx *.xls *.json"),
            ("모든 파일", "*.*"),
        ]
        files = filedialog.askopenfilenames(title="변환할 파일 선택", filetypes=filetypes)
        if not files:
            return
        count = 0
        for f in files:
            fp = Path(f)
            if fp not in self.file_vars:
                self.file_vars[fp] = tk.BooleanVar(value=True)
                count += 1
        self._refresh_file_list()
        self.status_label.config(text=f"{count}개 파일 추가됨")

    def _remove_selected(self):
        to_remove = [fp for fp, var in self.file_vars.items() if var.get()]
        for fp in to_remove:
            del self.file_vars[fp]
        self._refresh_file_list()
        self.status_label.config(text=f"{len(to_remove)}개 파일 제거됨")

    def _clear_all(self):
        self.file_vars.clear()
        self._refresh_file_list()
        self.status_label.config(text="목록 초기화됨")

    def _select_all(self):
        for var in self.file_vars.values():
            var.set(True)

    def _deselect_all(self):
        for var in self.file_vars.values():
            var.set(False)

    def _refresh_file_list(self):
        for widget in self.scroll_frame.winfo_children():
            widget.destroy()
        groups: dict[str, list[Path]] = {}
        for fp in sorted(self.file_vars.keys()):
            groups.setdefault(str(fp.parent), []).append(fp)
        for folder, files in groups.items():
            ttk.Label(self.scroll_frame, text=f"  {folder}",
                      font=("맑은 고딕", 9, "bold"),
                      background="white").pack(anchor="w", padx=8, pady=(8, 2))
            for fp in files:
                var = self.file_vars[fp]
                size = fp.stat().st_size / 1024
                ext = fp.suffix.lower()
                icon = self._get_icon(ext)
                ttk.Checkbutton(self.scroll_frame,
                                 text=f"{icon} {fp.name}  ({size:.1f} KB)",
                                 variable=var).pack(anchor="w", padx=24)
        checked = sum(1 for v in self.file_vars.values() if v.get())
        self.file_count_label.config(text=f"{checked}/{len(self.file_vars)}개 선택")

    def _get_icon(self, ext: str) -> str:
        if ext in CODE_EXTENSIONS:
            return "[C]"
        elif ext in {".docx"}:
            return "[W]"
        elif ext in {".xlsx", ".xls", ".csv"}:
            return "[X]"
        elif ext == ".pdf":
            return "[P]"
        elif ext == ".json":
            return "[J]"
        elif ext in {".md"}:
            return "[M]"
        return "[F]"

    def _start_convert(self):
        selected = [fp for fp, var in self.file_vars.items() if var.get()]
        if not selected:
            messagebox.showwarning("알림", "변환할 파일을 선택하세요.")
            return
        vault = Path(self.vault_path.get())
        if not vault.exists():
            messagebox.showerror("오류", f"볼트 경로가 존재하지 않습니다:\n{vault}")
            return

        self.convert_btn.config(state="disabled")
        self.status_label.config(text="변환 중...")

        def do_convert():
            converted = []
            errors = []
            for fp in selected:
                try:
                    body = convert_file(fp)
                    ext = fp.suffix.lower()
                    fm = {
                        "tags": "",
                        "aliases": "",
                        "date": datetime.now().strftime("%Y-%m-%d"),
                        "version": "",
                        "type": _guess_type(ext),
                        "description": fp.stem,
                        "status": "draft",
                    }
                    converted.append(ConvertedFile(
                        source_path=fp, filename=fp.stem + ".md",
                        frontmatter=fm, body=body,
                    ))
                except Exception as e:
                    errors.append(f"{fp.name}: {e}")

            def open_preview():
                self.convert_btn.config(state="normal")
                if errors:
                    messagebox.showwarning("일부 실패",
                                            f"{len(errors)}개 파일 변환 실패:\n" + "\n".join(errors))
                if converted:
                    self.status_label.config(
                        text=f"{len(converted)}개 파일 변환 완료 — 미리보기 열림")
                    PreviewWindow(self.root, converted, vault,
                                  self.subfolder_name.get().strip(),
                                  self.vault_ctx,
                                  self.api_key_var)
                else:
                    self.status_label.config(text="변환된 파일이 없습니다")

            self.root.after(0, open_preview)

        threading.Thread(target=do_convert, daemon=True).start()


# ── 메인 ───────────────────────────────────────────────────
def main():
    root = tk.Tk()
    app = ObsidianConverterApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()
