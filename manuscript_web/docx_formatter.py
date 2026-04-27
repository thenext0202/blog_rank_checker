"""원고 제작기 — 서식 파싱 (ㄴ 지시) + Word 출력 (.docx)"""
import re


# ╔══════════════════════════════════════════════════════════════╗
# ║  서식 파싱 (ㄴ 지시 → 서식 딕셔너리)                          ║
# ╚══════════════════════════════════════════════════════════════╝


def _split_slash(word):
    """v2.1.6~: 타겟 단어가 `/` 로 구분된 여러 조각이면 분할.

    LLM이 `ㄴ "문장1 / 문장2 / 문장3" 빨간색, 볼드` 식으로 한 ㄴ 안에 여러 조각을
    슬래시로 나열하는 경우, 통째로 매칭하면 본문(줄바꿈으로 쪼개진 상태)에서
    절대 매칭 안 됨. 조각 단위로 분리해 각각 별도 타겟으로 등록한다.

    `.strip('*')` — `'**단어**'` 마크다운 껍질도 함께 제거.
    """
    raw = (word or "").strip().strip('*').strip()
    if not raw:
        return []
    if '/' not in raw:
        return [raw]
    return [p.strip().strip('*').strip() for p in raw.split('/') if p.strip()]


def parse_annotation(annotation_text):
    """ㄴ 서식 지시 줄 → 서식 딕셔너리
    예: 'ㄴ 글자 크기 16, 글꼴 두껍게, 노란 형광펜'
    """
    from docx.enum.text import WD_COLOR_INDEX
    text = annotation_text.lstrip('ㄴ').strip()
    fmt = {
        'font_size': None,
        'bold': False,
        'italic': False,
        'underline': False,
        'colored_words': [],   # [(word, color_name), ...]
        'full_text_color': None,  # 전체 글자색 (옅은회색 등)
        'full_text_color_hex': None,  # 헥스 직접 지정: '0000FF' 등
        'highlight': None,     # WD_COLOR_INDEX 값
        'quote': None,
        'link': False,
        'multi_line': 1,
        'is_image_desc': False,
        'target_words': [],    # — "단어" / "단어" — 형태로 지정된 타겟 단어들
    }

    # 이미지 설명: ㄴ (혈압 측정하는 모습 사진)
    if text.startswith('(') and text.endswith(')'):
        fmt['is_image_desc'] = True
        return fmt

    # 인용구 N번 — 글자 크기 미지정 시 기본 19pt 주입 (LLM이 종종 누락)
    m = re.search(r'인용구\s*(\d+)\s*번', text)
    if m:
        fmt['quote'] = int(m.group(1))
        if fmt['font_size'] is None:
            fmt['font_size'] = 19

    # 글자 크기 (유효 크기만 허용) — '글자 크기 N' 또는 'Npt' 표기 수용
    VALID_FONT_SIZES = [11, 12, 13, 15, 16, 19, 24, 28]
    m = re.search(r'글자\s*크기\s*(\d+)', text)
    if not m:
        m = re.search(r'(\d+)\s*pt\b', text, re.IGNORECASE)
    if m:
        requested = int(m.group(1))
        fmt['font_size'] = min(VALID_FONT_SIZES, key=lambda x: abs(x - requested))

    # `'단어' 볼드` / `'단어' 밑줄` — 단어별 볼드/밑줄 추출 (v2.1.4~)
    # 여러 그룹(`'A' 볼드, 'B' 볼드`)도 모두 수집하도록 finditer 사용
    # `\s*[가-힣]*` — '단어'에 / '단어' 에 / '단어'는 같은 조사(공백 포함) 허용
    # `.strip('*')` — LLM이 ㄴ 지시 안에 넣는 `'**단어**'` 마크다운 볼드 껍질 제거
    # (본문 visible_text에는 `**` 가 없어서 껍질 붙은 채로는 단어 매칭 실패)
    for m in re.finditer(r"((?:'[^']+'\s*[가-힣]*\s*,?\s*)+)\s*(?:글꼴\s*두껍게|두껍게|볼드|bold)", text, re.IGNORECASE):
        for w in re.findall(r"'([^']+)'", m.group(1)):
            for sub in _split_slash(w):
                fmt.setdefault('bolded_words', []).append(sub)
    for m in re.finditer(r"((?:'[^']+'\s*[가-힣]*\s*,?\s*)+)\s*(?:밑줄|underline)", text, re.IGNORECASE):
        for w in re.findall(r"'([^']+)'", m.group(1)):
            for sub in _split_slash(w):
                fmt.setdefault('underlined_words', []).append(sub)

    # 볼드 — '글꼴 두껍게' / '두껍게' / '볼드' / 'bold' (앞에 '단어' 없는 경우만 문단 전체)
    _bold_m = re.search(r"(?<!['\"])\s*(글꼴\s*두껍게|두껍게|볼드|bold)", text, re.IGNORECASE)
    if _bold_m:
        _before = text[:_bold_m.start()].rstrip()
        # 바로 앞이 `'단어'` 또는 `'단어'에` / `'단어' 에` 식으로 끝나면 bolded_words 용
        if not re.search(r"""['\"][^'\"]+['\"]\s*[가-힣]*\s*,?\s*$""", _before):
            fmt['bold'] = True

    # 밑줄 — 동일 패턴
    _under_m = re.search(r"(?<!['\"])\s*(밑줄|underline)", text, re.IGNORECASE)
    if _under_m:
        _before = text[:_under_m.start()].rstrip()
        if not re.search(r"""['\"][^'\"]+['\"]\s*[가-힣]*\s*,?\s*$""", _before):
            fmt['underline'] = True

    # 이탤릭 / 기울임
    if re.search(r'이탤릭|기울임|글꼴\s*기울임', text):
        fmt['italic'] = True

    # 글자색 전체 적용: "글자색 옅은 회색", "글자색 파란색" 등 (특정 단어가 아닌 전체)
    full_color_map = {
        '옅은 회색': '옅은회색', '많이 옅은 회색': '많이옅은회색',
        '회색': '회색', '진한 회색': '진한회색',
        '빨간색': '빨간색', '파란색': '파란색', '청록색': '청록색',
        '초록색': '초록색', '보라색': '보라색', '주황색': '주황색',
    }
    for pattern, color_key in full_color_map.items():
        if re.search(rf'글자\s*색\s*{re.escape(pattern)}', text):
            fmt['full_text_color'] = color_key
            break

    # 색상 키워드: '단어' 빨간색 형태 ('단어1', '단어2' 파란색 도 지원)
    # 하늘색·노란색은 본래 형광펜 색으로만 쓰이지만 v2.1~ 표준 통일로 동일 형식 지원
    color_names = ['빨간색', '파란색', '청록색', '초록색', '보라색', '주황색', '회색', '하늘색', '노란색']
    for color_name in color_names:
        # 뒤에 "형광펜"이 붙으면 하이라이트 매칭 쪽으로 넘기기 (하늘색/노란색 중복 매칭 방지)
        # 여러 단어 그룹(`'A' 빨간색, 'B' 빨간색`)도 모두 수집
        # `\s*[가-힣]*` — '단어'에 / '단어' 에 / '단어'는 같은 조사(공백 포함) 허용
        # `.strip('*')` — `'**단어**'` 마크다운 껍질 제거 (본문 매칭 실패 방지)
        for m in re.finditer(rf"((?:'[^']+'\s*[가-힣]*\s*,?\s*)+)\s*{color_name}(?!\s*형광펜)", text):
            for w in re.findall(r"'([^']+)'", m.group(1)):
                for sub in _split_slash(w):
                    fmt['colored_words'].append((sub, color_name))

    # 헥스 직접 지정: '빨간색(FF0000)', '파란색(0000FF)', '파란색(1155CC)' 등
    m = re.search(
        r'(빨간색|파란색|청록색|초록색|보라색|주황색|회색|하늘색|노란색)'
        r'\s*\(\s*#?([0-9A-Fa-f]{6})\s*\)', text)
    if m:
        fmt['full_text_color_hex'] = m.group(2).upper()

    # 따옴표 없는 색상명 단독 표기 → 전체 글자색
    # 예: 'ㄴ 파란색, 볼드' → full_text_color = 파란색
    # 예: 'ㄴ 빨간색, \'단어\' 파란색' → full_text_color = 빨간색 + colored_words = [(단어, 파란색)]
    #   (두 ㄴ 주석이 _merged로 합쳐진 경우 전체 색상 + 특정 단어 색상이 공존)
    # 가드:
    #   - 색상명 바로 앞에 '단어' 형식이 있으면 colored_words 용이므로 skip
    #   - 뒤에 '형광펜'이 오면 하이라이트 용이므로 skip
    if not fmt['full_text_color_hex'] and not fmt['full_text_color']:
        for cn in ['빨간색', '파란색', '청록색', '초록색', '보라색', '주황색', '회색', '하늘색', '노란색']:
            m2 = re.search(rf"(?<![\'가-힣]){re.escape(cn)}(?=[\s,()]|$)", text)
            if m2:
                before = text[:m2.start()].rstrip()
                # 바로 앞이 `'단어'` 또는 `"단어"` 로 끝나면 colored_words 용 → skip
                # `\s*[가-힣]*` — '단어'에, '단어' 에 같은 조사(공백 포함) 붙은 형태도 skip
                if re.search(r"""['\"][^'\"]+['\"]\s*[가-힣]*\s*$""", before):
                    continue
                tail = text[m2.end():].strip()
                if tail.startswith('형광펜'):
                    continue
                fmt['full_text_color'] = cn
                break

    # 흰 글자 / 흰색 글자 / 흰 글씨 — 검정 형광펜 위 흰색 글자 지시
    if re.search(r'흰\s*(글자|색|글씨|색\s*글자)', text):
        fmt['full_text_color_hex'] = 'FFFFFF'

    # 형광펜 (노란/검정/파란/하늘/빨간/초록/청록)
    highlight_map = {
        '노란|노랑': WD_COLOR_INDEX.YELLOW,
        '검정|검은': WD_COLOR_INDEX.BLACK,
        '하늘': WD_COLOR_INDEX.TURQUOISE,   # 하늘색(cyan) → Word 표준 터콰이즈
        '파란|파랑': WD_COLOR_INDEX.BLUE,
        '빨간|빨강': WD_COLOR_INDEX.RED,
        '초록': WD_COLOR_INDEX.GREEN,
        '청록': WD_COLOR_INDEX.TEAL,
    }
    # v2.1.4~: 먼저 `'단어' 색상 형광펜` 형태를 단어별 형광펜으로 추출.
    # 남은(단독 `ㄴ 색상 형광펜`) 지시는 fmt['highlight']로 문단 전체 적용.
    # `\s*[가-힣]*` — '단어'에, '단어' 에 같은 조사(공백 포함) 허용 (예: '블러디션 배합' 에 검정 형광펜)
    # `.strip('*')` — `'**단어**'` 마크다운 껍질 제거 (본문 매칭 실패 방지)
    for hl_pattern, hl_val in highlight_map.items():
        # 여러 단어 그룹(`'A' 노란 형광펜, 'B' 노란 형광펜`)도 모두 수집
        for m in re.finditer(rf"((?:'[^']+'\s*[가-힣]*\s*,?\s*)+)\s*(?:{hl_pattern})색?\s*형광펜", text):
            for w in re.findall(r"'([^']+)'", m.group(1)):
                for sub in _split_slash(w):
                    fmt.setdefault('highlighted_words', []).append((sub, hl_val))
    for hl_pattern, hl_val in highlight_map.items():
        # 단어 추출 뒤에 남은 일반 형광펜 — 앞에 `'..'`가 없는 경우만
        m2 = re.search(rf"(?<!['\"])(?:{hl_pattern})색?\s*형광펜", text)
        if m2:
            # 이 매치 바로 앞에 `'단어'` 가 붙어 있으면 이미 highlighted_words로 잡힌 것 → skip
            # `\s*[가-힣]*` — '단어'에, '단어' 에 같은 조사(공백 포함) 붙은 형태도 skip
            before = text[:m2.start()].rstrip()
            if re.search(r"""['\"][^'\"]+['\"]\s*[가-힣]*\s*$""", before):
                continue
            fmt['highlight'] = hl_val
            break

    # 링크 도구로 삽입/연결 (공백·표현 변형 허용)
    if re.search(r'링크\s*도구\s*로\s*(삽입|연결)', text):
        fmt['link'] = True

    # N줄 모두 (두/세/네/다섯)
    num_map = {'두': 2, '세': 3, '네': 4, '다섯': 5}
    m = re.search(r'(두|세|네|다섯)\s*줄\s*모두', text)
    if m:
        fmt['multi_line'] = num_map.get(m.group(1), 1)

    # 타겟 단어 추출 — 대시(—/–/-) 주변의 큰따옴표 단어
    # 예: 'ㄴ 하늘색 형광펜, 볼드 — "오메가3추천"' → target_words=['오메가3추천']
    # 예: 'ㄴ "블러디션 배합" — 검정 형광펜, 볼드' → target_words=['블러디션 배합']
    target_words_found = []
    for m in re.finditer(r'[—–\-]\s*"([^"]+)"', text):
        target_words_found.append(m.group(1))
    for m in re.finditer(r'"([^"]+)"\s*[—–\-]', text):
        target_words_found.append(m.group(1))
    # 괄호 안 단어 타겟 인식 — LLM이 자주 쓰는 표기
    # 1) 따옴표 있는 케이스 (기존):
    #    "ㄴ 하늘색 형광펜 ('천연멜라토닌'만)" → ['천연멜라토닌']
    #    "('단어1', '단어2')" → ['단어1', '단어2']
    # 2) v2.1.7~: 따옴표 없는 케이스 — `ㄴ 색상[형광펜], 볼드 (단어)` LLM 표기 흔함
    #    "ㄴ 노란 형광펜, 볼드 (rTG 오메가3)" → ['rTG 오메가3']
    #    "ㄴ 검정 형광펜, 볼드 (블러디션 배합)" → ['블러디션 배합']
    #    "ㄴ 파란색, 볼드 (정상이 150 미만)" → ['정상이 150 미만']
    #    inner에 메타 키워드(색상/형광펜/볼드/크기 등) 있으면 skip — 메타 자체를 단어로 오인 방지
    # 전체가 괄호인 경우는 위의 is_image_desc 경로에서 이미 return되어 여기 도달 못 함
    _META_IN_PAREN = re.compile(
        r'(빨간색|파란색|청록색|초록색|보라색|주황색|회색|하늘색|노란색|'
        r'검정|검은|흰|볼드|bold|밑줄|underline|형광펜|두껍게|이탤릭|'
        r'기울임|글자\s*크기|\d+\s*pt|인용구)',
        re.IGNORECASE,
    )
    _paren_target_added = False  # v2.1.7.1: 괄호 추출로 단어가 등록됐는지 표시
    for m_paren in re.finditer(r'\(([^()]*)\)', text):
        inner = m_paren.group(1).strip()
        if not inner:
            continue
        # 따옴표(작은/큰) 있으면 그것만 등록
        quoted = re.findall(r"""['"]([^'"]+)['"]""", inner)
        if quoted:
            for q in quoted:
                target_words_found.append(q)
            continue
        # 따옴표 없는 — 메타 키워드 없을 때만 콤마 split 후 등록
        if _META_IN_PAREN.search(inner):
            continue
        for piece in re.split(r'[,，]', inner):
            p = piece.strip().strip('*').strip()
            if p:
                target_words_found.append(p)
                _paren_target_added = True
    # 중복 제거(순서 유지) — 슬래시 분할 후 등록
    seen = set()
    for w in target_words_found:
        for sub in _split_slash(w):
            if sub and sub not in seen:
                fmt['target_words'].append(sub)
                seen.add(sub)

    # v2.1.9~: 단순 모드 — 단어 단위 자동 색칠 비활성화.
    # LLM 출력의 단어 표기 모호성으로 단어 의도를 정규식 후처리로 추측하다 6회 회귀 발생.
    # 단어 단위 정보(colored_words/highlighted_words/bolded_words/underlined_words/target_words)를
    # 단락 단위(full_text_color/highlight/bold/underline)로 승격하고 단어 정보는 비움.
    # → _apply_formatting_to_para 가 항상 경로 B (단락 통째 적용) 로 동작 = 케이스 분기 없음.
    # 단어 단위 강조가 필요하면 사용자가 편집창에서 직접 (드래그+색 버튼).
    # 본문 `**단어**` 마크다운은 그대로 작동 → ** 범위만 적용.
    if fmt.get('colored_words'):
        if not fmt.get('full_text_color') and not fmt.get('full_text_color_hex'):
            fmt['full_text_color'] = fmt['colored_words'][0][1]
        fmt['colored_words'] = []
    if fmt.get('highlighted_words'):
        if not fmt.get('highlight'):
            fmt['highlight'] = fmt['highlighted_words'][0][1]
        fmt['highlighted_words'] = []
    if fmt.get('bolded_words'):
        fmt['bold'] = True
        fmt['bolded_words'] = []
    if fmt.get('underlined_words'):
        fmt['underline'] = True
        fmt['underlined_words'] = []
    fmt['target_words'] = []

    return fmt


def _is_self_reference_annotation(text):
    """ㄴ 주석 자신의 표시 스펙만 담긴 줄인지 판별.

    예: 'ㄴ 초록 형광펜' / 'ㄴ 초록 형광펜, 24pt, 볼드'
    — ㄴ 주석은 이미 초록 형광펜 24pt 볼드로 자동 표시되므로 이런 줄은 무시해야 함.
    '초록 형광펜'이 들어있고, 나머지가 크기·pt·볼드·구분자뿐이면 True.
    """
    s = text.lstrip('ㄴ').strip()
    if not re.search(r'초록\s*형광펜', s):
        return False
    cleaned = re.sub(r'초록\s*형광펜', '', s)
    cleaned = re.sub(r'\d+\s*pt', '', cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r'글자\s*크기\s*\d+', '', cleaned)
    # 볼드/두껍게/bold 키워드도 self-reference 판정 시 제거
    cleaned = re.sub(r'글꼴\s*두껍게|두껍게|볼드|bold', '', cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r'[,\s]', '', cleaned)
    return cleaned == ''


def _augment_label_with_bold(display, targets_text):
    """ㄴ 초록 라벨 표시 텍스트 보정 — targets에 `**..**` 볼드가 있는데
    ㄴ 지시에 볼드 키워드가 없으면 '라벨에만' ', 글꼴 두껍게'를 붙여 시각적 힌트를 제공.
    실제 서식 적용은 바꾸지 않음.
    """
    if not targets_text:
        return display
    has_md_bold = any(re.search(r'\*\*[^*]+\*\*', t or '') for t in targets_text)
    if not has_md_bold:
        return display
    if re.search(r'글꼴\s*두껍게|두껍게|볼드|bold', display, re.IGNORECASE):
        return display
    return display + ', 글꼴 두껍게'


def _annotation_display_text(ann):
    """ㄴ 주석 줄을 Word 화면에 표시할 때 쓰는 정제 텍스트.
    - 색상명 뒤 (헥스코드) 제거: '빨간색(FF0000)' → '빨간색'
    - 링크 지시 포함되면 'ㄴ 링크 도구로 삽입' 단독으로 치환
    - LLM이 ㄴ 라인 끝에 축약 메타로 붙이는 '(만)' / '(단어 한정)' 꼬리 제거
    """
    if re.search(r'링크\s*도구\s*로\s*(삽입|연결)', ann):
        return 'ㄴ 링크 도구로 삽입'
    ann = re.sub(r'\s*\(\s*(만|단어\s*한정|단어만)\s*\)\s*$', '', ann)
    return re.sub(
        r'(빨간색|파란색|청록색|초록색|보라색|주황색|회색|하늘색|노란색)\s*\(\s*#?[0-9A-Fa-f]{6}\s*\)',
        r'\1', ann)


def _is_format_annotation(text):
    """ㄴ로 시작하는 줄이 서식 지시인지 콘텐츠인지 판별.
    서식 키워드가 있으면 True, 없으면 False (일반 콘텐츠)."""
    stripped = text.lstrip('ㄴ').strip()
    if stripped.startswith('(') and stripped.endswith(')'):
        return True
    if re.search(r'글자\s*크기|글꼴\s*두껍게|두껍게|형광펜|인용구|이탤릭|기울임|링크\s*도구|줄\s*모두|글자\s*색', stripped):
        return True
    # 볼드/밑줄/pt 표기
    if re.search(r'볼드|bold|밑줄|underline|\d+\s*pt\b', stripped, re.IGNORECASE):
        return True
    # 'N단어' 형태 색상 (기존)
    if re.search(r"'[^']+'\s*(빨간색|파란색|청록색|초록색|보라색|주황색|회색)", stripped):
        return True
    # 색상명 + 헥스 괄호 / 색상명 + 쉼표·괄호·줄끝 (단독 색상 지시)
    if re.search(
            r'(빨간색|파란색|청록색|초록색|보라색|주황색|회색|하늘색|노란색)\s*(\(|,|$)',
            stripped):
        return True
    return False


# ╔══════════════════════════════════════════════════════════════╗
# ║  색상/스타일 헬퍼                                             ║
# ╚══════════════════════════════════════════════════════════════╝

def _get_color_name_to_rgb():
    from docx.shared import RGBColor
    return {
        '빨간색': RGBColor(0xFF, 0x00, 0x00),
        '파란색': RGBColor(0x00, 0x70, 0xC0),
        '하늘색': RGBColor(0x87, 0xCE, 0xEB),
        '노란색': RGBColor(0xFF, 0xC0, 0x00),
        '청록색': RGBColor(0x00, 0x80, 0x80),
        '초록색': RGBColor(0x00, 0x80, 0x00),
        '보라색': RGBColor(0x70, 0x30, 0xA0),
        '주황색': RGBColor(0xED, 0x7D, 0x31),
        '회색': RGBColor(0x80, 0x80, 0x80),
        '많이옅은회색': RGBColor(0xC0, 0xC0, 0xC0),
        '옅은회색': RGBColor(0xA0, 0xA0, 0xA0),
        '진한회색': RGBColor(0x50, 0x50, 0x50),
    }


def _split_colored_words_across_targets(targets, colored_words):
    """colored_words 텍스트가 여러 문단에 걸칠 때 문단별로 분리.
    Returns: dict[int, list[(word, color)]] or None (분리 불필요 시)
    """
    if not colored_words or len(targets) <= 1:
        return None
    spanning = [w for w, _ in colored_words if not any(w in t for _, t in targets)]
    if not spanning:
        return None

    all_texts = [t for _, t in targets]
    joined = ' '.join(all_texts)
    char_colors = [None] * len(joined)
    for word, color_name in colored_words:
        for m in re.finditer(re.escape(word), joined):
            for j in range(m.start(), m.end()):
                char_colors[j] = color_name
    result = {}
    pos = 0
    for idx, (_, para_text) in enumerate(targets):
        para_cw = []
        pc = char_colors[pos:pos + len(para_text)]
        i = 0
        while i < len(para_text):
            if pc[i] is not None:
                color = pc[i]
                j = i
                while j < len(para_text) and pc[j] == color:
                    j += 1
                para_cw.append((para_text[i:j], color))
                i = j
            else:
                i += 1
        result[idx] = para_cw
        pos += len(para_text) + 1  # +1 for space separator
    return result


# ╔══════════════════════════════════════════════════════════════╗
# ║  Word 출력 (.docx)                                           ║
# ╚══════════════════════════════════════════════════════════════╝

def _build_styled_segments(original_text, colored_words, highlighted_words=None,
                           bolded_words=None, underlined_words=None):
    """텍스트를 마크다운 볼드/이탤릭 + 단어별 색상/형광펜/볼드/밑줄 기준으로 세그먼트 분할"""
    md_re = re.compile(r'\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*')

    chunks = []  # (text, is_bold, is_italic)
    last = 0
    for m in md_re.finditer(original_text):
        if m.start() > last:
            chunks.append((original_text[last:m.start()], False, False))
        if m.group(1):
            chunks.append((m.group(1), True, True))
        elif m.group(2):
            chunks.append((m.group(2), True, False))
        elif m.group(3):
            chunks.append((m.group(3), False, True))
        last = m.end()
    if last < len(original_text):
        chunks.append((original_text[last:], False, False))

    highlighted_words = highlighted_words or []
    bolded_words = bolded_words or []
    underlined_words = underlined_words or []
    if not colored_words and not highlighted_words and not bolded_words and not underlined_words:
        return [(t, {'bold': b, 'italic': it, 'color': None, 'highlight': None,
                     'word_bold': False, 'word_underline': False})
                for t, b, it in chunks]

    visible_text = ''.join(c[0] for c in chunks)
    char_colors = [None] * len(visible_text)
    for word, color_name in colored_words:
        for m in re.finditer(re.escape(word), visible_text):
            for j in range(m.start(), m.end()):
                char_colors[j] = color_name
    char_highlights = [None] * len(visible_text)
    for word, hl_val in highlighted_words:
        for m in re.finditer(re.escape(word), visible_text):
            for j in range(m.start(), m.end()):
                char_highlights[j] = hl_val
    char_bold = [False] * len(visible_text)
    for word in bolded_words:
        for m in re.finditer(re.escape(word), visible_text):
            for j in range(m.start(), m.end()):
                char_bold[j] = True
    char_underline = [False] * len(visible_text)
    for word in underlined_words:
        for m in re.finditer(re.escape(word), visible_text):
            for j in range(m.start(), m.end()):
                char_underline[j] = True

    segments = []
    pos = 0
    for chunk_text, is_bold, is_italic in chunks:
        i = 0
        while i < len(chunk_text):
            cur_color = char_colors[pos + i]
            cur_hl = char_highlights[pos + i]
            cur_wb = char_bold[pos + i]
            cur_wu = char_underline[pos + i]
            j = i
            while (j < len(chunk_text)
                   and char_colors[pos + j] == cur_color
                   and char_highlights[pos + j] == cur_hl
                   and char_bold[pos + j] == cur_wb
                   and char_underline[pos + j] == cur_wu):
                j += 1
            segments.append((chunk_text[i:j], {
                'bold': is_bold, 'italic': is_italic,
                'color': cur_color, 'highlight': cur_hl,
                'word_bold': cur_wb, 'word_underline': cur_wu,
            }))
            i = j
        pos += len(chunk_text)

    return segments


def _clear_paragraph_runs(para):
    """기존 run 제거 (pPr 유지)"""
    from docx.oxml.ns import qn
    p_elem = para._element
    for r in list(p_elem.findall(qn('w:r'))):
        p_elem.remove(r)


def _apply_quote_border(paragraph, quote_num):
    """인용구 스타일 — 왼쪽 컬러 테두리"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    color_map = {
        1: '4472C4', 2: '70AD47', 3: 'ED7D31',
        4: 'FFC000', 5: '5B9BD5', 6: '7030A0',
    }
    border_color = color_map.get(quote_num, '4472C4')
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '18')
    left.set(qn('w:space'), '8')
    left.set(qn('w:color'), border_color)
    pBdr.append(left)
    pPr.append(pBdr)


def _apply_formatting_to_para(para, original_text, fmt):
    """ㄴ 서식 딕셔너리를 해당 문단에 실제 적용.

    규칙: 원문에 `**..**` 볼드 마크다운이 하나라도 있으면 **색/형광펜/밑줄/볼드는 그 범위에만** 적용.
    없으면 단락 전체에 적용. 글자 크기와 인용구는 항상 단락 전체.
    target_words가 있으면 해당 단어에만 색/형광펜/볼드/밑줄 적용 (문단 전체 X).
    """
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_COLOR_INDEX
    BLUE_C = RGBColor(0x00, 0x70, 0xC0)
    WHITE_C = RGBColor(0xFF, 0xFF, 0xFF)

    is_quote = bool(fmt.get('quote'))
    target_words = fmt.get('target_words') or []

    # v2.1.7~: target_words 경로는 fmt['highlight']/fmt['full_text_color']만 단어에 적용.
    # highlighted_words / colored_words 가 같이 채워진 경우 (예: `ㄴ '단어' 색 형광펜` 의
    # 단어가 괄호 추출과 단일따옴표 추출 양쪽에 잡힘), 첫 번째 색을 효과 단락-단위 색으로 사용.
    # — fmt 자체는 수정하지 않음. _build_document 의 has_para_only_fmt 체크가 이 승격에
    #   영향받아 비매칭 라인까지 칠해지는 부작용 방지.
    eff_highlight = fmt.get('highlight')
    if not eff_highlight and fmt.get('highlighted_words'):
        eff_highlight = fmt['highlighted_words'][0][1]
    eff_full_color = fmt.get('full_text_color')
    eff_full_color_hex = fmt.get('full_text_color_hex')
    if not eff_full_color and not eff_full_color_hex and fmt.get('colored_words'):
        eff_full_color = fmt['colored_words'][0][1]

    # ── target_words 경로: 해당 단어에만 서식 적용 ──
    if target_words and not is_quote:
        # 마크다운 ** .. ** / * .. * 를 파싱해 visible text와 char별 볼드/이탤릭 맵 구성
        # (target_words 경로도 마크다운은 그대로 해석해서 렌더링해야 별표가 리터럴로 나가지 않음)
        visible_parts = []
        bold_map = []
        italic_map = []
        md_pat = re.compile(r'\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*')
        _last = 0
        for m in md_pat.finditer(original_text):
            if m.start() > _last:
                chunk = original_text[_last:m.start()]
                visible_parts.append(chunk)
                bold_map.extend([False] * len(chunk))
                italic_map.extend([False] * len(chunk))
            if m.group(1):
                chunk = m.group(1)
                visible_parts.append(chunk)
                bold_map.extend([True] * len(chunk))
                italic_map.extend([True] * len(chunk))
            elif m.group(2):
                chunk = m.group(2)
                visible_parts.append(chunk)
                bold_map.extend([True] * len(chunk))
                italic_map.extend([False] * len(chunk))
            elif m.group(3):
                chunk = m.group(3)
                visible_parts.append(chunk)
                bold_map.extend([False] * len(chunk))
                italic_map.extend([True] * len(chunk))
            _last = m.end()
        if _last < len(original_text):
            tail = original_text[_last:]
            visible_parts.append(tail)
            bold_map.extend([False] * len(tail))
            italic_map.extend([False] * len(tail))
        visible_text = ''.join(visible_parts)

        # 이 para 안에 target_word가 하나도 없으면 건드리지 않음
        if not any(w in visible_text for w in target_words):
            return
        _clear_paragraph_runs(para)

        char_is_target = [False] * len(visible_text)
        for w in target_words:
            for m in re.finditer(re.escape(w), visible_text):
                for j in range(m.start(), m.end()):
                    char_is_target[j] = True

        _color_map = _get_color_name_to_rgb()
        i = 0
        while i < len(visible_text):
            cur_tgt = char_is_target[i]
            cur_b = bold_map[i]
            cur_it = italic_map[i]
            j = i
            while (j < len(visible_text)
                   and char_is_target[j] == cur_tgt
                   and bold_map[j] == cur_b
                   and italic_map[j] == cur_it):
                j += 1
            seg_text = visible_text[i:j]
            run = para.add_run(seg_text)
            if fmt.get('font_size'):
                run.font.size = Pt(fmt['font_size'])
            # 마크다운 볼드/이탤릭은 타겟 여부와 상관없이 적용
            if cur_b:
                run.bold = True
            if cur_it:
                run.italic = True
            if cur_tgt:
                if fmt.get('bold'):
                    run.bold = True
                if fmt.get('italic'):
                    run.italic = True
                if fmt.get('underline'):
                    run.underline = True
                has_color = False
                if eff_full_color_hex:
                    run.font.color.rgb = RGBColor.from_string(eff_full_color_hex)
                    has_color = True
                elif eff_full_color and eff_full_color in _color_map:
                    run.font.color.rgb = _color_map[eff_full_color]
                    has_color = True
                # 검정 형광펜 + 글자색 미지정 → 자동 흰 글자 (가독성 보정)
                if eff_highlight == WD_COLOR_INDEX.BLACK and not has_color:
                    run.font.color.rgb = WHITE_C
                if eff_highlight:
                    run.font.highlight_color = eff_highlight
            i = j
        return

    _clear_paragraph_runs(para)
    segments = _build_styled_segments(
        original_text,
        fmt.get('colored_words', []),
        fmt.get('highlighted_words', []),
        fmt.get('bolded_words', []),
        fmt.get('underlined_words', []),
    )
    has_md_bold_spans = bool(re.search(r'\*\*[^*]+\*\*', original_text))

    for seg_text, seg_props in segments:
        run = para.add_run(seg_text)
        # 크기는 단락 전체
        if fmt.get('font_size'):
            run.font.size = Pt(fmt['font_size'])

        # 볼드: 인용구·세그먼트 볼드·colored_words 범위는 그대로,
        # 그 외 fmt.bold는 ** 범위가 없을 때만 전체 적용.
        # v2.1.4~: word_bold(단어별 볼드)가 있으면 해당 세그먼트만 볼드.
        if is_quote:
            run.bold = True
        elif seg_props.get('bold'):
            run.bold = True
        elif seg_props.get('word_bold'):
            run.bold = True
        elif fmt.get('bold') and seg_props.get('color'):
            run.bold = True
        elif fmt.get('bold') and not fmt.get('colored_words') and not fmt.get('bolded_words') and not has_md_bold_spans:
            run.bold = True

        if fmt.get('italic') or seg_props.get('italic'):
            run.italic = True

        # 시각 서식(색/형광펜/밑줄) 적용 범위 결정
        # v2.1.4~: 명시적 visual 지시(색/형광펜/밑줄/세그먼트 색)는 **..** 범위 무관 항상 적용.
        # v2.1.8~: 단, 단어 명시(colored_words/highlighted_words/bolded_words)가 하나도 없고
        #   단락 단위 fmt(highlight/full_text_color)만 있고 본문에 ** 마크다운이 있으면
        #   → ** 범위에만 적용. LLM이 `ㄴ 색 형광펜, 볼드` 만 달고 본문에 `**단어**` 로 단어를 표시한
        #   케이스에서 문장 전체가 칠해지는 부작용 방지.
        no_word_specs = (
            not fmt.get('colored_words')
            and not fmt.get('highlighted_words')
            and not fmt.get('bolded_words')
        )
        has_seg_visual = bool(seg_props.get('color') or seg_props.get('highlight'))
        has_para_visual = bool(
            fmt.get('full_text_color')
            or fmt.get('full_text_color_hex')
            or fmt.get('highlight')
            or fmt.get('underline')
        )
        if has_md_bold_spans and no_word_specs and has_para_visual and not has_seg_visual:
            # ** 범위에만 적용 (v2.1.4 이전 동작 복원, 단어 specs 없을 때 한정)
            apply_visual = bool(seg_props.get('bold'))
        else:
            apply_visual = (not has_md_bold_spans) or seg_props.get('bold') or has_seg_visual or has_para_visual

        # 밑줄: 세그먼트 word_underline 우선, 그 외 fmt.underline은 underlined_words 없고 apply_visual 일 때
        if seg_props.get('word_underline'):
            run.underline = True
        elif apply_visual and fmt.get('underline') and not fmt.get('underlined_words'):
            run.underline = True

        if not is_quote and apply_visual:
            _color_map = _get_color_name_to_rgb()
            color_name = seg_props.get('color')
            has_any_color = False
            if color_name and color_name in _color_map:
                run.font.color.rgb = _color_map[color_name]
                has_any_color = True
            elif fmt.get('full_text_color_hex'):
                # 헥스 직접 지정 — '빨간색(FF0000)', '파란색(1155CC)' 등
                run.font.color.rgb = RGBColor.from_string(fmt['full_text_color_hex'])
                has_any_color = True
            elif fmt.get('full_text_color'):
                ftc = fmt['full_text_color']
                if ftc in _color_map:
                    run.font.color.rgb = _color_map[ftc]
                    has_any_color = True
            # 검정 형광펜 + 글자색 미지정 → 자동 흰 글자 (가독성 보정)
            seg_hl = seg_props.get('highlight')
            eff_hl = seg_hl if seg_hl is not None else fmt.get('highlight')
            if eff_hl == WD_COLOR_INDEX.BLACK and not has_any_color:
                run.font.color.rgb = WHITE_C
            # 세그먼트별 형광펜 우선, 없으면 문단 전체 형광펜
            if seg_hl is not None:
                run.font.highlight_color = seg_hl
            elif fmt.get('highlight'):
                run.font.highlight_color = fmt['highlight']

    if fmt.get('link'):
        for run in para.runs:
            # 헥스가 명시돼 있으면 그대로 유지, 아니면 기본 파란색
            if not fmt.get('full_text_color_hex'):
                run.font.color.rgb = BLUE_C
            run.underline = True


def _add_text_runs(para, text):
    """마크다운 볼드/이탤릭 처리하여 run 추가"""
    md_re = re.compile(r'\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*')
    last = 0
    for m in md_re.finditer(text):
        if m.start() > last:
            para.add_run(text[last:m.start()])
        if m.group(1):
            run = para.add_run(m.group(1))
            run.bold = True
            run.italic = True
        elif m.group(2):
            run = para.add_run(m.group(2))
            run.bold = True
        elif m.group(3):
            run = para.add_run(m.group(3))
            run.italic = True
        last = m.end()
    if last < len(text):
        para.add_run(text[last:])


def _add_blogger_request_box(doc, lines):
    """★ 블로거 요청사항 → 빨간 테두리 + 노란 배경 테이블 박스"""
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)

    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'FFF8E1')
    shading.set(qn('w:val'), 'clear')
    cell._element.get_or_add_tcPr().append(shading)

    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '12')
        border.set(qn('w:color'), 'FF0000')
        border.set(qn('w:space'), '0')
        tcBorders.append(border)
    cell._element.get_or_add_tcPr().append(tcBorders)

    first_para = cell.paragraphs[0]
    for i, line in enumerate(lines):
        para = cell.add_paragraph() if i > 0 else first_para
        run = para.add_run(line)
        run.bold = True
        run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
        run.font.size = Pt(11)

    doc.add_paragraph('')


# ╔══════════════════════════════════════════════════════════════╗
# ║  줄 길이 정규화 (pre-processing) — 10~28자 타겟               ║
# ╚══════════════════════════════════════════════════════════════╝
# 목표: 한 줄이 공백 포함 10~28자가 되도록 긴 줄은 쪼개고 짧은 줄은 합침.
# 기존 5가지 파싱 로직은 건드리지 않고 _build_document 진입 직전 전처리로만 동작.

_NORM_MIN_LEN = 10
_NORM_MAX_LEN = 28
_NORM_MAX_LEN_LATIN = 40  # v2.1.10: 영문 우세 라인 한도 (한글의 약 1.5배)


def _is_latin_dominant(s):
    """라인이 영문/숫자/공백 위주인지 — 한글 char 비율 30% 미만."""
    if not s:
        return False
    total = sum(1 for c in s if not c.isspace())
    if total == 0:
        return False
    hangul = sum(1 for c in s if '가' <= c <= '힣')
    return (hangul / total) < 0.30


def _max_len_for(s):
    return _NORM_MAX_LEN_LATIN if _is_latin_dominant(s) else _NORM_MAX_LEN


_NORM_MIN_LEN_LATIN = 16  # v2.1.10: 영문 우세 라인 최소 (한글 10자와 시각 폭 비슷)


def _min_len_for(s):
    return _NORM_MIN_LEN_LATIN if _is_latin_dominant(s) else _NORM_MIN_LEN
_NORM_GUARD_LEN = 5            # 의도적 짧은 줄 기준 (네?, 정말? 등)
_NORM_ORPHAN_MIN = 8           # 분할 후 뒤조각이 이보다 짧으면 orphan — 분할 위치를 앞으로 당김
_NORM_END_PUNCT = '.?!~…'      # 의도적 종결부호 — 마침표 포함: 문장 끝난 줄은 다음 줄과 합치지 않음

_CONNECTIVE_ENDINGS = re.compile(
    r'(어서|아서|고|며|면서|지만|는데|은데|니까|'
    r'므로|기에|면|도록|거나|든지|아도|어도|아야|'
    r'어야|라서|라도|'
    # 자주 쓰는 동사 축약형 표면형
    r'해서|해야|해도|봐서|봐야|봐도|'
    r'와서|와야|와도|돼서|돼야|돼도|'
    r'가서|가야|가도|여서|켜서|'
    r'려고|려면)$'
)

_SEPARATOR_RE = re.compile(r'^[─—–\-]{5,}$')


def _ends_with_intentional_punct(line):
    return bool(line) and line[-1] in _NORM_END_PUNCT


def _is_intentional_short(line):
    return len(line) < _NORM_GUARD_LEN and _ends_with_intentional_punct(line)


def _is_norm_boundary(stripped):
    """정규화 블록 경계 — 이 줄들은 정규화 대상에서 제외, 블록도 끊음."""
    if not stripped:
        return True
    if stripped.startswith('ㄴ'):
        return True
    if '★' in stripped:
        return True
    if re.match(r'^\d{1,2}$', stripped):
        return True
    if re.match(r'^#{1,6}\s', stripped):
        return True
    # 제목 라벨 — 길어도 한 줄 유지 ("제목 :", "제목:", "제목：" 등)
    if re.match(r'^제목\s*[:：]', stripped):
        return True
    # 본문 + 중간 ㄴ 지시 (예: "본문 ㄴ 빨간색")
    if re.search(r'\s+ㄴ\s+', stripped):
        return True
    # ─── 구분선
    if _SEPARATOR_RE.match(stripped):
        return True
    return False


def _find_split_point(line, max_len=None, min_len=None):
    """연결어미 우선, 없으면 max_len 이하 가장 뒤 공백에서 분할.
    max_len/min_len 미지정 시 기본 상수 사용.
    반환: (앞 조각, 뒤 조각) 또는 None (분할 불가)
    """
    if max_len is None:
        max_len = _NORM_MAX_LEN
    if min_len is None:
        min_len = _NORM_MIN_LEN
    words = line.split(' ')
    if len(words) < 2:
        return None
    best_connective = None
    best_fallback = None
    cum = 0
    for i, w in enumerate(words[:-1]):
        cum += len(w)
        if cum > max_len:
            break
        if _CONNECTIVE_ENDINGS.search(w) and cum >= min_len:
            best_connective = i
        best_fallback = i
        cum += 1
    split_idx = best_connective if best_connective is not None else best_fallback
    if split_idx is None:
        return None

    # orphan 방지 — 뒤조각이 너무 짧으면(예: "먹었다.") 분할 위치를 한 칸씩 앞으로 당김
    tail = ' '.join(words[split_idx + 1:])
    while len(tail) < _NORM_ORPHAN_MIN and split_idx > 0:
        candidate_idx = split_idx - 1
        head_check = ' '.join(words[:candidate_idx + 1])
        if len(head_check) < min_len:
            break  # 더 당기면 앞조각이 너무 짧아짐 — 중단
        split_idx = candidate_idx
        tail = ' '.join(words[split_idx + 1:])

    # 그래도 뒤조각이 짧으면 분할 보류 (28자 약간 초과는 허용)
    if len(tail) < _NORM_ORPHAN_MIN:
        return None

    return ' '.join(words[:split_idx + 1]), tail


def _split_before_short(lines):
    """앞 줄이 짧고 미완성인데 합치면 초과하는 경우, 뒷 줄을 선제 분할."""
    result = []
    prev = None
    for line in lines:
        if (prev is not None
                and len(prev) < _NORM_MIN_LEN
                and not _ends_with_intentional_punct(prev)
                and not _is_intentional_short(line)
                and len(prev) + 1 + len(line) > _NORM_MAX_LEN):
            max_allowed = _NORM_MAX_LEN - len(prev) - 1
            if max_allowed > 0:
                parts = _find_split_point(line, max_len=max_allowed, min_len=1)
                if parts and parts[0]:
                    result.append(parts[0])
                    line = parts[1]
        result.append(line)
        prev = line
    return result


def _split_long_lines(lines):
    """초과 줄을 연결어미 우선 기준으로 분할 (한글 28자 / 영문우세 40자 한도, v2.1.10)."""
    result = []
    for line in lines:
        max_len = _max_len_for(line)
        while len(line) > max_len:
            parts = _find_split_point(line, max_len=max_len)
            if not parts:
                break
            result.append(parts[0])
            line = parts[1]
        result.append(line)
    return result


def _merge_short_lines(lines):
    """10자 미만 → 다음 줄과 무조건 합침 (v2.1.10: 한도 가드 풀림 — 사용자 결정).
    의도적 짧은 줄 보호. 합쳐진 결과는 한글 28 / 영문우세 40 한도 내에서 우선 시도하지만,
    초과해도 강행 (사용자: '윗 문장 한도 넘어도 한 문장 10자 이하 금지')."""
    result = []
    i = 0
    while i < len(lines):
        curr = lines[i]
        j = i + 1
        while (len(curr) < _min_len_for(curr)
               and not _ends_with_intentional_punct(curr)
               and j < len(lines)
               and not _is_intentional_short(lines[j])):
            curr = curr + ' ' + lines[j]
            j += 1
        result.append(curr)
        i = j
    # v2.1.10: 결과에서 여전히 짧은 줄(마침표 끝남 가드 때문에 합쳐지지 않은 케이스)을
    # 앞 줄에 흡수. 사용자 요구 "한 문장 최소 10자(영문 16자) 이상" 강제.
    final = []
    for ln in result:
        if (final
                and len(ln) < _min_len_for(ln)
                and not _is_intentional_short(ln)):
            final[-1] = final[-1].rstrip() + ' ' + ln
        else:
            final.append(ln)
    return final


def _absorb_short_label_paragraphs(text):
    """v2.1.10: 짧은 `**라벨**` 단독 단락(예: '**홍국**', '**블러디션 배합**')이
    빈 줄로 분리되어 있으면 위 단락 마지막 줄 끝에 흡수.

    조건: 본문 한 줄, 형식 `**...**` 단독, 10자 미만 (한글 우세 기준).
    위 단락이 본문(ㄴ/★/구분선/제목 아님)일 때만.
    빈 줄 한 개를 흡수에 사용. 다음 ㄴ 라인은 그 자리 유지 → 합쳐진 단락에 ㄴ 적용,
    `**` 마크다운 부분만 색칠 (v2.1.9 경로 B 동작).
    """
    if not text:
        return text
    label_pat = re.compile(r'^\s*\*\*[^*]+\*\*\s*$')
    lines = text.split('\n')
    out = list(lines)
    in_box = False
    for i in range(1, len(out)):
        s = (out[i] or '').strip()
        if '★' in s:
            in_box = True; continue
        if in_box:
            if _SEPARATOR_RE.match(s):
                in_box = False
            continue
        if not label_pat.match(s):
            continue
        if len(s) >= _NORM_MIN_LEN:  # 10자 이상이면 그대로
            continue
        # 위가 빈 줄이어야 흡수 (본문 한가운데 라벨은 건드리지 않음)
        if i == 0 or (out[i-1] or '').strip():
            continue
        # 그 위 본문 라인 찾기 (빈 줄 건너)
        k = i - 2
        while k >= 0 and not (out[k] or '').strip():
            k -= 1
        if k < 0:
            continue
        prev = out[k].strip()
        # 위가 본문(ㄴ/★/구분선/제목/헤딩 아님) 일 때만
        if (prev.startswith('ㄴ') or '★' in prev or _SEPARATOR_RE.match(prev)
                or re.match(r'^#{1,6}\s', prev) or re.match(r'^제목\s*[:：]', prev)):
            continue
        # 흡수: 위 본문 끝에 라벨 붙이고, 라벨 라인 + 그 위 빈 줄 제거
        out[k] = out[k].rstrip() + ' ' + s
        out[i] = ''         # 라벨 라인 비움
        out[i-1] = None     # 빈 줄 마커 — 아래에서 None 라인 제거
    return '\n'.join(ln for ln in out if ln is not None)


def _normalize_line_lengths(text):
    """블록 경계를 보존하며 각 블록 내에서 긴 줄 분할 + 짧은 줄 합침 적용.
    ★ 블로거 요청사항 박스 내부는 통째로 건드리지 않음.
    v2.1.10: 짧은 `**라벨**` 단독 단락은 위 단락에 미리 흡수 (10자 미만 금지).
    """
    text = _absorb_short_label_paragraphs(text)
    lines = text.split('\n')
    result = []
    block = []
    in_blogger_req = False

    for line in lines:
        stripped = line.strip()

        # ★ 박스 시작
        if '★' in stripped:
            if block:
                result.extend(_merge_short_lines(_split_before_short(_split_long_lines(block))))
                block = []
            in_blogger_req = True
            result.append(line)
            continue

        # ★ 박스 내부 — 통과
        if in_blogger_req:
            result.append(line)
            if _SEPARATOR_RE.match(stripped):
                in_blogger_req = False
            continue

        # 일반 블록 경계
        if _is_norm_boundary(stripped):
            if block:
                result.extend(_merge_short_lines(_split_before_short(_split_long_lines(block))))
                block = []
            result.append(line)
        else:
            block.append(stripped)

    if block:
        result.extend(_merge_short_lines(_split_before_short(_split_long_lines(block))))
    return '\n'.join(result)


def _ann_color_signature(line):
    """ㄴ 라인의 색·형광펜·인용구 시그니처 추출 — 합쳐도 안전한지 판정용 (v2.1.8).

    같은 시그니처 ㄴ 끼리만 lift/합치기. 다른 색·다른 형광펜이면 별도 라인 유지.
    그래야 합쳐진 라벨이 parse_annotation 에서 단어 색을 뒤섞지 않음.
    """
    s = (line or '').lstrip('ㄴ').strip()
    # 형광펜 (색 + 형광펜 키워드)
    hl = tuple(sorted(set(re.findall(
        r'(노란|노랑|검정|검은|하늘|파란|파랑|빨간|빨강|초록|청록)\s*(?:색)?\s*형광펜', s))))
    # 색상 (형광펜 아닌 단독 색)
    s_no_hl = re.sub(r'(노란|노랑|검정|검은|하늘|파란|파랑|빨간|빨강|초록|청록)\s*(?:색)?\s*형광펜', '', s)
    colors = tuple(sorted(set(re.findall(
        r'(빨간색|파란색|청록색|초록색|보라색|주황색|회색|하늘색|노란색)', s_no_hl))))
    quote = bool(re.search(r'인용구', s))
    return (hl, colors, quote)


def _lift_mid_paragraph_annotations(text):
    """문단 중간에 끼인 ㄴ 주석을 문단 끝으로 이동.
    v2.1.8~: 단락 내 ㄴ 들이 모두 같은 색 시그니처일 때만 lift. 다른 색이면 자리 유지
    (각 ㄴ이 자기 바로 위 본문만 잡도록 — 합쳐서 색 섞이는 부작용 방지).
    """
    if not text:
        return text
    lines = text.split('\n')
    out = []
    in_box = False
    i = 0
    while i < len(lines):
        line = lines[i]
        s = line.strip()
        if '★' in s:
            in_box = True
            out.append(line); i += 1; continue
        if in_box:
            out.append(line)
            if _SEPARATOR_RE.match(s):
                in_box = False
            i += 1; continue
        if s and not s.startswith('ㄴ'):
            # 단락 수집
            block = []  # 원본 라인 그대로
            anns_only = []
            j = i
            while j < len(lines):
                t = lines[j]
                ts = t.strip()
                if not ts or '★' in ts:
                    break
                block.append(t)
                if ts.startswith('ㄴ'):
                    anns_only.append(t)
                j += 1
            # ㄴ 들의 시그니처가 모두 같은가?
            sigs = {_ann_color_signature(a) for a in anns_only}
            if len(anns_only) >= 2 and len(sigs) == 1:
                # 같은 색 → lift (본문 먼저, ㄴ 끝으로)
                bodies = [b for b in block if not b.strip().startswith('ㄴ')]
                out.extend(bodies)
                out.extend(anns_only)
            else:
                # 다른 색 또는 ㄴ 1개 이하 → 원본 순서 유지
                out.extend(block)
            i = j
        else:
            out.append(line); i += 1
    return '\n'.join(out)


# 여러 줄에 걸친 `**...**` 탐지 — DOTALL 필요 (. 가 \n 매치)
_MULTILINE_BOLD_RE = re.compile(r'\*\*([^*]+?)\*\*', re.DOTALL)


def _rewrap_multiline_bold(text):
    """줄 길이 정규화로 `**A\\nB\\nC**` 형태가 된 볼드를 각 줄 개별 래핑으로 복구.

    워드·HTML 렌더러는 줄 단위로 `**...**` 쌍을 찾기 때문에, 여는 `**`는 첫 줄에
    닫는 `**`는 마지막 줄에만 있으면 엔진이 짝을 못 찾고 `*` 기호가 그대로 출력됨.
    → `**A**\\n**B**\\n**C**` 로 재래핑해 각 줄이 독립적으로 볼드 처리되게 만듦.
    ㄴ 지시 줄과 빈 줄은 래핑 대상에서 제외.
    """
    if not text or '**' not in text:
        return text

    def _repl(m):
        inner = m.group(1)
        if '\n' not in inner:
            return m.group(0)
        out = []
        for ln in inner.split('\n'):
            s = ln.strip()
            if not s:
                out.append(ln)
            elif s.startswith('ㄴ'):
                out.append(ln)
            else:
                leading = ln[:len(ln) - len(ln.lstrip())]
                trailing = ln[len(ln.rstrip()):]
                out.append(f'{leading}**{ln.strip()}**{trailing}')
        return '\n'.join(out)

    return _MULTILINE_BOLD_RE.sub(_repl, text)


# ── _build_document: 텍스트 → docx.Document 객체 (서식 적용) ──
def _build_document(text, normalize=True):
    """텍스트 → docx 문서.

    normalize=True: 줄 길이 정규화를 적용 (워드 다운로드 기본 동작 유지).
    normalize=False: 편집창에서 내려준 본문을 1:1 그대로 렌더링 — 미리보기/사용자가
    수동으로 쪼갠 줄바꿈을 존중하기 위함. 사용자는 /generate 시점에 이미 정규화된
    본문을 받으므로 이후 편집창 내용은 절대 진실로 취급.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '맑은 고딕'
    style.font.size = Pt(11)

    annotation_re = re.compile(r'^ㄴ\s*')
    image_num_re = re.compile(r'^\d{1,2}$')

    GREEN = RGBColor(0x00, 0x80, 0x00)
    BLUE = RGBColor(0x00, 0x70, 0xC0)

    # 줄 길이 정규화 (10~28자 타겟) — 기존 파싱 전 전처리
    if normalize:
        text = _normalize_line_lengths(text)
    # 여러 줄에 걸친 `**...**` 복구 — 줄 쪼개기가 여는/닫는 쌍을 토막 낸 경우
    # 엔진이 `*` 을 리터럴로 출력하는 것을 막기 위해 각 줄 개별 래핑으로 변환.
    # normalize=False (미리보기/다운로드) 경로에서도 편집창 body에 이미 split된
    # `**A\\nB\\nC**` 가 들어있을 수 있으므로 normalize와 무관하게 항상 실행.
    text = _rewrap_multiline_bold(text)
    # v2.1~ 변경:
    # - `_normalize_highlight_ann_targets` 호출/함수 삭제 — color_names에 하늘/노란 추가 후 불필요.
    # - 레거시 em-dash 타겟은 output_parser._em_dash_to_quote 가 진입 시점에 변환.
    # - parse()에서는 `_merge_same_target_annotations` 호출 제거, 여기서만 한 번 실행.
    #   (편집창 사용자 수동 편집으로 생긴 중복 ㄴ 정리는 미리보기·다운로드 시점에만 필요.)
    # v2.1.4~: 문단 중간 ㄴ 주석을 문단 끝으로 이동 (적용 대상 통합 + 같은 para 중복 apply 방지)
    text = _lift_mid_paragraph_annotations(text)
    from output_parser import _merge_same_target_annotations
    text = _merge_same_target_annotations(text)

    lines = text.split('\n')
    # 연속된 ㄴ 서식 라인 병합 — 같은 단락에 두 번 apply되면 _clear_paragraph_runs가
    # 첫 번째 서식을 날려버리므로 하나의 ㄴ 라인으로 합쳐서 한 번에 적용한다.
    _merged = []
    _i = 0
    while _i < len(lines):
        _cur = lines[_i]
        _s = _cur.strip()
        # ㄴ 주석 자기-참조 줄(ㄴ 초록 형광펜 등)은 완전 제거
        if _s.startswith('ㄴ') and _is_self_reference_annotation(_s):
            _i += 1
            continue
        if _s.startswith('ㄴ') and _is_format_annotation(_s):
            _combined = _s
            _cur_sig = _ann_color_signature(_s)
            _j = _i + 1
            while _j < len(lines):
                _nxt = lines[_j].strip()
                if _nxt.startswith('ㄴ') and _is_self_reference_annotation(_nxt):
                    _j += 1  # 병합 대상에서 제외, 라인 자체도 drop
                    continue
                if _nxt.startswith('ㄴ') and _is_format_annotation(_nxt):
                    # v2.1.8~: 같은 색 시그니처일 때만 합침. 다른 색이면 별도 라인 유지
                    # (parse_annotation 한 번에 여러 색이 섞여 본문에 색이 새는 부작용 방지)
                    if _ann_color_signature(_nxt) != _cur_sig:
                        break
                    _combined = _combined + ', ' + _nxt.lstrip('ㄴ').strip()
                    _j += 1
                else:
                    break
            _merged.append(_combined)
            _i = _j
        else:
            _merged.append(_cur)
            _i += 1
    lines = _merged
    recent = []  # (paragraph, original_text) 버퍼
    pending_fmts = []  # 아래 텍스트에 적용할 대기 서식
    blogger_req_lines = []
    in_blogger_req = False

    for line in lines:
        stripped = line.strip()

        # ── ★ 블로거 요청사항 수집 ──
        if '★' in stripped:
            in_blogger_req = True
            blogger_req_lines.append(stripped)
            continue

        if in_blogger_req:
            # 구분선(─/—/-) 5개 이상에서 박스 종료 (박스 밖에 구분선 텍스트 출력 X)
            if re.match(r'^[─—–\-]{5,}$', stripped):
                _add_blogger_request_box(doc, blogger_req_lines)
                blogger_req_lines = []
                in_blogger_req = False
                recent.append((doc.paragraphs[-1] if doc.paragraphs else None, ''))
                continue
            # ㄴ 서식 지시 줄은 박스 전체 서식용 표시이므로 drop (박스 내 노출 X)
            if stripped.startswith('ㄴ'):
                continue
            # 박스 내부의 빈 줄은 무시 (박스 닫지 않음)
            if not stripped:
                continue
            blogger_req_lines.append(stripped)
            continue

        # ── 빈 줄 ──
        if not stripped:
            p = doc.add_paragraph('')
            recent.append((p, ''))
            continue

        # ── 제목 라인 (Phase E body 첫 머리에 남는 "제목 : ..." 라벨) ──
        # v2.1.6~: 라벨로만 표시하고 단어별 서식 매칭 대상에서 제외.
        # recent에 등록하지 않으면 후속 ㄴ 주석의 search_words / content_group 수집에서
        # 자연스럽게 빠짐 → 키워드 형광펜이 제목 라인 대신 본문 첫 라인에 적용됨.
        if re.match(r'^\*{0,2}\s*제목\s*\*{0,2}\s*[:：]', stripped):
            p = doc.add_paragraph()
            _add_text_runs(p, stripped)
            continue

        # ── 본문 중간에 ㄴ 서식이 섞인 경우 ──
        mid_ann = re.match(r'^(.+?)\s+ㄴ\s+(.+)$', stripped)
        if mid_ann and _is_format_annotation('ㄴ ' + mid_ann.group(2)):
            content_part = mid_ann.group(1).strip()
            ann_part = 'ㄴ ' + mid_ann.group(2).strip()
            p = doc.add_paragraph()
            _add_text_runs(p, content_part)
            recent.append((p, content_part))
            mid_fmt = parse_annotation(ann_part)
            if mid_fmt.get('colored_words'):
                if all(w in content_part for w, _ in mid_fmt['colored_words']):
                    _apply_formatting_to_para(p, content_part, mid_fmt)
                else:
                    pending_fmts.append((mid_fmt, []))
            elif not mid_fmt['is_image_desc']:
                _apply_formatting_to_para(p, content_part, mid_fmt)
            ap = doc.add_paragraph()
            ann_display = _annotation_display_text(ann_part)
            ann_display = _augment_label_with_bold(ann_display, [content_part])
            run = ap.add_run(ann_display)
            run.bold = True
            run.font.color.rgb = GREEN
            run.font.size = Pt(24)
            run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
            recent.append((ap, ann_display))
            continue

        # ── ㄴ 서식 지시 줄 (서식 키워드가 있는 것만) ──
        if annotation_re.match(stripped) and _is_format_annotation(stripped):
            fmt = parse_annotation(stripped)

            if fmt['is_image_desc']:
                continue

            content_paras = [(p, t) for p, t in recent
                             if t.strip() and not (re.match(r'^ㄴ\s*', t.strip()) and _is_format_annotation(t.strip()))]
            target_count = fmt['multi_line']

            # Approach B — 의미 블록 스캔:
            # ㄴ 위로 빈 줄/직전 ㄴ 서식 라인까지 거슬러 올라가 한 블록을 모은다.
            # 블록 안에 `**..**` 단락이 하나라도 있으면 **볼드가 있는 단락만** 타겟(볼드 없는 일반 문장은 스킵).
            # 없으면 블록 마지막 한 단락만 타겟.
            # 이 로직은 사용자가 '두 줄 모두' 등 multi_line을 명시하지 않고 colored_words도 없을 때만 동작.
            # v2.1~ 타겟 산정 로직:
            # - 먼저 "연속 ㄴ 주석들을 건너뛰고 그 위 본문 그룹"을 수집 (content_group).
            # - search_words 있음: 그 그룹 안에서 각 단어의 첫 매칭 라인만 타겟.
            #   못 찾은 단어는 조용히 스킵 (fallback 없음).
            # - search_words 없음 (타겟 미지정): 그룹 안 **..** 볼드 단락만, 없으면 그룹 마지막 1줄만.
            search_words = (
                [w for w, _ in fmt.get('colored_words', [])]
                + list(fmt.get('target_words', []))
            )

            # content_group 수집: 현재 ㄴ 위의 "본문 문단" 수집.
            # v2.1.4~: 중간 ㄴ 주석은 경계 아닌 "단어 서식일 뿐"으로 취급 → 건너뛰고 위 본문까지 계속.
            # 문단 경계는 **빈 줄**만으로 판정.
            content_group = []
            for p_r, t_r in reversed(recent):
                t_s = (t_r or '').strip()
                is_ann = bool(re.match(r'^ㄴ\s*', t_s) and _is_format_annotation(t_s))
                if not t_s:
                    # 빈 줄: 수집 시작 전이면 skip, 이미 수집 중이면 중단 (문단 경계)
                    if content_group:
                        break
                    continue
                if is_ann:
                    # ㄴ 주석: 경계 아닌 단어 서식일 뿐. 건너뛰고 위 본문 계속 수집.
                    continue
                content_group.append((p_r, t_r))
            content_group.reverse()

            if target_count > 1:
                # multi_line 명시 (예: "두 줄 모두") — content_paras 기준 뒤에서 N줄
                targets = content_paras[-target_count:] if content_paras else []
            elif search_words:
                # 각 단어별로 content_group 안에서 매칭 라인 찾기
                matched = []
                seen_ids = set()
                for word in search_words:
                    for p_r, t_r in content_group:
                        clean = re.sub(r'\*+', '', t_r)
                        if word in clean:
                            if id(p_r) not in seen_ids:
                                matched.append((p_r, t_r))
                                seen_ids.add(id(p_r))
                            break
                targets = matched
            else:
                # 타겟 단어 미지정 ㄴ 주석 — 블록 전체 적용.
                # 편집창 JS가 여러 줄 선택 시 블록 위 빈 줄을 자동 삽입해 경계 확보하므로
                # 윗 문단 흡수는 발생하지 않음. **..** 명시 범위가 있으면 그것만 우선.
                if content_group:
                    bold_paras = [(p, t) for p, t in content_group
                                  if re.search(r'\*\*[^*]+\*\*', t.strip())]
                    targets = bold_paras if bold_paras else content_group
                else:
                    targets = []

            applied = False
            if search_words:
                if targets:
                    applied = True
                # else: 매칭된 단어가 하나도 없음 → 스킵 (fallback 없음)
            else:
                # 타겟 단어 미지정 ㄴ 주석 (예: ㄴ 파란색 단독)은 targets에 적용
                applied = True

            if applied:
                per_para_cw = _split_colored_words_across_targets(targets, fmt.get('colored_words', []))
                for idx, (para, para_text) in enumerate(targets):
                    if per_para_cw and idx in per_para_cw:
                        para_fmt = dict(fmt)
                        para_fmt['colored_words'] = per_para_cw[idx]
                        _apply_formatting_to_para(para, para_text, para_fmt)
                    else:
                        _apply_formatting_to_para(para, para_text, fmt)

                # v2.1.4~: 단어 타겟이 있어도 문단 전체 서식(full_text_color/highlight/underline)이
                # 함께 있으면 content_group의 비타겟 라인에도 문단 전체 서식만 별도 적용.
                # 예) `ㄴ '치' 파란색, 빨간색` → 매칭 라인은 치=파랑+나머지=빨강, 비매칭 라인은 전체 빨강.
                has_para_only_fmt = (
                    fmt.get('full_text_color')
                    or fmt.get('full_text_color_hex')
                    or fmt.get('highlight')
                    or fmt.get('underline')
                    or fmt.get('bold')
                )
                if search_words and has_para_only_fmt and content_group:
                    target_ids = {id(p) for p, _ in targets}
                    for para, para_text in content_group:
                        if id(para) in target_ids:
                            continue
                        bare_fmt = dict(fmt)
                        bare_fmt['colored_words'] = []
                        bare_fmt['highlighted_words'] = []
                        bare_fmt['bolded_words'] = []
                        bare_fmt['underlined_words'] = []
                        bare_fmt['target_words'] = []
                        _apply_formatting_to_para(para, para_text, bare_fmt)

            # ㄴ 줄 자체 → 초록색 주석 (헥스괄호 제거, 링크는 단독 치환)
            display = _annotation_display_text(stripped)
            display = _augment_label_with_bold(display, [t for _, t in targets])
            p = doc.add_paragraph()
            run = p.add_run(display)
            run.bold = True
            run.font.color.rgb = GREEN
            run.font.size = Pt(24)
            run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
            recent.append((p, display))
            continue

        # ── 이미지 번호 (00→0, 01→1, 02→2...) ──
        if image_num_re.match(stripped):
            display_num = str(int(stripped))
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(display_num)
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = BLUE
            recent.append((p, display_num))
            continue

        # ── 마크다운 헤딩 ──
        if stripped.startswith('### '):
            p = doc.add_heading(stripped.lstrip('# ').strip(), level=3)
            recent.append((p, stripped))
            continue
        if stripped.startswith('## '):
            p = doc.add_heading(stripped.lstrip('# ').strip(), level=2)
            recent.append((p, stripped))
            continue
        if stripped.startswith('# '):
            p = doc.add_heading(stripped.lstrip('# ').strip(), level=1)
            recent.append((p, stripped))
            continue

        # ── 일반 텍스트 ──
        p = doc.add_paragraph()
        _add_text_runs(p, stripped)
        recent.append((p, stripped))

        # ── 대기 중인 서식 적용 (ㄴ이 텍스트 위에 있었던 경우) ──
        if pending_fmts:
            new_pending = []
            for pfmt, collected in pending_fmts:
                collected.append((p, stripped))
                p_search_words = (
                    [w for w, _ in pfmt.get('colored_words', [])]
                    + list(pfmt.get('target_words', []))
                )
                if p_search_words:
                    all_text = ' '.join(t for _, t in collected)
                    if all(w in all_text for w in p_search_words):
                        per_para_cw = _split_colored_words_across_targets(collected, pfmt.get('colored_words', []))
                        for cidx, (cp, ct) in enumerate(collected):
                            if per_para_cw and cidx in per_para_cw:
                                p_fmt = dict(pfmt)
                                p_fmt['colored_words'] = per_para_cw[cidx]
                                _apply_formatting_to_para(cp, ct, p_fmt)
                            else:
                                _apply_formatting_to_para(cp, ct, pfmt)
                    elif len(collected) < 8:
                        new_pending.append((pfmt, collected))
                else:
                    _apply_formatting_to_para(p, stripped, pfmt)
            pending_fmts = new_pending

        if len(recent) > 15:
            recent = recent[-15:]

    # 남은 블로거 요청사항 처리
    if blogger_req_lines:
        _add_blogger_request_box(doc, blogger_req_lines)

    return doc


# ── 외부 인터페이스 ──
def save_as_docx(text, filepath):
    """텍스트 → 서식 적용된 .docx 파일 저장."""
    _build_document(text).save(filepath)


def build_docx_bytes_from_text(text, normalize=False):
    """텍스트 → 서식 적용된 .docx bytes (메모리 반환).

    기본 normalize=False: 편집창 = 최종본이라는 전제. 사용자가 수동으로 쪼갠
    줄바꿈을 그대로 보존. /generate 단계에서 이미 normalize_text로 한 번 처리했음.
    """
    from io import BytesIO
    buf = BytesIO()
    _build_document(text, normalize=normalize).save(buf)
    return buf.getvalue()


def normalize_text(text):
    """편집창 표시용 정규화 — 미리보기/워드 출력 결과와 동일한 줄바꿈 보장.

    /generate 응답에서 호출. 사용자가 편집창에서 보는 줄바꿈과
    미리보기 줄바꿈이 같아야 어디를 고쳐야 할지 가늠 가능.
    줄 쪼개기로 인해 `**...**` 가 토막 난 경우도 여기서 같이 복구.
    """
    if not text:
        return ''
    t = _normalize_line_lengths(text)
    t = _rewrap_multiline_bold(t)
    return t
