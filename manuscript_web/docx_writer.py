"""원고 → Word(.docx) 변환.

Phase E 출력(★블로거 요청사항★ + ㄴ 서식 지시)이면 docx_formatter 서식 엔진 사용.
그 외(Phase C만 있는 폴백)는 제목 + 본문 단순 단락 분리.
"""
import re
from io import BytesIO

from docx import Document
from docx.shared import Pt

from docx_formatter import build_docx_bytes_from_text


def _is_phase_e_text(body):
    """Phase E 완성본인지 판별 — ★블로거 요청사항★ 블록 존재 여부."""
    return '★' in (body or '')


def _sanitize_markup(text):
    """LLM이 섞어 쓴 HTML/확장 마크다운 문법 정리.

    docx_formatter가 인식하는 것만 남긴다 (**볼드** + ㄴ 주석 + 순수 텍스트).
    포매터 입력 직전 1회 실행. 포매터 자체는 건드리지 않는다.
    """
    if not text:
        return text

    # <br> → 줄바꿈
    text = re.sub(r"<\s*br\s*/?\s*>", "\n", text, flags=re.IGNORECASE)
    # HTML/XML 태그 제거 (내용물 유지) — <span>, <u>, <b>, <div>, <p>, <strong>, <em> 등
    text = re.sub(r"</?[a-zA-Z][^>]*>", "", text)
    # ~~취소선~~ → 내용만
    text = re.sub(r"~~([^~\n]+?)~~", r"\1", text)
    # ___단어___ → **단어** (3개 먼저 처리해야 2개 패턴과 충돌 안 함)
    text = re.sub(r"___([^_\n]+?)___", r"**\1**", text)
    # __단어__ → **단어**
    text = re.sub(r"__([^_\n]+?)__", r"**\1**", text)
    # ==단어== → **단어**
    text = re.sub(r"==([^=\n]+?)==", r"**\1**", text)

    return text


def build_docx_bytes(title, body):
    """제목 + 본문 → .docx bytes.

    Phase E 본문은 자체가 완전한 문서(★ 블록 + '제목 :' + ㄴ 지시)이므로
    title을 무시하고 body를 서식 엔진에 그대로 넘긴다.
    """
    body = (body or "").strip()

    if _is_phase_e_text(body):
        body = _sanitize_markup(body)
        return build_docx_bytes_from_text(body)

    # 폴백: Phase C만 있는 경우 — 단순 제목 + 단락
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '맑은 고딕'
    style.font.size = Pt(11)

    if title:
        doc.add_heading(title.strip(), level=1)

    if body:
        for block in re.split(r'\n\s*\n', body):
            b = block.strip()
            if b:
                doc.add_paragraph(b)
    else:
        doc.add_paragraph("(원고 없음)")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def save_docx(title, body, path):
    with open(path, 'wb') as f:
        f.write(build_docx_bytes(title, body))
    return path
