"""
원고 배정 프로그램 v1.3
- 발행리스트에서 업체별 배정 건수 확인 + 폴더 생성
- 원고리스트 ↔ 미배정 폴더 매칭 + 자동 검수
- 5일 키워드 겹침 체크 + 자동 배정
- 04.블로그 원고 모음 + 05.발행요청에 복사, 미배정 삭제
- 원고리스트 J열 파라미터 → 자사 발행리스트 H열 자동 기입
"""

import sys, os, re, shutil, threading, zipfile, json
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import datetime, timedelta
from pathlib import Path

import tkinter as tk
from tkinter import ttk, messagebox, scrolledtext

import gspread
from google.oauth2.service_account import Credentials
from docx import Document


# ─────────────────────────────────────────────
# 1. 설정 / 상수
# ─────────────────────────────────────────────
VERSION = "1.4"

SHEET_ID = "1jflcdbmBjQsY4hp8rNULXGGVqb64fGno4kudlTJoqM4"
KEYWORD_SHEET_ID = "1xJAogt0alaQ8A5OctxltPaF3kg_0PFSF5Z0ePxMw3tY"

# EXE 실행 시 EXE 옆, 아니면 기존 경로에서 credentials.json 찾기
if getattr(sys, 'frozen', False):
    _base_dir = os.path.dirname(sys.executable)
else:
    _base_dir = os.path.dirname(os.path.abspath(__file__))
CRED_FILE = os.path.join(_base_dir, "credentials.json")
if not os.path.exists(CRED_FILE):
    CRED_FILE = r"C:\Users\iamhy\Desktop\프로그램 개발\manuscript_generator\credentials.json"

EXCLUDE_FILE = os.path.join(_base_dir, "exclude_keywords.json")
MANUAL_PASS_FILE = os.path.join(_base_dir, "manual_pass.json")

BASE_PATH = r"G:\공유 드라이브\더넥스트\더넥스트\01.마케팅팀\01.바이럴 마케팅"
UNASSIGNED_PATH = os.path.join(BASE_PATH, "03.발행", "03.미발행", "02.미배정")
ARCHIVE_PATH = os.path.join(BASE_PATH, "03.발행", "04.블로그 원고 모음")
REQUEST_PATH = os.path.join(BASE_PATH, "03.발행", "05.발행요청")

PRODUCT_CODE_MAP = {
    "bc": ("블러드싸이클", "01"),
    "gc": ("혈당컷", "02"),
    "pt": ("판토오틴", "03"),
    "sc": ("상어연골환", "04"),
    "pf": ("활성엽산", "05"),
    "mt": ("멜라토닌", "06"),
    "hc": ("헬리컷", "07"),
    "gt": ("글루타치온", "09"),
}

COMPANY_FOLDERS = {
    "블로거": "01.블로거",
    "긍정애드": "02.긍정애드",
    "궁서": "03.궁서",
    "건배": "04.건배",
    "굳음마케팅": "05.굳음마케팅",
    "케이엘": "06.케이엘",
    "솔민랩스": "07.솔민랩스",
}

IGNORE_FILES = {"desktop.ini", "photothumb.db", "thumbs.db", ".ds_store"}
IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp"}

SKIP_FOLDERS = {
    "00.기존 원고_과대광고 전", "01.글램웰", "02.내부작가",
    "03.수정 원고", "04.AI 변환 원고", "수정 요청", "251216 원고",
}


# ─────────────────────────────────────────────
# 2. Google Sheets
# ─────────────────────────────────────────────
def connect_sheet(sheet_id, cred_path=CRED_FILE):
    scope = [
        "https://www.googleapis.com/auth/spreadsheets",
        "https://www.googleapis.com/auth/drive",
    ]
    creds = Credentials.from_service_account_file(cred_path, scopes=scope)
    gc = gspread.authorize(creds)
    return gc.open_by_key(sheet_id)


def load_publish_list(spreadsheet, target_date_str):
    ws = spreadsheet.worksheet("자사 발행리스트")
    rows = ws.get_all_values()
    result = {}
    for i, r in enumerate(rows[1:], start=2):
        date_val = r[0].strip() if len(r) > 0 else ""
        if date_val != target_date_str:
            continue
        company = r[13].strip() if len(r) > 13 else ""
        product = r[1].strip() if len(r) > 1 else ""
        keyword = r[4].strip() if len(r) > 4 else ""
        topic = r[5].strip() if len(r) > 5 else ""
        parameter = r[7].strip() if len(r) > 7 else ""
        author = r[10].strip() if len(r) > 10 else ""
        title = r[11].strip() if len(r) > 11 else ""
        category = r[27].strip() if len(r) > 27 else ""
        if company:
            result.setdefault(company, []).append({
                "product": product, "keyword": keyword, "row": i,
                "topic": topic, "parameter": parameter,
                "author": author, "title": title,
                "category": category, "company": company,
            })
    return result


def load_publish_keywords_5days(spreadsheet, target_date):
    """발행리스트에서 전후 5일간 업체별 키워드."""
    ws = spreadsheet.worksheet("자사 발행리스트")
    rows = ws.get_all_values()
    date_strs = set()
    for delta in range(-5, 6):
        d = target_date + timedelta(days=delta)
        date_strs.add(f"{d.month}/{d.day}")

    company_keywords = {}
    for r in rows[1:]:
        date_val = r[0].strip() if len(r) > 0 else ""
        if date_val not in date_strs:
            continue
        company = r[13].strip() if len(r) > 13 else ""
        keyword = r[4].strip() if len(r) > 4 else ""
        if company and keyword:
            company_keywords.setdefault(company, set()).add(keyword)
    return company_keywords


def load_manuscript_list(spreadsheet):
    ws = spreadsheet.worksheet("원고리스트")
    rows = ws.get_all_values()
    unassigned = []
    for i, r in enumerate(rows[1:], start=2):
        title = r[5].strip() if len(r) > 5 else ""
        pub_date = r[6].strip() if len(r) > 6 else ""
        if title and not pub_date:
            filename = r[32].strip() if len(r) > 32 else ""
            keyword = r[3].strip() if len(r) > 3 else ""
            parameter = r[9].strip() if len(r) > 9 else ""
            unassigned.append({
                "row": i,
                "submit_date": r[0].strip() if len(r) > 0 else "",
                "product": r[1].strip() if len(r) > 1 else "",
                "keyword": keyword,
                "title": title,
                "author": r[7].strip() if len(r) > 7 else "",
                "filename": filename,
                "parameter": parameter,
            })
    return unassigned


def update_publish_parameters(spreadsheet, param_updates):
    """자사 발행리스트 H열(8열)에 파라미터 기입."""
    ws = spreadsheet.worksheet("자사 발행리스트")
    for row_num, param_val in param_updates:
        cell = f"H{row_num}"
        ws.update(cell, [[param_val]], value_input_option="USER_ENTERED")


def update_keyword_sheet(cred_path, year, entries):
    """키워드 배정 시트 제품 탭에 데이터 추가.
    entries: [{"product", "keyword", "topic", "parameter", "author", "category", "title"}, ...]
    """
    ss = connect_sheet(KEYWORD_SHEET_ID, cred_path)
    by_product = {}
    for e in entries:
        by_product.setdefault(e["product"], []).append(e)

    results = []
    for product, items in by_product.items():
        try:
            ws = ss.worksheet(product)
        except gspread.exceptions.WorksheetNotFound:
            results.append(f"  ✗ '{product}' 탭 없음 ({len(items)}건 스킵)")
            continue

        rows = []
        for it in items:
            rows.append([
                year,               # A: 년도
                it["keyword"],      # B: 키워드
                it["topic"],        # C: 주제
                it["parameter"],    # D: 파라미터값
                it["author"],       # E: 작가
                it["category"],     # F: 분류
                it["title"],        # G: 제목
            ])
        ws.append_rows(rows, value_input_option="USER_ENTERED")
        results.append(f"  ✓ {product}: {len(items)}건 추가")
    return results



# ─────────────────────────────────────────────
# 3. 파일 시스템 작업
# ─────────────────────────────────────────────
def get_product_code(folder_name):
    parts = folder_name.rsplit("_", 1)
    return parts[-1].lower() if len(parts) > 1 else ""


def extract_keyword_from_folder(folder_name):
    parts = folder_name.split("_")
    if len(parts) >= 2:
        m = re.match(r"\d{6}(.+)", parts[1])
        if m:
            return m.group(1)
    return ""


def list_unassigned_folders():
    if not os.path.exists(UNASSIGNED_PATH):
        return []
    # ZIP 파일 자동 압축 해제
    for entry in os.scandir(UNASSIGNED_PATH):
        if entry.is_file() and entry.name.lower().endswith(".zip"):
            folder_name = entry.name[:-4]  # .zip 제거
            folder_path = os.path.join(UNASSIGNED_PATH, folder_name)
            if not os.path.exists(folder_path):
                try:
                    with zipfile.ZipFile(entry.path, "r") as zf:
                        zf.extractall(folder_path)
                    os.remove(entry.path)
                except Exception:
                    pass
    folders = []
    for entry in os.scandir(UNASSIGNED_PATH):
        if entry.is_dir() and entry.name not in SKIP_FOLDERS:
            if re.search(r"\d{6}", entry.name):
                folders.append(entry.name)
    return sorted(folders)


def create_date_folder(company, date_mmdd):
    folder_name = COMPANY_FOLDERS.get(company)
    if not folder_name:
        return None, f"알 수 없는 업체: {company}"
    path = os.path.join(REQUEST_PATH, folder_name, date_mmdd)
    os.makedirs(path, exist_ok=True)
    return path, None


def copy_to_archive(folder_name, product_code):
    info = PRODUCT_CODE_MAP.get(product_code)
    if not info:
        return None, f"알 수 없는 제품코드: {product_code}"
    product_name, folder_num = info
    m = re.search(r"(\d{6})", folder_name)
    month = m.group(1)[2:4] if m else f"{datetime.now().month:02d}"
    month_folder = f"{month}월"

    product_folder = None
    if os.path.exists(ARCHIVE_PATH):
        for entry in os.scandir(ARCHIVE_PATH):
            if entry.is_dir() and entry.name.startswith(f"{folder_num}."):
                product_folder = entry.name
                break
    if not product_folder:
        product_folder = f"{folder_num}.{product_name}"

    dest_dir = os.path.join(ARCHIVE_PATH, product_folder, month_folder)
    os.makedirs(dest_dir, exist_ok=True)
    src = os.path.join(UNASSIGNED_PATH, folder_name)
    dest = os.path.join(dest_dir, folder_name)
    if os.path.exists(dest):
        shutil.rmtree(dest)
    shutil.copytree(src, dest)
    return dest, None


def copy_to_request_as_zip(folder_name, company, date_mmdd):
    company_folder = COMPANY_FOLDERS.get(company)
    if not company_folder:
        return None, f"알 수 없는 업체: {company}"
    dest_dir = os.path.join(REQUEST_PATH, company_folder, date_mmdd)
    os.makedirs(dest_dir, exist_ok=True)
    src = os.path.join(UNASSIGNED_PATH, folder_name)
    zip_path = os.path.join(dest_dir, f"{folder_name}.zip")

    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for root, dirs, files in os.walk(src):
            for f in files:
                if f.lower() in IGNORE_FILES:
                    continue
                full = os.path.join(root, f)
                arcname = os.path.relpath(full, src)
                zf.write(full, arcname)
    return zip_path, None


def delete_from_unassigned(folder_name):
    path = os.path.join(UNASSIGNED_PATH, folder_name)
    if os.path.exists(path):
        shutil.rmtree(path)


# ─────────────────────────────────────────────
# 4. 원고 검수
# ─────────────────────────────────────────────
def review_manuscript(folder_name):
    folder_path = os.path.join(UNASSIGNED_PATH, folder_name)
    result = {
        "folder": folder_name,
        "passed": True,
        "image_ok": [],
        "image_issues": [],
        "link_ok": [],
        "link_issues": [],
        "product_image_ok": [],
        "product_image_issues": [],
        "errors": [],
    }

    docx_files = list(Path(folder_path).glob("*.docx"))
    if not docx_files:
        result["errors"].append("docx 파일 없음")
        result["passed"] = False
        return result

    try:
        doc = Document(str(docx_files[0]))
    except Exception as e:
        result["errors"].append(f"docx 읽기 실패: {e}")
        result["passed"] = False
        return result

    paragraphs = [(i, p.text.strip()) for i, p in enumerate(doc.paragraphs)]

    image_files = []
    for f in os.listdir(folder_path):
        if f.lower() in IGNORE_FILES or f.lower().endswith(".docx"):
            continue
        if Path(f).suffix.lower() in IMAGE_EXTS:
            image_files.append(f)

    # 이미지 체크
    referenced_nums = set()
    for idx, text in paragraphs:
        # 쉼표 포함 숫자(1,000 / 2,000 등)는 이미지 번호가 아님 — 제거 후 판단
        cleaned = re.sub(r"\d{1,3}(,\d{3})+", "", text).strip()
        if re.match(r"^\d+$", cleaned):
            if int(cleaned) <= 30:
                referenced_nums.add(cleaned)
        elif re.match(r"^\d+[\s、]+\d+", cleaned):
            for m in re.finditer(r"\d+", cleaned):
                if int(m.group()) <= 30:
                    referenced_nums.add(m.group())

    file_nums = {}
    for f in image_files:
        name = Path(f).stem
        if re.match(r"^[a-z]{2}\d+$", name.lower()):
            continue
        m = re.match(r"^(\d+)", name)
        if m:
            file_nums[m.group(1)] = f

    for num in sorted(referenced_nums, key=lambda x: int(x)):
        if num in file_nums:
            result["image_ok"].append(f"이미지 {num}번({file_nums[num]}): OK")
        else:
            result["image_issues"].append(f"이미지 {num}번: docx에 있으나 폴더에 파일 없음")

    for num, fname in sorted(file_nums.items(), key=lambda x: int(x[0])):
        if num not in referenced_nums:
            result["image_issues"].append(f"이미지 {num}번({fname}): 폴더에 있으나 docx에 참조 없음")

    # 링크 체크
    link_pattern = r"https://mkt\.shopping\.naver\.com/link/"
    link_ok_keywords = ["링크 도구", "링크도구", "링크 배너", "링크배너", "링크 삽입", "링크삽입"]
    for idx, text in paragraphs:
        if re.search(link_pattern, text):
            next_text = ""
            for j in range(idx + 1, len(doc.paragraphs)):
                nt = doc.paragraphs[j].text.strip()
                if nt:
                    next_text = nt
                    break
            if any(kw in next_text for kw in link_ok_keywords):
                result["link_ok"].append(
                    f"링크({text}): OK ({next_text[:30]})")
            else:
                result["link_issues"].append(
                    f"링크({text}): 링크 삽입 멘트 없음"
                )

    # 광고 이미지 체크
    product_img_refs = set()
    for idx, text in paragraphs:
        m = re.match(r"^([a-z]{2}\d+)", text.lower())
        if m:
            product_img_refs.add(m.group(1))

    for ref in product_img_refs:
        matching = [f for f in image_files if Path(f).stem.lower() == ref]
        if matching:
            result["product_image_ok"].append(
                f"광고이미지 {ref}({matching[0]}): OK")
        else:
            result["product_image_issues"].append(
                f"광고이미지 {ref}: docx에 있으나 폴더에 파일 없음"
            )

    if result["image_issues"] or result["link_issues"] or result["product_image_issues"]:
        result["passed"] = False

    return result


# ─────────────────────────────────────────────
# 5. 유틸리티
# ─────────────────────────────────────────────
def date_to_md(d):
    return f"{d.month}/{d.day}"

def date_to_mmdd(d):
    return f"{d.month:02d}{d.day:02d}"

def date_to_yymmdd(d):
    return d.strftime("%y%m%d")

def normalize_keyword(kw):
    kw = re.sub(r"\(\d+\)$", "", kw)
    kw = re.sub(r"\s+", "", kw)
    return kw.lower().strip()

def load_exclusions():
    if os.path.exists(EXCLUDE_FILE):
        try:
            with open(EXCLUDE_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            pass
    return {}


def save_exclusions(data):
    with open(EXCLUDE_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def load_manual_pass():
    if os.path.exists(MANUAL_PASS_FILE):
        try:
            with open(MANUAL_PASS_FILE, "r", encoding="utf-8") as f:
                return set(json.load(f))
        except Exception:
            pass
    return set()


def save_manual_pass(data):
    with open(MANUAL_PASS_FILE, "w", encoding="utf-8") as f:
        json.dump(sorted(data), f, ensure_ascii=False, indent=2)


def check_exclusion(exclusions, company, title, folder_name):
    """업체의 제외 키워드가 제목이나 폴더명에 포함되면 해당 키워드 반환."""
    kw_list = exclusions.get(company, [])
    text = (title + " " + folder_name).lower()
    for kw in kw_list:
        if kw.lower() in text:
            return kw
    return None


def keywords_overlap(kw1, kw2):
    n1 = normalize_keyword(kw1)
    n2 = normalize_keyword(kw2)
    if not n1 or not n2:
        return False
    return n1 == n2 or n1 in n2 or n2 in n1


# ─────────────────────────────────────────────
# 6. GUI 앱
# ─────────────────────────────────────────────
class ManuscriptAssignerApp:
    def __init__(self, root):
        self.root = root
        self.root.title(f"원고 배정 프로그램 v{VERSION}")
        self.root.geometry("1300x850")
        self.root.minsize(1000, 700)

        self.spreadsheet = None
        self.current_step = 0

        self.target_date = datetime.now().date() + timedelta(days=1)
        self.publish_data = {}
        self.company_keywords_5d = {}
        self.matched = []
        self.review_results = {}
        self.assignments = {}
        self.publish_row_map = {}
        self.param_updates = []
        self.written_rows = []  # Step 1 시트 기입 행 번호 (취소용)
        self._write_stop = False  # Step 1 시트 기입 중지 플래그
        self.exclusions = load_exclusions()

        self._setup_styles()
        self._build_ui()
        self.root.after(100, self._connect_sheet)

    def _setup_styles(self):
        style = ttk.Style()
        style.theme_use("clam")
        default_font = ("맑은 고딕", 10)
        style.configure(".", font=default_font)
        style.configure("Treeview", font=default_font, rowheight=26)
        style.configure("Treeview.Heading", font=("맑은 고딕", 10, "bold"))
        style.configure("Title.TLabel", font=("맑은 고딕", 14, "bold"))
        style.configure("Step.TLabel", font=("맑은 고딕", 10), padding=5)
        style.configure("CurrentStep.TLabel", font=("맑은 고딕", 10, "bold"),
                        foreground="white", background="#2196F3", padding=5)
        style.configure("DoneStep.TLabel", font=("맑은 고딕", 10),
                        foreground="white", background="#4CAF50", padding=5)

    def _build_ui(self):
        top = ttk.Frame(self.root, padding=10)
        top.pack(fill="x")

        ttk.Label(top, text="원고 배정 프로그램", style="Title.TLabel").pack(side="left")

        date_frame = ttk.Frame(top)
        date_frame.pack(side="left", padx=30)
        ttk.Label(date_frame, text="배정 날짜:").pack(side="left")
        self.date_var = tk.StringVar(value=date_to_md(self.target_date))
        ttk.Entry(date_frame, textvariable=self.date_var, width=8).pack(side="left", padx=5)
        ttk.Button(date_frame, text="변경", command=self._change_date).pack(side="left")

        ttk.Button(top, text="제외 키워드 설정",
                   command=self._open_exclusion_dialog).pack(side="right", padx=5)
        ttk.Button(top, text="🔄 새로고침", command=self._refresh).pack(side="right", padx=10)
        self.status_var = tk.StringVar(value="시트 연결 중...")
        ttk.Label(top, textvariable=self.status_var).pack(side="right")

        self.step_frame = ttk.Frame(self.root, padding=(10, 5))
        self.step_frame.pack(fill="x")

        self.step_names = [
            "1. 배정 건수",
            "2. 매칭 + 검수",
            "3. 자동 배정",
            "4. 실행",
        ]
        self.step_labels = []
        for name in self.step_names:
            lbl = ttk.Label(self.step_frame, text=name, style="Step.TLabel")
            lbl.pack(side="left", padx=2)
            self.step_labels.append(lbl)

        ttk.Separator(self.root, orient="horizontal").pack(fill="x", pady=5)

        # 하단 먼저 pack (expand 콘텐츠에 밀리지 않도록)
        bottom = ttk.Frame(self.root, padding=10)
        bottom.pack(fill="x", side="bottom")

        self.content_frame = ttk.Frame(self.root, padding=10)
        self.content_frame.pack(fill="both", expand=True)
        self.prev_btn = ttk.Button(bottom, text="◀ 이전", command=self._prev_step)
        self.prev_btn.pack(side="left")
        self.cancel_write_btn = ttk.Button(
            bottom, text="⚠ 시트 기입 취소",
            command=self._cancel_written_rows, state="disabled")
        self.cancel_write_btn.pack(side="left", padx=20)
        self.next_btn = ttk.Button(bottom, text="다음 ▶", command=self._next_step)
        self.next_btn.pack(side="right")

    def _update_step_indicator(self):
        for i, lbl in enumerate(self.step_labels):
            if i < self.current_step:
                lbl.configure(style="DoneStep.TLabel")
            elif i == self.current_step:
                lbl.configure(style="CurrentStep.TLabel")
            else:
                lbl.configure(style="Step.TLabel")

    def _clear_content(self):
        for w in self.content_frame.winfo_children():
            w.destroy()

    def _change_date(self):
        val = self.date_var.get().strip()
        try:
            parts = val.split("/")
            month, day = int(parts[0]), int(parts[1])
            self.target_date = datetime(datetime.now().year, month, day).date()
            self.status_var.set(f"날짜 변경: {date_to_md(self.target_date)}")
        except Exception:
            messagebox.showerror("오류", "날짜 형식: M/D (예: 3/18)")

    def _connect_sheet(self):
        def worker():
            try:
                self.spreadsheet = connect_sheet(SHEET_ID)
                self.root.after(0, lambda: self.status_var.set("시트 연결 완료"))
                self.root.after(0, lambda: self._go_to_step(0))
            except Exception as e:
                self.root.after(0, lambda: self.status_var.set(f"시트 연결 실패: {e}"))
        threading.Thread(target=worker, daemon=True).start()

    def _refresh(self):
        self.status_var.set("새로고침 중...")
        step = self.current_step
        def worker():
            try:
                self.spreadsheet = connect_sheet(SHEET_ID)
                self.root.after(0, lambda: self.status_var.set("새로고침 완료"))
                self.root.after(0, lambda: self._go_to_step(step))
            except Exception as e:
                self.root.after(0, lambda: self.status_var.set(f"새로고침 실패: {e}"))
        threading.Thread(target=worker, daemon=True).start()

    # ── 스텝 이동 ──
    def _go_to_step(self, step):
        self.current_step = step
        self._update_step_indicator()
        self._clear_content()
        self.prev_btn.configure(state="normal" if step > 0 else "disabled")
        self.next_btn.configure(state="normal", text="다음 ▶")

        builders = [
            self._build_step1,
            self._build_step2,
            self._build_step3,
            self._build_step4,
        ]
        if 0 <= step < len(builders):
            builders[step]()

    def _prev_step(self):
        if self.current_step > 0:
            self._go_to_step(self.current_step - 1)

    def _next_step(self):
        if self.current_step < 3:
            self._go_to_step(self.current_step + 1)

    # ──────────────────────────────────────
    # Step 1: 배정 건수 확인 + 폴더 생성
    # ──────────────────────────────────────
    def _build_step1(self):
        f = self.content_frame
        ttk.Label(f, text="Step 1. 업체별 배정 건수 설정",
                  style="Title.TLabel").pack(anchor="w", pady=(0, 10))

        # ── 상단: 발행 계획 직접 설정 ──
        plan_frame = ttk.LabelFrame(f, text="발행 계획 설정", padding=10)
        plan_frame.pack(fill="x", pady=(0, 5))

        self.company_counts = {}
        row_frame = ttk.Frame(plan_frame)
        row_frame.pack(fill="x")
        for i, (company, folder) in enumerate(COMPANY_FOLDERS.items()):
            col = i % 4
            r = i // 4
            cf = ttk.Frame(row_frame)
            cf.grid(row=r, column=col, padx=10, pady=3, sticky="w")
            ttk.Label(cf, text=company, width=10).pack(side="left")
            var = tk.IntVar(value=0)
            self.company_counts[company] = var
            ttk.Spinbox(cf, from_=0, to=30, width=4, textvariable=var).pack(side="left", padx=5)

        plan_btn_frame = ttk.Frame(plan_frame)
        plan_btn_frame.pack(fill="x", pady=(8, 0))
        ttk.Button(plan_btn_frame, text="시트 기입 + 불러오기",
                   command=self._write_plan_and_load).pack(side="left")
        self.stop_write_btn = ttk.Button(plan_btn_frame, text="■ 중지",
                                          command=self._stop_write, state="disabled")
        self.stop_write_btn.pack(side="left", padx=5)
        self.plan_status = tk.StringVar(value="")
        ttk.Label(plan_btn_frame, textvariable=self.plan_status).pack(side="left", padx=10)

        # ── 중단: 폴더 생성 ──
        btn_frame = ttk.Frame(f)
        btn_frame.pack(fill="x", pady=5)
        self.folder_btn = ttk.Button(btn_frame, text="날짜 폴더 생성",
                                     command=self._create_folders, state="disabled")
        self.folder_btn.pack(side="left", padx=10)

        tree_frame = ttk.Frame(f)
        tree_frame.pack(fill="both", expand=True)

        cols = ("company", "count", "products")
        self.pub_tree = ttk.Treeview(tree_frame, columns=cols, show="headings", height=8)
        self.pub_tree.heading("company", text="발행처")
        self.pub_tree.heading("count", text="건수")
        self.pub_tree.heading("products", text="제품(키워드) 내역")
        self.pub_tree.column("company", width=120)
        self.pub_tree.column("count", width=60, anchor="center")
        self.pub_tree.column("products", width=700)
        sb = ttk.Scrollbar(tree_frame, orient="vertical", command=self.pub_tree.yview)
        self.pub_tree.configure(yscrollcommand=sb.set)
        self.pub_tree.pack(side="left", fill="both", expand=True)
        sb.pack(side="right", fill="y")

        self.step1_log = scrolledtext.ScrolledText(f, height=5, state="disabled",
                                                    font=("맑은 고딕", 9))
        self.step1_log.pack(fill="x", pady=(5, 0))

    def _stop_write(self):
        """시트 기입 중지."""
        self._write_stop = True
        self.stop_write_btn.configure(state="disabled")
        self.plan_status.set("중지 요청됨...")

    def _write_plan_and_load(self):
        """발행 계획 기입 + 발행리스트 불러오기 통합.
        건수가 모두 0이면 기입 없이 불러오기만 실행."""
        plan = {}
        for company, var in self.company_counts.items():
            count = var.get()
            if count > 0:
                plan[company] = count

        # 건수 0 → 기입 없이 바로 불러오기
        if not plan:
            self._load_publish_data()
            return

        total = sum(plan.values())
        msg = "\n".join(f"  {c}: {n}건" for c, n in plan.items())
        if not messagebox.askyesno("확인",
            f"[{date_to_md(self.target_date)}] 총 {total}건\n{msg}\n\n"
            f"자사 발행리스트에 기입합니다."):
            return

        self._write_stop = False
        self.plan_status.set("시트 확인 중...")
        self.stop_write_btn.configure(state="normal")
        date_str = date_to_md(self.target_date)

        def worker():
            try:
                ws = self.spreadsheet.worksheet("자사 발행리스트")
                # A열+N열만 읽어서 해당 날짜 기존 데이터 확인
                a_col = ws.col_values(1)   # A열
                n_col = ws.col_values(14)  # N열
                existing_count = sum(1 for a, n in zip(a_col[1:], n_col[1:])
                                     if a.strip() == date_str and n.strip())
                if existing_count > 0:
                    self.root.after(0, lambda: self.plan_status.set(""))
                    self.root.after(0, lambda: self.stop_write_btn.configure(state="disabled"))
                    def _ask_existing():
                        if messagebox.askyesno("경고",
                            f"[{date_str}] 날짜에 이미 {existing_count}건이 있습니다.\n"
                            f"기입 없이 기존 데이터를 불러올까요?"):
                            self._load_publish_data()
                    self.root.after(0, _ask_existing)
                    return

                # H열 기준 마지막 데이터 행 찾기
                h_col = ws.col_values(8)  # H열 전체
                last_row = len(h_col)
                for i in range(len(h_col) - 1, -1, -1):
                    if h_col[i].strip():
                        last_row = i + 1
                        break
                start_row = last_row + 1

                # A열(날짜)과 N열(발행처)만 개별 기입 (다른 열 수식 보호)
                row_idx = start_row
                written_count = 0
                for company, count in plan.items():
                    for _ in range(count):
                        if self._write_stop:
                            break
                        ws.update(f"A{row_idx}", [[date_str]], value_input_option="USER_ENTERED")
                        ws.update(f"N{row_idx}", [[company]], value_input_option="USER_ENTERED")
                        row_idx += 1
                        written_count += 1
                        self.root.after(0, lambda wc=written_count:
                            self.plan_status.set(f"기입 중... {wc}/{total}건"))
                    if self._write_stop:
                        break

                self.written_rows = list(range(start_row, row_idx))
                self.root.after(0, lambda: self.stop_write_btn.configure(state="disabled"))

                if self._write_stop:
                    # 중지됨 — 이미 기입된 행만 기록
                    self.root.after(0, lambda:
                        self.cancel_write_btn.configure(state="normal"))
                    self.root.after(0, lambda wc=written_count:
                        self.plan_status.set(f"⚠ 중지됨 ({wc}/{total}건 기입)"))
                    self.root.after(0, lambda wc=written_count:
                        self._log(self.step1_log,
                            f"발행 계획 기입 중지: {date_str} {wc}/{total}건"))
                else:
                    self.root.after(0, lambda:
                        self.cancel_write_btn.configure(state="normal"))
                    self.root.after(0, lambda:
                        self.plan_status.set(f"✓ {total}건 기입 완료"))
                    self.root.after(0, lambda:
                        self._log(self.step1_log,
                            f"발행 계획 기입: {date_str} 총 {total}건\n{msg}"))
                    # 기입 완료 후 자동으로 발행리스트 불러오기
                    self.root.after(500, self._load_publish_data)
            except Exception as e:
                self.root.after(0, lambda: self.stop_write_btn.configure(state="disabled"))
                self.root.after(0, lambda:
                    self.plan_status.set(f"✗ 오류: {e}"))
        threading.Thread(target=worker, daemon=True).start()

    def _cancel_written_rows(self):
        """Step 1에서 기입한 시트 행을 삭제하여 취소."""
        rows = self.written_rows
        if not rows:
            messagebox.showinfo("알림", "취소할 시트 기입 내역이 없습니다.")
            return

        if not messagebox.askyesno("⚠ 시트 기입 취소",
            f"자사 발행리스트에 기입한 {len(rows)}행을 삭제합니다.\n\n"
            f"정말 취소하시겠습니까?"):
            return

        self.cancel_write_btn.configure(state="disabled")
        self.status_var.set("시트 기입 취소 중...")

        def worker():
            try:
                ws = self.spreadsheet.worksheet("자사 발행리스트")
                for row_num in sorted(rows, reverse=True):
                    ws.delete_rows(row_num)
                self.written_rows = []
                self.publish_data = {}
                self.root.after(0, lambda:
                    self.status_var.set(f"✓ {len(rows)}행 삭제 완료 (시트 기입 취소)"))
                self.root.after(0, lambda:
                    self.plan_status.set("시트 기입 취소됨"))
            except Exception as e:
                self.root.after(0, lambda:
                    self.cancel_write_btn.configure(state="normal"))
                self.root.after(0, lambda:
                    self.status_var.set(f"✗ 시트 기입 취소 오류: {e}"))
        threading.Thread(target=worker, daemon=True).start()

    def _load_publish_data(self):
        self.status_var.set("발행리스트 불러오는 중...")
        target = date_to_md(self.target_date)

        def worker():
            try:
                data = load_publish_list(self.spreadsheet, target)
                kw5d = load_publish_keywords_5days(self.spreadsheet, self.target_date)
                self.publish_data = data
                self.company_keywords_5d = kw5d
                self.root.after(0, lambda: self._display_publish_data(data, kw5d))
            except Exception as e:
                self.root.after(0, lambda: self.status_var.set(f"오류: {e}"))
        threading.Thread(target=worker, daemon=True).start()

    def _display_publish_data(self, data, kw5d):
        self.pub_tree.delete(*self.pub_tree.get_children())
        total = 0
        for company, items in sorted(data.items()):
            count = len(items)
            total += count
            products = ", ".join(f"{it['product']}({it['keyword']})" for it in items)
            self.pub_tree.insert("", "end", values=(company, count, products))

        self.status_var.set(f"총 {total}건 ({len(data)}개 업체)")
        self.folder_btn.configure(state="normal")

        log = f"[{date_to_md(self.target_date)}] 총 {total}건\n"
        for company, kws in sorted(kw5d.items()):
            log += f"  {company} 5일 키워드({len(kws)}개): {', '.join(list(kws)[:10])}\n"
        self._log(self.step1_log, log)

    def _create_folders(self):
        date_mmdd = date_to_mmdd(self.target_date)
        created = []
        errors = []
        for company in self.publish_data:
            path, err = create_date_folder(company, date_mmdd)
            if err:
                errors.append(err)
            else:
                created.append((company, path))

        log = f"\n{'='*50}\n폴더 생성 결과 ({len(created)}개)\n{'='*50}\n"
        for company, path in created:
            log += f"\n✓ {company}/{date_mmdd}\n"
            log += f"  경로: {path}\n"
            try:
                existing = [f for f in os.listdir(path)
                            if f.lower() not in IGNORE_FILES]
                if existing:
                    log += f"  기존 파일({len(existing)}): {', '.join(existing[:5])}\n"
                else:
                    log += f"  (비어 있음)\n"
            except Exception:
                pass
        if errors:
            log += f"\n오류: {', '.join(errors)}\n"
        self._log(self.step1_log, log)
        self.status_var.set(f"폴더 생성 완료: {len(created)}개")

    # ──────────────────────────────────────
    # Step 2: 매칭 + 검수 (한 번에)
    # ──────────────────────────────────────
    def _build_step2(self):
        f = self.content_frame
        ttk.Label(f, text="Step 2. 원고 매칭 + 검수",
                  style="Title.TLabel").pack(anchor="w", pady=(0, 10))

        top_bar = ttk.Frame(f)
        top_bar.pack(fill="x", pady=5)
        ttk.Button(top_bar, text="매칭 + 검수 실행",
                   command=self._match_and_review).pack(side="left")
        self.step2_progress = tk.StringVar(value="")
        ttk.Label(top_bar, textvariable=self.step2_progress).pack(side="left", padx=15)

        # 상세 보기 + 수동 통과 버튼 (하단, 먼저 pack)
        bottom_frame = ttk.Frame(f)
        bottom_frame.pack(fill="x", pady=(5, 0), side="bottom")

        self.pass_btn = ttk.Button(bottom_frame, text="선택 항목 검수 통과 처리",
                                    command=self._manual_pass, state="disabled")
        self.pass_btn.pack(anchor="e", pady=(0, 3))

        self.review_detail = scrolledtext.ScrolledText(
            bottom_frame, height=7, state="disabled", font=("맑은 고딕", 9))
        self.review_detail.pack(fill="x")

        # 트리뷰
        tree_frame = ttk.Frame(f)
        tree_frame.pack(fill="both", expand=True)

        cols = ("status", "author", "product", "keyword", "filename",
                "img", "link", "ad_img", "review")
        self.step2_tree = ttk.Treeview(tree_frame, columns=cols,
                                        show="headings", height=14)
        for col, text in [("status","매칭"), ("author","작가"), ("product","제품"),
                          ("keyword","키워드"), ("filename","파일명"), ("img","이미지"),
                          ("link","링크"), ("ad_img","광고"), ("review","검수")]:
            self.step2_tree.heading(col, text=text,
                command=lambda c=col: self._sort_step2_tree(c))
        self.step2_tree.column("status", width=60, anchor="center")
        self.step2_tree.column("author", width=60)
        self.step2_tree.column("product", width=80)
        self.step2_tree.column("keyword", width=120)
        self.step2_tree.column("filename", width=380)
        self.step2_tree.column("img", width=55, anchor="center")
        self.step2_tree.column("link", width=55, anchor="center")
        self.step2_tree.column("ad_img", width=45, anchor="center")
        self.step2_tree.column("review", width=55, anchor="center")

        self.step2_tree.tag_configure("pass", background="#E8F5E9")
        self.step2_tree.tag_configure("fail", background="#FFEBEE")
        self.step2_tree.tag_configure("drive_only", background="#FFF3E0")
        self.step2_tree.tag_configure("sheet_only", background="#F5F5F5")

        sb = ttk.Scrollbar(tree_frame, orient="vertical",
                           command=self.step2_tree.yview)
        self.step2_tree.configure(yscrollcommand=sb.set)
        self.step2_tree.pack(side="left", fill="both", expand=True)
        sb.pack(side="right", fill="y")

        self.step2_tree.bind("<<TreeviewSelect>>", self._show_review_detail)
        self.step2_tree.bind("<Double-1>", self._open_manuscript_folder)

        # Step 2 진입 시 자동으로 매칭+검수 실행
        self.root.after(100, self._match_and_review)

    def _match_and_review(self):
        """매칭 → 매칭된 건 자동 검수."""
        self.status_var.set("매칭 + 검수 중...")
        self.step2_progress.set("시트 불러오는 중...")

        def worker():
            try:
                # 1) 매칭
                ms_list = load_manuscript_list(self.spreadsheet)
                folders = list_unassigned_folders()
                folder_set = set(folders)

                matched = []
                sheet_only = []

                for ms in ms_list:
                    fname = ms["filename"]
                    if fname in folder_set:
                        matched.append({**ms, "folder_name": fname,
                                        "match_status": "매칭"})
                    elif fname:
                        found = False
                        for fn in folders:
                            if fn.startswith(fname):
                                matched.append({**ms, "folder_name": fn,
                                                "match_status": "매칭"})
                                found = True
                                break
                        if not found:
                            sheet_only.append({**ms, "folder_name": "",
                                               "match_status": "시트만"})

                matched_folders = {m["folder_name"] for m in matched}
                drive_only = []
                for fn in folders:
                    if fn not in matched_folders:
                        code = get_product_code(fn)
                        info = PRODUCT_CODE_MAP.get(code, ("", ""))
                        author = fn.split("_")[0] if "_" in fn else ""
                        kw = extract_keyword_from_folder(fn)
                        drive_only.append({
                            "row": None, "submit_date": "", "product": info[0],
                            "keyword": kw, "title": "", "author": author,
                            "filename": fn, "folder_name": fn,
                            "match_status": "폴더만",
                        })

                all_items = matched + drive_only + sheet_only
                self.matched = all_items

                self.root.after(0, lambda: self.step2_progress.set(
                    f"매칭 완료({len(matched)}건), 검수 중..."))

                # 2) 매칭된 건 + 폴더만 있는 건 검수 (병렬)
                self.review_results.clear()
                saved_pass = load_manual_pass()
                to_review = [m for m in all_items if m.get("folder_name")]
                total = len(to_review)
                done = [0]

                def _review(ms):
                    folder = ms["folder_name"]
                    result = review_manuscript(folder)
                    # 기인정: 이전에 수동 통과 저장된 폴더
                    if not result["passed"] and folder in saved_pass:
                        result["passed"] = True
                        result["pre_approved"] = True
                    done[0] += 1
                    self.root.after(0, lambda d=done[0]:
                        self.step2_progress.set(f"검수 중... {d}/{total}"))
                    return folder, result

                with ThreadPoolExecutor(max_workers=8) as pool:
                    futures = [pool.submit(_review, ms) for ms in to_review]
                    for fut in as_completed(futures):
                        folder, result = fut.result()
                        self.review_results[folder] = result

                # 3) 결과 표시
                self.root.after(0, lambda: self._display_step2(
                    all_items, len(matched), len(drive_only), len(sheet_only)))

            except Exception as e:
                self.root.after(0, lambda: self.status_var.set(f"오류: {e}"))

        threading.Thread(target=worker, daemon=True).start()

    def _display_step2(self, all_items, n_matched, n_drive, n_sheet):
        self.step2_tree.delete(*self.step2_tree.get_children())

        passed = 0
        for m in all_items:
            folder = m.get("folder_name", "")
            review = self.review_results.get(folder)

            if m["match_status"] == "시트만":
                tag = "sheet_only"
                img_s = link_s = ad_s = review_s = "-"
            elif review:
                img_s = "OK" if not review["image_issues"] else f"✗{len(review['image_issues'])}"
                link_s = "OK" if not review["link_issues"] else f"✗{len(review['link_issues'])}"
                ad_s = "OK" if not review["product_image_issues"] else f"✗{len(review['product_image_issues'])}"
                if review["errors"]:
                    review_s = "오류"
                    tag = "fail"
                elif review["passed"]:
                    review_s = "기인정" if review.get("pre_approved") else "통과"
                    tag = "pass"
                    passed += 1
                else:
                    review_s = "실패"
                    tag = "fail"
            else:
                img_s = link_s = ad_s = review_s = "-"
                tag = "drive_only" if m["match_status"] == "폴더만" else "sheet_only"

            display_name = folder or m["filename"]
            self.step2_tree.insert("", "end", values=(
                m["match_status"], m["author"], m["product"],
                m.get("keyword", ""), display_name,
                img_s, link_s, ad_s, review_s
            ), tags=(tag,))

        total_reviewed = len(self.review_results)
        self.step2_progress.set(
            f"매칭: {n_matched} | 폴더만: {n_drive} | 시트만: {n_sheet} | "
            f"검수 통과: {passed}/{total_reviewed}")
        self.status_var.set(f"매칭+검수 완료: 통과 {passed}/{total_reviewed}건")

    def _sort_step2_tree(self, col):
        """Step 2 트리뷰 컬럼 클릭 시 정렬."""
        items = [(self.step2_tree.set(k, col), k)
                 for k in self.step2_tree.get_children("")]
        reverse = getattr(self, '_step2_sort_reverse', False)
        items.sort(key=lambda x: x[0], reverse=reverse)
        for idx, (_, k) in enumerate(items):
            self.step2_tree.move(k, "", idx)
        self._step2_sort_reverse = not reverse

    def _open_manuscript_folder(self, event=None):
        """더블클릭 시 해당 원고 폴더를 탐색기로 열기."""
        item = self.step2_tree.focus()
        if not item:
            return
        vals = self.step2_tree.item(item, "values")
        folder_name = vals[4]  # filename column
        folder_path = os.path.join(UNASSIGNED_PATH, folder_name)
        if os.path.isdir(folder_path):
            os.startfile(folder_path)
        else:
            messagebox.showwarning("폴더 없음", f"폴더를 찾을 수 없습니다:\n{folder_name}")

    def _show_review_detail(self, event=None):
        item = self.step2_tree.focus()
        if not item:
            return
        vals = self.step2_tree.item(item, "values")
        folder = vals[4]  # filename column
        result = self.review_results.get(folder)

        # 검수 실패 항목이면 통과 버튼 활성화
        if result and not result["passed"]:
            self.pass_btn.configure(state="normal")
        else:
            self.pass_btn.configure(state="disabled")

        self.review_detail.configure(state="normal")
        self.review_detail.delete("1.0", "end")

        if not result:
            self.review_detail.insert("1.0", f"{folder}: 검수 대상 아님")
            self.review_detail.configure(state="disabled")
            return

        lines = [f"=== {folder} ==="]
        if result["errors"]:
            lines.append(f"[오류] {', '.join(result['errors'])}")

        # 이미지
        if result.get("image_ok") or result["image_issues"]:
            lines.append("[이미지]")
            for ok in result.get("image_ok", []):
                lines.append(f"  ✓ {ok}")
            for issue in result["image_issues"]:
                lines.append(f"  ✗ {issue}")

        # 링크
        if result.get("link_ok") or result["link_issues"]:
            lines.append("[링크]")
            for ok in result.get("link_ok", []):
                lines.append(f"  ✓ {ok}")
            for issue in result["link_issues"]:
                lines.append(f"  ✗ {issue}")

        # 광고이미지
        if result.get("product_image_ok") or result["product_image_issues"]:
            lines.append("[광고이미지]")
            for ok in result.get("product_image_ok", []):
                lines.append(f"  ✓ {ok}")
            for issue in result["product_image_issues"]:
                lines.append(f"  ✗ {issue}")

        lines.append(f"\n→ {'검수 통과' if result['passed'] else '검수 실패'}")

        self.review_detail.insert("1.0", "\n".join(lines))
        self.review_detail.configure(state="disabled")

    def _manual_pass(self):
        """선택된 검수 실패 항목을 수동으로 통과 처리."""
        item = self.step2_tree.focus()
        if not item:
            return
        vals = list(self.step2_tree.item(item, "values"))
        folder = vals[4]
        result = self.review_results.get(folder)
        if not result or result["passed"]:
            return

        result["passed"] = True
        self.pass_btn.configure(state="disabled")

        # 영구 저장
        saved_pass = load_manual_pass()
        saved_pass.add(folder)
        save_manual_pass(saved_pass)

        # 트리뷰 업데이트: 검수 → 통과, 태그 → pass
        vals[8] = "통과(수동)"
        self.step2_tree.item(item, values=vals, tags=("pass",))

        # 상세 보기 갱신
        self._show_review_detail()

        # 통과 건수 다시 세기
        passed = sum(1 for r in self.review_results.values() if r["passed"])
        total_reviewed = len(self.review_results)
        self.step2_progress.set(self.step2_progress.get().rsplit("검수 통과:", 1)[0]
                                + f"검수 통과: {passed}/{total_reviewed}")

    # ──────────────────────────────────────
    # Step 3: 자동 배정 (진입 시 바로 실행)
    # ──────────────────────────────────────
    def _build_step3(self):
        f = self.content_frame
        ttk.Label(f, text="Step 3. 자동 배정 (5일 키워드 겹침 체크)",
                  style="Title.TLabel").pack(anchor="w", pady=(0, 10))

        info_frame = ttk.Frame(f)
        info_frame.pack(fill="x", pady=5)
        ttk.Button(info_frame, text="재배정",
                   command=self._auto_assign).pack(side="left")
        self.assign_info_var = tk.StringVar(value="")
        ttk.Label(info_frame, textvariable=self.assign_info_var).pack(
            side="left", padx=15)

        # 하단: 수동 변경 (먼저 pack)
        ctrl_frame = ttk.Frame(f)
        ctrl_frame.pack(fill="x", pady=(5, 0), side="bottom")
        ttk.Label(ctrl_frame, text="선택 원고 업체 변경:").pack(side="left")
        companies = ["(선택 안 함)"] + list(COMPANY_FOLDERS.keys())
        self.company_combo = ttk.Combobox(ctrl_frame, values=companies,
                                          width=12, state="readonly")
        self.company_combo.pack(side="left", padx=5)
        ttk.Button(ctrl_frame, text="변경",
                   command=self._change_company).pack(side="left")
        self.remaining_var = tk.StringVar()
        ttk.Label(ctrl_frame, textvariable=self.remaining_var).pack(side="right")

        # 트리뷰
        tree_frame = ttk.Frame(f)
        tree_frame.pack(fill="both", expand=True)

        cols = ("folder", "product", "keyword", "company", "conflict", "excluded")
        self.assign_tree = ttk.Treeview(tree_frame, columns=cols,
                                         show="headings", height=18)
        self.assign_tree.heading("folder", text="원고")
        self.assign_tree.heading("product", text="제품")
        self.assign_tree.heading("keyword", text="키워드")
        self.assign_tree.heading("company", text="배정 업체")
        self.assign_tree.heading("conflict", text="키워드 겹침")
        self.assign_tree.heading("excluded", text="제외")
        self.assign_tree.column("folder", width=370)
        self.assign_tree.column("product", width=90)
        self.assign_tree.column("keyword", width=120)
        self.assign_tree.column("company", width=100, anchor="center")
        self.assign_tree.column("conflict", width=100, anchor="center")
        self.assign_tree.column("excluded", width=80, anchor="center")

        self.assign_tree.tag_configure("ok", background="#E8F5E9")
        self.assign_tree.tag_configure("conflict", background="#FFF3E0")
        self.assign_tree.tag_configure("excluded", background="#FFCDD2")

        sb = ttk.Scrollbar(tree_frame, orient="vertical",
                           command=self.assign_tree.yview)
        self.assign_tree.configure(yscrollcommand=sb.set)
        self.assign_tree.pack(side="left", fill="both", expand=True)
        sb.pack(side="right", fill="y")

        # 진입 시 바로 자동 배정 실행
        self._auto_assign()

    def _auto_assign(self):
        """검수 통과 원고를 업체에 자동 배정 (5일 키워드 체크)."""
        self.assign_tree.delete(*self.assign_tree.get_children())
        self.assignments.clear()
        self.publish_row_map = {}  # folder → 자사 발행리스트 행 번호

        passed = [folder for folder, r in self.review_results.items()
                  if r["passed"]]

        if not passed:
            self.assign_info_var.set("검수 통과 원고 없음. Step 2를 먼저 실행하세요.")
            return

        if not self.publish_data:
            self.assign_info_var.set("발행리스트 없음. Step 1을 먼저 실행하세요.")
            return

        # 업체별 필요 건수 (발행리스트 행 번호 포함, deep copy로 재배정 가능)
        import copy
        needs = {}
        for company, items in self.publish_data.items():
            if company in COMPANY_FOLDERS:
                needs[company] = {
                    "remaining": len(items),
                    "products": [it["product"] for it in items],
                    "available_rows": copy.deepcopy(items),
                    "used_keywords": set(self.company_keywords_5d.get(company, set())),
                }

        # 자동 배정
        conflict_count = 0
        for folder in passed:
            code = get_product_code(folder)
            info = PRODUCT_CODE_MAP.get(code, ("", ""))
            product_name = info[0]

            ms = next((m for m in self.matched
                       if m.get("folder_name") == folder), None)
            kw = ms.get("keyword", "") if ms else extract_keyword_from_folder(folder)

            # 원고 제목 (매칭 정보에서)
            ms_title = ms.get("title", "") if ms else ""

            best = None
            best_score = -1
            best_conflict = False

            for company, need in needs.items():
                if need["remaining"] <= 0:
                    continue
                # 제외 키워드 체크: 해당 업체에 배정 불가
                if check_exclusion(self.exclusions, company, ms_title, folder):
                    continue
                has_conflict = any(
                    keywords_overlap(kw, ek)
                    for ek in need["used_keywords"]
                )
                score = 0
                if product_name in need["products"]:
                    score += 10
                score += need["remaining"]
                if has_conflict:
                    score -= 100
                if score > best_score:
                    best_score = score
                    best = company
                    best_conflict = has_conflict

            # 제외된 업체 목록 표시
            excluded_companies = []
            for company in needs:
                exc_kw = check_exclusion(self.exclusions, company, ms_title, folder)
                if exc_kw:
                    excluded_companies.append(f"{company}({exc_kw})")
            excluded_str = ", ".join(excluded_companies)

            if best:
                self.assignments[folder] = best
                needs[best]["remaining"] -= 1
                needs[best]["used_keywords"].add(kw)
                # 발행리스트 행 매핑: 제품 일치 우선, 없으면 첫 번째
                avail = needs[best]["available_rows"]
                picked = None
                for idx, it in enumerate(avail):
                    if it["product"] == product_name:
                        picked = idx
                        break
                if picked is None and avail:
                    picked = 0
                if picked is not None:
                    picked_item = avail.pop(picked)
                    self.publish_row_map[folder] = picked_item
                conflict_str = "겹침" if best_conflict else ""
                if best_conflict:
                    conflict_count += 1
                tag = "conflict" if best_conflict else "ok"
            else:
                self.assignments[folder] = ""
                conflict_str = ""
                best = "(건수 초과)"
                tag = "conflict"

            if excluded_str:
                tag = "excluded" if tag == "ok" else tag

            self.assign_tree.insert("", "end", iid=folder, values=(
                folder, product_name, kw, best, conflict_str, excluded_str
            ), tags=(tag,))

        n = len(passed)
        self.assign_info_var.set(
            f"검수 통과 {n}건 배정 완료"
            + (f" (키워드 겹침 {conflict_count}건)" if conflict_count else ""))
        self._update_remaining()
        self.status_var.set("자동 배정 완료")

    def _open_exclusion_dialog(self):
        """업체별 제외 키워드 관리 다이얼로그."""
        dlg = tk.Toplevel(self.root)
        dlg.title("업체별 제외 키워드 설정")
        dlg.geometry("500x400")
        dlg.transient(self.root)
        dlg.grab_set()

        # 업체 선택
        top = ttk.Frame(dlg, padding=10)
        top.pack(fill="x")
        ttk.Label(top, text="업체:").pack(side="left")
        company_var = tk.StringVar()
        company_cb = ttk.Combobox(top, textvariable=company_var,
                                   values=list(COMPANY_FOLDERS.keys()),
                                   width=15, state="readonly")
        company_cb.pack(side="left", padx=5)

        # 키워드 목록
        list_frame = ttk.Frame(dlg, padding=(10, 0))
        list_frame.pack(fill="both", expand=True)
        ttk.Label(list_frame, text="제외 키워드 목록:").pack(anchor="w")

        lb = tk.Listbox(list_frame, font=("맑은 고딕", 10), height=12)
        lb.pack(fill="both", expand=True, pady=(3, 0))

        # 추가/삭제
        input_frame = ttk.Frame(dlg, padding=10)
        input_frame.pack(fill="x")
        ttk.Label(input_frame, text="키워드:").pack(side="left")
        kw_entry = ttk.Entry(input_frame, width=20)
        kw_entry.pack(side="left", padx=5)
        ttk.Label(input_frame, text="(쉼표로 여러 개 입력 가능)",
                  font=("맑은 고딕", 8)).pack(side="left", padx=5)

        def refresh_list(*_):
            lb.delete(0, "end")
            company = company_var.get()
            if company:
                for kw in self.exclusions.get(company, []):
                    lb.insert("end", kw)

        company_cb.bind("<<ComboboxSelected>>", refresh_list)

        def add_keywords():
            company = company_var.get()
            text = kw_entry.get().strip()
            if not company or not text:
                return
            new_kws = [k.strip() for k in text.replace("、", ",").split(",") if k.strip()]
            if not new_kws:
                return
            existing = self.exclusions.setdefault(company, [])
            for kw in new_kws:
                if kw not in existing:
                    existing.append(kw)
            save_exclusions(self.exclusions)
            kw_entry.delete(0, "end")
            refresh_list()

        def remove_keyword():
            company = company_var.get()
            sel = lb.curselection()
            if not company or not sel:
                return
            kw = lb.get(sel[0])
            kw_list = self.exclusions.get(company, [])
            if kw in kw_list:
                kw_list.remove(kw)
                if not kw_list:
                    del self.exclusions[company]
                save_exclusions(self.exclusions)
            refresh_list()

        btn_frame = ttk.Frame(dlg, padding=(10, 0, 10, 10))
        btn_frame.pack(fill="x")
        ttk.Button(btn_frame, text="추가", command=add_keywords).pack(side="left")
        ttk.Button(btn_frame, text="선택 삭제", command=remove_keyword).pack(side="left", padx=5)
        ttk.Button(btn_frame, text="닫기", command=dlg.destroy).pack(side="right")

        kw_entry.bind("<Return>", lambda e: add_keywords())

        # 첫 업체 자동 선택
        if COMPANY_FOLDERS:
            company_cb.current(0)
            refresh_list()

    def _change_company(self):
        item = self.assign_tree.focus()
        company = self.company_combo.get()
        if not item or not company:
            return

        vals = list(self.assign_tree.item(item, "values"))

        # 기존 매핑 해제
        old_company = self.assignments.get(item, "")
        old_pub = self.publish_row_map.pop(item, None)

        if company == "(선택 안 함)":
            self.assignments[item] = ""
            vals[3] = "(선택 안 함)"
            vals[4] = ""
            vals[5] = ""
            self.assign_tree.item(item, values=vals, tags=("conflict",))
        else:
            self.assignments[item] = company
            vals[3] = company
            kw = vals[2]
            existing_kws = self.company_keywords_5d.get(company, set())
            has_conflict = any(keywords_overlap(kw, ek) for ek in existing_kws)
            vals[4] = "겹침" if has_conflict else ""
            # 제외 키워드 체크
            ms = next((m for m in self.matched if m.get("folder_name") == item), None)
            ms_title = ms.get("title", "") if ms else ""
            exc_kw = check_exclusion(self.exclusions, company, ms_title, item)
            vals[5] = f"{company}({exc_kw})" if exc_kw else ""
            tag = "conflict" if has_conflict else "ok"
            if exc_kw:
                tag = "excluded"
            self.assign_tree.item(item, values=vals, tags=(tag,))

            # 새 업체의 발행리스트 행 매핑
            if company in self.publish_data:
                # 이미 다른 폴더에 매핑된 행 제외
                used_rows = {info["row"] for info in self.publish_row_map.values()}
                for pub_item in self.publish_data[company]:
                    if pub_item["row"] not in used_rows:
                        self.publish_row_map[item] = pub_item
                        break

        self._update_remaining()

    def _update_remaining(self):
        assigned_counts = {}
        for folder, company in self.assignments.items():
            if company:
                assigned_counts[company] = assigned_counts.get(company, 0) + 1

        parts = []
        for company, items in sorted(self.publish_data.items()):
            need = len(items)
            done = assigned_counts.get(company, 0)
            parts.append(f"{company}: {done}/{need}")

        total_assigned = sum(assigned_counts.values())
        total_needed = sum(len(v) for v in self.publish_data.values())
        self.remaining_var.set(
            f"배정 현황 ({total_assigned}/{total_needed}): "
            + " | ".join(parts))

    # ──────────────────────────────────────
    # Step 4: 실행
    # ──────────────────────────────────────
    def _build_step4(self):
        f = self.content_frame
        ttk.Label(f, text="Step 4. 배정 실행",
                  style="Title.TLabel").pack(anchor="w", pady=(0, 10))

        btn_frame = ttk.Frame(f)
        btn_frame.pack(fill="x", pady=5)
        self.exec_btn = ttk.Button(btn_frame, text="배정 실행",
                                    command=self._execute_assignment)
        self.exec_btn.pack(side="left")
        self.sheet_btn = ttk.Button(btn_frame, text="시트 파라미터 업데이트",
                                     command=self._update_sheet_dates,
                                     state="disabled")
        self.sheet_btn.pack(side="left", padx=10)
        self.kw_sheet_btn = ttk.Button(btn_frame, text="키워드 시트 기입",
                                        command=self._update_keyword_sheet,
                                        state="disabled")
        self.kw_sheet_btn.pack(side="left", padx=5)

        self.rollback_btn = ttk.Button(btn_frame, text="⚠ 작업 모두 취소",
                                        command=self._rollback_assignment,
                                        state="disabled")
        self.rollback_btn.pack(side="right", padx=5)

        self.exec_progress_var = tk.StringVar(value="")
        ttk.Label(btn_frame, textvariable=self.exec_progress_var).pack(
            side="left", padx=10)

        self.exec_log = scrolledtext.ScrolledText(
            f, height=30, state="disabled", font=("맑은 고딕", 9))
        self.exec_log.pack(fill="both", expand=True)

        self.next_btn.configure(state="disabled")

        # 배정 요약
        assigned = {fo: co for fo, co in self.assignments.items() if co}
        self._append_exec_log(f"배정 대상: {len(assigned)}건")
        for fo, co in assigned.items():
            code = get_product_code(fo)
            info = PRODUCT_CODE_MAP.get(code, ("", ""))
            self._append_exec_log(f"  {fo} → {co} ({info[0]})")

    def _execute_assignment(self):
        assigned = {fo: co for fo, co in self.assignments.items() if co}
        if not assigned:
            messagebox.showinfo("알림", "배정된 원고가 없습니다.")
            return

        if not messagebox.askyesno("확인",
            f"{len(assigned)}건의 원고를 배정합니다.\n"
            f"(원고모음 복사 + 발행요청 ZIP + 미배정 삭제)\n\n계속?"):
            return

        self.exec_btn.configure(state="disabled")
        date_mmdd = date_to_mmdd(self.target_date)
        total = len(assigned)

        def worker():
            success = 0
            errors = []
            param_updates = []
            rollback_items = []  # 되돌리기용 기록

            for i, (folder, company) in enumerate(assigned.items()):
                self.root.after(0, lambda i=i:
                    self.exec_progress_var.set(f"처리 중... {i+1}/{total}"))

                code = get_product_code(folder)
                log = [f"\n[{i+1}/{total}] {folder} → {company}"]
                rb = {"folder": folder}

                dest, err = copy_to_archive(folder, code)
                if err:
                    log.append(f"  ✗ 원고 모음: {err}")
                    errors.append(err)
                else:
                    log.append(f"  ✓ 원고 모음 복사")
                    rb["archive_path"] = dest

                zp, err2 = copy_to_request_as_zip(folder, company, date_mmdd)
                if err2:
                    log.append(f"  ✗ 발행요청 ZIP: {err2}")
                    errors.append(err2)
                else:
                    log.append(f"  ✓ 발행요청 ZIP")
                    rb["zip_path"] = zp

                try:
                    delete_from_unassigned(folder)
                    log.append(f"  ✓ 미배정 삭제")
                    rb["deleted"] = True
                except Exception as e:
                    log.append(f"  ✗ 삭제 실패: {e}")
                    errors.append(str(e))

                ms = next((m for m in self.matched
                           if m.get("folder_name") == folder), None)

                pub_info = getattr(self, 'publish_row_map', {}).get(folder)
                if pub_info:
                    param_val = ms.get("parameter", "") if ms else ""
                    param_updates.append((pub_info["row"], param_val))
                    if param_val:
                        log.append(f"  ✓ 파라미터: {param_val}")
                    else:
                        log.append(f"  - 파라미터 빈 값 기입 (원고리스트 J열 비어있음)")
                    rb["pub_row"] = pub_info["row"]
                else:
                    log.append(f"  - 발행리스트 매핑 없음")

                rollback_items.append(rb)

                if not err and not err2:
                    success += 1

                self.root.after(0, lambda l=log: self._append_exec_log("\n".join(l)))

            self.param_updates = param_updates
            self.rollback_items = rollback_items

            summary = f"\n{'='*50}\n완료: {success}/{total}건"
            if errors:
                summary += f" | 오류 {len(errors)}건"

            self.root.after(0, lambda: self._append_exec_log(summary))
            self.root.after(0, lambda: self.exec_progress_var.set(
                f"완료: {success}/{total}건"))
            self.root.after(0, lambda: self.sheet_btn.configure(state="normal"))
            self.root.after(0, lambda: self.kw_sheet_btn.configure(state="normal"))
            self.root.after(0, lambda: self.rollback_btn.configure(state="normal"))
            self.root.after(0, lambda: self.status_var.set(
                f"배정 완료: {success}/{total}건"))
            # 배정 완료 후 자동으로 시트 파라미터 + 키워드 시트 기입
            self.root.after(500, lambda: self._auto_update_sheets())

        threading.Thread(target=worker, daemon=True).start()

    def _auto_update_sheets(self):
        """배정 실행 후 시트 파라미터 + 키워드 시트 자동 순차 기입."""
        param_updates = getattr(self, 'param_updates', [])
        self.sheet_btn.configure(state="disabled")
        self.kw_sheet_btn.configure(state="disabled")

        def worker():
            # 1) 시트 파라미터 업데이트
            if param_updates:
                try:
                    self.root.after(0, lambda:
                        self.exec_progress_var.set("시트 파라미터 업데이트 중..."))
                    update_publish_parameters(self.spreadsheet, param_updates)
                    self.root.after(0, lambda:
                        self._append_exec_log(
                            f"\n✓ 발행리스트 파라미터 업데이트: {len(param_updates)}건"))
                except Exception as e:
                    self.root.after(0, lambda:
                        self._append_exec_log(f"\n✗ 시트 파라미터 오류: {e}"))

            # 2) 키워드 시트 기입
            pub_map = getattr(self, 'publish_row_map', {})
            assigned = {fo: co for fo, co in self.assignments.items() if co}
            row_nums = [pub_map[fo]["row"] for fo in assigned if pub_map.get(fo)]

            if row_nums:
                try:
                    self.root.after(0, lambda:
                        self.exec_progress_var.set("키워드 시트 기입 중..."))
                    import time
                    time.sleep(3)  # 수식 반영 대기

                    ws = self.spreadsheet.worksheet("자사 발행리스트")
                    all_rows = ws.get_all_values()
                    year = str(self.target_date.year)

                    entries = []
                    for row_num in row_nums:
                        r = all_rows[row_num - 1]
                        product = r[1].strip() if len(r) > 1 else ""
                        if not product or product == "-":
                            continue
                        entries.append({
                            "product": product,
                            "keyword": r[4].strip() if len(r) > 4 else "",
                            "topic": r[5].strip() if len(r) > 5 else "",
                            "parameter": r[7].strip() if len(r) > 7 else "",
                            "author": r[10].strip() if len(r) > 10 else "",
                            "category": r[27].strip() if len(r) > 27 else "",
                            "title": r[11].strip() if len(r) > 11 else "",
                        })

                    if entries:
                        results = update_keyword_sheet(CRED_FILE, year, entries)
                        msg = "\n키워드 시트 기입:\n" + "\n".join(results)
                        self.root.after(0, lambda: self._append_exec_log(msg))
                except Exception as e:
                    self.root.after(0, lambda:
                        self._append_exec_log(f"\n✗ 키워드 시트 오류: {e}"))

            self.root.after(0, lambda: self.sheet_btn.configure(state="normal"))
            self.root.after(0, lambda: self.kw_sheet_btn.configure(state="normal"))
            self.root.after(0, lambda: self.exec_progress_var.set("모든 작업 완료"))
            self.root.after(0, lambda: self.status_var.set("모든 작업 완료"))

        threading.Thread(target=worker, daemon=True).start()

    def _update_sheet_dates(self):
        param_updates = getattr(self, 'param_updates', [])
        if not param_updates:
            messagebox.showinfo("알림", "업데이트할 항목이 없습니다.")
            return

        if not messagebox.askyesno("확인",
            f"발행리스트 파라미터 {len(param_updates)}건을 시트에 기입합니다."):
            return

        def worker():
            try:
                update_publish_parameters(self.spreadsheet, param_updates)
                self.root.after(0, lambda:
                    self._append_exec_log(
                        f"\n✓ 발행리스트 파라미터 업데이트: {len(param_updates)}건"))
                self.root.after(0, lambda:
                    self.status_var.set("시트 업데이트 완료"))
            except Exception as e:
                self.root.after(0, lambda:
                    self._append_exec_log(f"\n✗ 시트 오류: {e}"))
        threading.Thread(target=worker, daemon=True).start()

    def _update_keyword_sheet(self):
        """키워드 배정 시트에 배정 결과 기입 (시트에서 최신 값 다시 읽기)."""
        pub_map = getattr(self, 'publish_row_map', {})
        assigned = {fo: co for fo, co in self.assignments.items() if co}

        if not any(pub_map.get(fo) for fo in assigned):
            messagebox.showinfo("알림", "기입할 항목이 없습니다.")
            return

        if not messagebox.askyesno("확인",
            f"키워드 배정 시트에 {len(assigned)}건을 기입합니다.\n"
            f"(자사 발행리스트에서 최신 값을 다시 읽습니다)"):
            return

        year = str(self.target_date.year)
        self.kw_sheet_btn.configure(state="disabled")
        row_nums = [pub_map[fo]["row"] for fo in assigned if pub_map.get(fo)]

        def worker():
            try:
                # 자사 발행리스트에서 최신 값 다시 읽기
                ws = self.spreadsheet.worksheet("자사 발행리스트")
                all_rows = ws.get_all_values()

                entries = []
                for row_num in row_nums:
                    r = all_rows[row_num - 1]  # 0-indexed
                    product = r[1].strip() if len(r) > 1 else ""
                    if not product or product == "-":
                        continue
                    entries.append({
                        "product": product,
                        "keyword": r[4].strip() if len(r) > 4 else "",
                        "topic": r[5].strip() if len(r) > 5 else "",
                        "parameter": r[7].strip() if len(r) > 7 else "",
                        "author": r[10].strip() if len(r) > 10 else "",
                        "category": r[27].strip() if len(r) > 27 else "",
                        "title": r[11].strip() if len(r) > 11 else "",
                    })

                if not entries:
                    self.root.after(0, lambda:
                        self._append_exec_log("\n✗ 키워드 시트: 유효한 항목 없음 (제품명 확인 필요)"))
                    self.root.after(0, lambda:
                        self.kw_sheet_btn.configure(state="normal"))
                    return

                results = update_keyword_sheet(CRED_FILE, year, entries)
                msg = "\n키워드 시트 기입:\n" + "\n".join(results)
                self.root.after(0, lambda: self._append_exec_log(msg))
                self.root.after(0, lambda:
                    self.status_var.set("키워드 시트 기입 완료"))
            except Exception as e:
                self.root.after(0, lambda:
                    self._append_exec_log(f"\n✗ 키워드 시트 오류: {e}"))
                self.root.after(0, lambda:
                    self.kw_sheet_btn.configure(state="normal"))
        threading.Thread(target=worker, daemon=True).start()

    def _rollback_assignment(self):
        """배정 실행 결과를 모두 되돌리기."""
        items = getattr(self, 'rollback_items', [])
        if not items:
            messagebox.showinfo("알림", "되돌릴 항목이 없습니다.")
            return

        if not messagebox.askyesno("⚠ 작업 취소 확인",
            f"{len(items)}건의 배정을 되돌립니다.\n\n"
            f"• 원고모음 → 미배정으로 복원\n"
            f"• 발행요청 ZIP 삭제\n"
            f"• 자사 발행리스트 행 삭제\n\n"
            f"정말 취소하시겠습니까?"):
            return

        self.rollback_btn.configure(state="disabled")

        def worker():
            self.root.after(0, lambda:
                self._append_exec_log(f"\n{'='*50}\n작업 취소 시작"))

            for rb in items:
                folder = rb["folder"]
                log = [f"\n[취소] {folder}"]

                # 1) 원고모음 → 미배정으로 복원
                archive = rb.get("archive_path")
                if archive and os.path.exists(archive):
                    dest = os.path.join(UNASSIGNED_PATH, folder)
                    try:
                        if os.path.exists(dest):
                            shutil.rmtree(dest)
                        shutil.copytree(archive, dest)
                        shutil.rmtree(archive)
                        log.append("  ✓ 미배정으로 복원")
                    except Exception as e:
                        log.append(f"  ✗ 복원 실패: {e}")
                else:
                    log.append("  - 원고모음 경로 없음 (복원 스킵)")

                # 2) 발행요청 ZIP 삭제
                zp = rb.get("zip_path")
                if zp and os.path.exists(zp):
                    try:
                        os.remove(zp)
                        log.append("  ✓ 발행요청 ZIP 삭제")
                    except Exception as e:
                        log.append(f"  ✗ ZIP 삭제 실패: {e}")

                self.root.after(0, lambda l=log: self._append_exec_log("\n".join(l)))

            # 3) 자사 발행리스트 행 삭제 (큰 번호부터 삭제해야 행 번호 안 밀림)
            pub_rows = sorted(
                [rb["pub_row"] for rb in items if rb.get("pub_row")],
                reverse=True)
            if pub_rows:
                try:
                    ws = self.spreadsheet.worksheet("자사 발행리스트")
                    for row_num in pub_rows:
                        ws.delete_rows(row_num)
                    self.root.after(0, lambda:
                        self._append_exec_log(
                            f"\n  ✓ 자사 발행리스트 {len(pub_rows)}행 삭제"))
                except Exception as e:
                    self.root.after(0, lambda:
                        self._append_exec_log(f"\n  ✗ 시트 행 삭제 오류: {e}"))

            self.rollback_items = []
            self.param_updates = []
            self.written_rows = []
            self.root.after(0, lambda:
                self.cancel_write_btn.configure(state="disabled"))
            self.root.after(0, lambda:
                self._append_exec_log(f"\n{'='*50}\n✓ 작업 취소 완료"))
            self.root.after(0, lambda:
                self.status_var.set("작업 취소 완료"))

        threading.Thread(target=worker, daemon=True).start()

    def _append_exec_log(self, text):
        self.exec_log.configure(state="normal")
        self.exec_log.insert("end", text + "\n")
        self.exec_log.see("end")
        self.exec_log.configure(state="disabled")

    def _log(self, widget, msg):
        widget.configure(state="normal")
        widget.insert("end", msg + "\n")
        widget.see("end")
        widget.configure(state="disabled")


# ─────────────────────────────────────────────
# 7. 실행
# ─────────────────────────────────────────────
if __name__ == "__main__":
    root = tk.Tk()
    app = ManuscriptAssignerApp(root)
    root.mainloop()
