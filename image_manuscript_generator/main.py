"""
이미지 기반 원고 제작기 v1.0
- 이미지를 먼저 선택 → 이미지에 맞는 원고 생성
- 치환 모드: 기존 원고를 다른 제품/이미지에 맞게 치환
- tkinter GUI (4탭: 신규생성, 치환모드, 설정, 이력)
"""
import os
import sys
import re
import threading
import tkinter as tk
from tkinter import ttk, scrolledtext, filedialog, messagebox

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import lib_common as lc
from image_metadata import ImageMetadataStore, ThumbnailCache
from image_selector import ImageSelector, ImageSlot
import substitution as sub

VERSION = "1.0"


# ╔══════════════════════════════════════════════════════════════╗
# ║  save_as_docx — 기존 원고제작기(manuscript_generator) 완전판    ║
# ╚══════════════════════════════════════════════════════════════╝


def _build_styled_segments(original_text, colored_words):
    """텍스트를 마크다운 볼드/이탤릭 + 색상 키워드 기준으로 세그먼트 분할"""
    md_re = re.compile(r'\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*')

    chunks = []  # (text, is_bold, is_italic)
    last = 0
    for m in md_re.finditer(original_text):
        if m.start() > last:
            chunks.append((original_text[last:m.start()], False, False))
        if m.group(1):       # ***bold italic***
            chunks.append((m.group(1), True, True))
        elif m.group(2):     # **bold**
            chunks.append((m.group(2), True, False))
        elif m.group(3):     # *italic*
            chunks.append((m.group(3), False, True))
        last = m.end()
    if last < len(original_text):
        chunks.append((original_text[last:], False, False))

    if not colored_words:
        return [(t, {'bold': b, 'italic': it, 'color': None})
                for t, b, it in chunks]

    visible_text = ''.join(c[0] for c in chunks)
    char_colors = [None] * len(visible_text)
    for word, color_name in colored_words:
        for m in re.finditer(re.escape(word), visible_text):
            for j in range(m.start(), m.end()):
                char_colors[j] = color_name

    segments = []
    pos = 0
    for chunk_text, is_bold, is_italic in chunks:
        i = 0
        while i < len(chunk_text):
            current_color = char_colors[pos + i]
            j = i
            while j < len(chunk_text) and char_colors[pos + j] == current_color:
                j += 1
            segments.append((chunk_text[i:j], {'bold': is_bold, 'italic': is_italic, 'color': current_color}))
            i = j
        pos += len(chunk_text)

    return segments


def _clear_paragraph_runs(para):
    """기존 run 제거 (pPr 유지)"""
    from docx.oxml.ns import qn
    p_elem = para._element
    for r in list(p_elem.findall(qn('w:r'))):
        p_elem.remove(r)


def _apply_quote_border(paragraph, quote_num):
    """인용구 스타일 — 왼쪽 컬러 테두리"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    color_map = {
        1: '4472C4', 2: '70AD47', 3: 'ED7D31',
        4: 'FFC000', 5: '5B9BD5', 6: '7030A0',
    }
    border_color = color_map.get(quote_num, '4472C4')
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '18')
    left.set(qn('w:space'), '8')
    left.set(qn('w:color'), border_color)
    pBdr.append(left)
    pPr.append(pBdr)


def _apply_formatting_to_para(para, original_text, fmt):
    """ㄴ 서식 딕셔너리를 해당 문단에 실제 적용"""
    from docx.shared import Pt, RGBColor
    BLUE_C = RGBColor(0x00, 0x70, 0xC0)
    _clear_paragraph_runs(para)

    segments = _build_styled_segments(original_text, fmt.get('colored_words', []))

    is_quote = bool(fmt.get('quote'))

    for seg_text, seg_props in segments:
        run = para.add_run(seg_text)
        if fmt.get('font_size'):
            run.font.size = Pt(fmt['font_size'])
        # 인용구는 무조건 볼드 / 일반: 색상 지정된 단어만 볼드
        if is_quote:
            run.bold = True
        elif seg_props.get('bold'):
            run.bold = True
        elif fmt.get('bold') and seg_props.get('color'):
            run.bold = True
        if fmt.get('italic') or seg_props.get('italic'):
            run.italic = True
        # 인용구에는 글자색/형광펜 적용하지 않음 (볼드만 가능)
        if not is_quote:
            _color_map = lc.get_color_name_to_rgb()
            # 개별 단어 색상
            color_name = seg_props.get('color')
            if color_name and color_name in _color_map:
                run.font.color.rgb = _color_map[color_name]
            # 전체 글자색 (개별 색상이 없을 때만)
            elif fmt.get('full_text_color'):
                ftc = fmt['full_text_color']
                if ftc in _color_map:
                    run.font.color.rgb = _color_map[ftc]
            if fmt.get('highlight'):
                run.font.highlight_color = fmt['highlight']

    if fmt.get('link'):
        for run in para.runs:
            run.font.color.rgb = BLUE_C
            run.underline = True


def _add_text_runs(para, text):
    """마크다운 볼드/이탤릭 처리하여 run 추가"""
    md_re = re.compile(r'\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*')
    last = 0
    for m in md_re.finditer(text):
        if m.start() > last:
            para.add_run(text[last:m.start()])
        if m.group(1):
            run = para.add_run(m.group(1))
            run.bold = True
            run.italic = True
        elif m.group(2):
            run = para.add_run(m.group(2))
            run.bold = True
        elif m.group(3):
            run = para.add_run(m.group(3))
            run.italic = True
        last = m.end()
    if last < len(text):
        para.add_run(text[last:])


def _add_blogger_request_box(doc, lines):
    """★ 블로거 요청사항 → 빨간 테두리 + 노란 배경 테이블 박스"""
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)

    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'FFF8E1')
    shading.set(qn('w:val'), 'clear')
    cell._element.get_or_add_tcPr().append(shading)

    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '12')
        border.set(qn('w:color'), 'FF0000')
        border.set(qn('w:space'), '0')
        tcBorders.append(border)
    cell._element.get_or_add_tcPr().append(tcBorders)

    first_para = cell.paragraphs[0]
    for i, line in enumerate(lines):
        para = cell.add_paragraph() if i > 0 else first_para
        run = para.add_run(line)
        run.bold = True
        run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
        run.font.size = Pt(11)

    doc.add_paragraph('')


def save_as_docx(text, filepath):
    """생성된 원고를 Word 파일로 저장 (ㄴ서식 적용) — 완전판"""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '맑은 고딕'
    style.font.size = Pt(11)

    annotation_re = re.compile(r'^ㄴ\s*')
    image_num_re = re.compile(r'^\d{1,2}$')

    GREEN = RGBColor(0x00, 0x80, 0x00)
    BLUE = RGBColor(0x00, 0x70, 0xC0)

    lines = text.split('\n')
    recent = []  # (paragraph, original_text) 버퍼
    pending_fmts = []  # 아래 텍스트에 적용할 대기 서식 [(fmt, collected_paras)]
    blogger_req_lines = []
    in_blogger_req = False

    for line in lines:
        stripped = line.strip()

        # ── ★ 블로거 요청사항 수집 ──
        if '★' in stripped:
            in_blogger_req = True
            blogger_req_lines.append(stripped)
            continue

        if in_blogger_req:
            if not stripped:
                _add_blogger_request_box(doc, blogger_req_lines)
                blogger_req_lines = []
                in_blogger_req = False
                recent.append((doc.paragraphs[-1] if doc.paragraphs else None, ''))
                continue
            else:
                blogger_req_lines.append(stripped)
                continue

        # ── 빈 줄 ──
        if not stripped:
            p = doc.add_paragraph('')
            recent.append((p, ''))
            continue

        # ── 본문 중간에 ㄴ 서식이 섞인 경우 ("그래서 ㄴ '2mg' 파란색") ──
        mid_ann = re.match(r'^(.+?)\s+ㄴ\s+(.+)$', stripped)
        if mid_ann and lc.is_format_annotation('ㄴ ' + mid_ann.group(2)):
            content_part = mid_ann.group(1).strip()
            ann_part = 'ㄴ ' + mid_ann.group(2).strip()
            p = doc.add_paragraph()
            _add_text_runs(p, content_part)
            recent.append((p, content_part))
            mid_fmt = lc.parse_annotation(ann_part)
            if mid_fmt.get('colored_words'):
                if all(w in content_part for w, _ in mid_fmt['colored_words']):
                    _apply_formatting_to_para(p, content_part, mid_fmt)
                else:
                    pending_fmts.append((mid_fmt, []))
            elif not mid_fmt['is_image_desc']:
                _apply_formatting_to_para(p, content_part, mid_fmt)
            ap = doc.add_paragraph()
            run = ap.add_run(ann_part)
            run.bold = True
            run.font.color.rgb = GREEN
            run.font.size = Pt(24)
            run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
            recent.append((ap, ann_part))
            continue

        # ── ㄴ 서식 지시 줄 (서식 키워드가 있는 것만) ──
        if annotation_re.match(stripped) and lc.is_format_annotation(stripped):
            fmt = lc.parse_annotation(stripped)

            # 이미지 설명(ㄴ (설명))은 문서에서 제외
            if fmt['is_image_desc']:
                continue

            # ㄴ 줄 자체와 빈 줄을 제외한 콘텐츠 문단만 추출
            content_paras = [(p, t) for p, t in recent
                             if t.strip() and not (re.match(r'^ㄴ\s*', t.strip()) and lc.is_format_annotation(t.strip()))]
            target_count = fmt['multi_line']
            targets = content_paras[-target_count:] if content_paras else []

            applied = False
            # colored_words가 대상 문단에 있는지 확인
            if fmt.get('colored_words'):
                if targets:
                    all_target_text = ' '.join(t for _, t in targets)
                    missing = any(w not in all_target_text for w, _ in fmt['colored_words'])
                    if missing and len(content_paras) > target_count:
                        found = False
                        for ext in range(target_count + 1, min(target_count + 5, len(content_paras) + 1)):
                            targets = content_paras[-ext:]
                            all_target_text = ' '.join(t for _, t in targets)
                            if all(w in all_target_text for w, _ in fmt['colored_words']):
                                found = True
                                break
                        if found:
                            applied = True
                        else:
                            # 위에서 못 찾음 → 아래 텍스트에 적용 대기
                            pending_fmts.append((fmt, []))
                    elif not missing:
                        applied = True
                    else:
                        # targets가 1개인데 못 찾음 → 대기
                        pending_fmts.append((fmt, []))
                else:
                    # 위에 콘텐츠가 없음 → 아래에 적용 대기
                    pending_fmts.append((fmt, []))
            else:
                applied = True

            if applied:
                for para, para_text in targets:
                    _apply_formatting_to_para(para, para_text, fmt)

            # ㄴ 줄 자체 → 초록색 주석
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.bold = True
            run.font.color.rgb = GREEN
            run.font.size = Pt(24)
            run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
            recent.append((p, stripped))
            continue

        # ── 이미지 번호 (00→0, 01→1, 02→2...) ──
        if image_num_re.match(stripped):
            display_num = str(int(stripped))  # 앞자리 0 제거
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(display_num)
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = BLUE
            recent.append((p, display_num))
            continue

        # ── 마크다운 헤딩 ──
        if stripped.startswith('### '):
            p = doc.add_heading(stripped.lstrip('# ').strip(), level=3)
            recent.append((p, stripped))
            continue
        if stripped.startswith('## '):
            p = doc.add_heading(stripped.lstrip('# ').strip(), level=2)
            recent.append((p, stripped))
            continue
        if stripped.startswith('# '):
            p = doc.add_heading(stripped.lstrip('# ').strip(), level=1)
            recent.append((p, stripped))
            continue

        # ── 일반 텍스트 ──
        p = doc.add_paragraph()
        _add_text_runs(p, stripped)
        recent.append((p, stripped))

        # ── 대기 중인 서식 적용 (ㄴ이 텍스트 위에 있었던 경우) ──
        if pending_fmts:
            new_pending = []
            for pfmt, collected in pending_fmts:
                collected.append((p, stripped))
                if pfmt.get('colored_words'):
                    all_text = ' '.join(t for _, t in collected)
                    if all(w in all_text for w, _ in pfmt['colored_words']):
                        for cp, ct in collected:
                            _apply_formatting_to_para(cp, ct, pfmt)
                    elif len(collected) < 5:
                        new_pending.append((pfmt, collected))
                else:
                    _apply_formatting_to_para(p, stripped, pfmt)
            pending_fmts = new_pending

        if len(recent) > 15:
            recent = recent[-15:]

    # 남은 블로거 요청사항 처리
    if blogger_req_lines:
        _add_blogger_request_box(doc, blogger_req_lines)

    doc.save(filepath)


# ╔══════════════════════════════════════════════════════════════╗
# ║  build_prompt — 이미지 시퀀스 포함 프롬프트 조립               ║
# ╚══════════════════════════════════════════════════════════════╝
# 기존 manuscript_generator/main.py의 build_prompt를 수정
# 변경점: 섹션 7-3 삭제, 섹션 7-4/7-5 추가

def build_prompt_with_images(sheet_data, product_name, prompt_type, style_name,
                             tone, font_size, alignment, quote_num, keywords, sub_keywords,
                             selected_refs, extra_instructions, include_toc,
                             image_selector,  # ImageSelector 인스턴스
                             persona_text="", title_text="", product_link="",
                             hashtags="", char_count="2000~2200",
                             color_positive="파란색", color_negative="빨간색",
                             highlight_emphasis="노란 형광펜", color_product="청록색",
                             highlight_product="노란 형광펜", title_repeat=True,
                             emphasis_fontsize="14", samples_dir=""):
    """이미지 시퀀스를 포함한 원고 생성 프롬프트"""

    prompt_template = sheet_data["prompts"].get(prompt_type, "")
    style_desc = sheet_data["styles"].get(style_name, "")
    guidelines = sheet_data["guidelines"]
    product_guide = sheet_data["products"].get(product_name, "")

    toc_instruction = "- 목차를 포함합니다." if include_toc else "- 목차를 넣지 않습니다."
    title_repeat_instruction = ("이미지 00 다음에 SEO 타이틀을 3줄 반복하고, "
        "'ㄴ 세 줄 모두, 글자 크기 11, 아주 옅은 회색' 서식 지시를 넣습니다.") if title_repeat \
        else "SEO 타이틀을 3번 반복하지 않습니다."

    parts = []

    # 1) 프롬프트 템플릿
    if prompt_template:
        parts.append(prompt_template)

    # 1-1) 샘플 원고
    sample_fname, sample_text = lc.load_sample_for_type(samples_dir, prompt_type, product_name, sheet_data)
    if sample_text:
        parts.append("\n\n===== 참고 원고 예시 (톤/구조/서식 참고용) =====")
        parts.append("이 원고의 톤, 문장 길이, 줄 끊기, 서식 지시(ㄴ), 이미지 배치, 전체 흐름을 참고하세요.")
        parts.append(f"\n--- 예시 원고 ---\n{sample_text}\n--- 예시 끝 ---")

    # 2) 작가 스타일
    if style_desc:
        parts.append(f"\n\n===== 작가 스타일 =====\n{style_desc}")

    # 3) 톤
    if tone == "반말":
        parts.append("\n\n===== 문체 =====\n반말(~거든, ~잖아)로, 친구에게 말하듯 편안하게.")
    else:
        parts.append("\n\n===== 문체 =====\n존댓말(~입니다, ~했어요)로, 정중하면서도 친근하게.")

    # 4) 제품 정보
    if product_guide:
        parts.append(f"\n\n===== 제품 정보: {product_name} =====\n{product_guide}")

    # 4-1) 목차 규칙
    parts.append("\n\n===== 목차 규칙 =====")
    parts.append("목차에는 제품명, 허브키워드를 넣지 마세요. 독자의 궁금증 중심으로.")

    # 5) 공통지침
    if guidelines:
        parts.append("\n\n===== 공통 작성 지침 =====")
        for i, g in enumerate(guidelines, 1):
            parts.append(f"{i}. {g}")

    # 6) 메인 키워드
    if keywords:
        parts.append(f"\n\n===== 메인 키워드 =====\n{keywords}")

    # 7) 연관 키워드
    if sub_keywords:
        parts.append(f"\n\n===== 연관 키워드 =====\n{sub_keywords}")

    # 7-1) 해시태그
    if hashtags:
        tags = " ".join(f"#{t.strip().lstrip('#').replace(' ', '')}" for t in hashtags.split(",") if t.strip())
        parts.append(f"\n\n===== 해시태그 =====\n{tags}")
        parts.append("★블로거 요청사항★ 안에만 넣으세요. 본문에는 넣지 마세요.")

    # 7-2) 글자수
    parts.append(f"\n\n===== 글자수 =====\n총 {char_count}자 범위로 작성하세요.")

    # 7-2b) 줄 끊기 규칙
    parts.append("\n\n===== 줄 끊기 규칙 =====")
    parts.append("네이버 블로그 스타일: 한 줄 공백포함 20자 이내. 줄바꿈으로 구분.")

    # 7-2c) ㄴ서식 배치 규칙
    parts.append("\n\n===== ㄴ 서식 지시 배치 규칙 =====")
    parts.append("ㄴ 서식 지시는 서식 적용할 텍스트 바로 아래 줄에. 독립 줄에 단독 작성.")

    # 7-2d) 색상 활용 지침
    parts.append("\n\n===== 색상 활용 지침 =====")
    parts.append("색상을 적극 활용. 구절/문장 단위로 색상 입히기. 색상 지시 10~15회 이상.")

    # ★ 7-4) 이미지 시퀀스 (신규)
    parts.append(image_selector.build_image_sequence_prompt())

    # ★ 7-5) 이미지-텍스트 연결 지침 (신규)
    parts.append(image_selector.build_image_text_guidelines())

    # 8) 참고자료
    papers = sheet_data.get("papers", {}).get(product_name, [])
    if selected_refs or papers:
        parts.append("\n\n===== 참고자료 =====")
        parts.append("참고자료의 논문/연구를 자연스럽게 인용하세요. 없는 논문은 지어내지 마세요.")
        for fname, content in selected_refs.items():
            if len(content) > 8000:
                content = content[:8000] + "\n..."
            parts.append(f"\n--- {fname} ---\n{content}")
        if papers:
            parts.append("\n--- 참고 논문 ---")
            for i, paper in enumerate(papers, 1):
                parts.append(f"\n[논문 {i}]\n{paper}")

    # 9) 서식 규칙
    fmt_template = sheet_data.get("format_instructions") or ""
    if fmt_template:
        link_text = product_link if product_link else "(제품 링크)"
        hl_emphasis = highlight_emphasis if highlight_emphasis != "없음" else "글꼴 두껍게"
        hl_product = highlight_product if highlight_product != "없음" else "글꼴 두껍게"
        try:
            parts.append(fmt_template.format(
                font_size=font_size, align_text=alignment, quote_num=quote_num,
                toc_instruction=toc_instruction, product_link=link_text,
                color_positive=color_positive, color_negative=color_negative,
                highlight_emphasis=hl_emphasis, color_product=color_product,
                highlight_product=hl_product, title_repeat=title_repeat_instruction,
                emphasis_fontsize=emphasis_fontsize,
            ))
        except KeyError as e:
            parts.append(f"\n\n[서식규칙 오류: {e}]")
    parts.append("\n[중요] ㄴ 서식에서 글자 크기는 11, 13, 15, 16, 19, 24, 28 중 하나만.")

    # 10) 페르소나
    if persona_text:
        parts.append(f"\n\n===== 블로거 페르소나 =====\n{persona_text}")
        if "제3자" in prompt_type:
            parts.append("위 페르소나는 '관찰 대상'. 제3자 시점에서 작성하세요.")
        else:
            parts.append("위 페르소나의 시점에서 작성하세요.")

    # 11) 제목
    if title_text:
        parts.append(f"\n\n===== 원고 제목 =====\n{title_text}")

    # 12) 추가 지시사항
    if extra_instructions:
        parts.append(f"\n\n===== 추가 지시사항 =====\n{extra_instructions}")

    # 13) 최종 체크리스트
    if guidelines:
        parts.append("\n\n===== 최종 체크리스트 =====")
        for i, g in enumerate(guidelines, 1):
            parts.append(f"□ {i}. {g}")

    # 14) 자연스러움 규칙
    parts.append("\n\n===== 자연스러움 규칙 =====")
    parts.append("이 글은 '이야기'입니다. 성분 나열 금지. 에피소드 5~8줄 이상. 감정 흐름 자연스럽게.")

    return "\n".join(parts), sample_fname


# ╔══════════════════════════════════════════════════════════════╗
# ║  GUI 앱                                                      ║
# ╚══════════════════════════════════════════════════════════════╝

class ImageManuscriptApp:

    THEME = lc.THEME

    def __init__(self, root):
        self.root = root
        self.root.title(f"이미지 기반 원고 제작기 v{VERSION}")
        self.root.geometry("1500x1000")
        self.root.minsize(1200, 800)
        self.root.configure(bg=self.THEME["bg"])

        # 상태
        self.paths = lc.make_paths()
        self.sheet_data = {"prompts": {}, "styles": {}, "guidelines": [], "products": {}, "product_codes": {}, "product_links": {}, "format_instructions": "", "papers": {}}
        self.spreadsheet = None
        self.drive_service = None
        self.is_generating = False
        self.batch_count = 0
        self.batch_current = 0
        self.batch_keywords = []

        # 이미지 관련
        self.metadata_store = ImageMetadataStore()
        self.thumb_cache = ThumbnailCache(self.paths["image_cache"])
        self.image_selector = ImageSelector(self.metadata_store)
        self._tk_images = {}  # 참조 유지용

        lc.setup_styles()
        self._build_ui()
        self._init_load()

        self.root.protocol("WM_DELETE_WINDOW", self._on_close)

    def _on_close(self):
        if self.is_generating:
            if not messagebox.askyesno("생성 중", "원고 생성이 진행 중입니다.\n종료하시겠습니까?"):
                return
        self.root.destroy()

    # ── UI 빌드 ──

    def _build_ui(self):
        self.notebook = ttk.Notebook(self.root)
        self.notebook.pack(fill='both', expand=True, padx=5, pady=5)

        self._build_tab_generate()
        self._build_tab_substitute()
        self._build_tab_settings()
        self._build_tab_history()

    def _build_tab_generate(self):
        """탭 1: 신규 생성"""
        tab = ttk.Frame(self.notebook)
        self.notebook.add(tab, text="  신규 생성  ")

        # 상단: 설정 + 이미지 프리뷰
        top = ttk.PanedWindow(tab, orient='horizontal')
        top.pack(fill='both', expand=True, padx=5, pady=5)

        # 좌측: 설정 패널
        left = ttk.LabelFrame(top, text="설정", padding=10)
        top.add(left, weight=1)

        row = 0
        # 제품
        ttk.Label(left, text="제품:").grid(row=row, column=0, sticky='w', pady=2)
        self.product_var = tk.StringVar()
        self.product_combo = ttk.Combobox(left, textvariable=self.product_var, state='readonly', width=20)
        self.product_combo.grid(row=row, column=1, sticky='ew', pady=2)
        row += 1

        # 원고 유형
        ttk.Label(left, text="유형:").grid(row=row, column=0, sticky='w', pady=2)
        self.type_var = tk.StringVar()
        self.type_combo = ttk.Combobox(left, textvariable=self.type_var, state='readonly', width=20)
        self.type_combo.grid(row=row, column=1, sticky='ew', pady=2)
        row += 1

        # 작가 스타일
        ttk.Label(left, text="스타일:").grid(row=row, column=0, sticky='w', pady=2)
        self.style_var = tk.StringVar()
        self.style_combo = ttk.Combobox(left, textvariable=self.style_var, state='readonly', width=20)
        self.style_combo.grid(row=row, column=1, sticky='ew', pady=2)
        row += 1

        # 톤
        ttk.Label(left, text="톤:").grid(row=row, column=0, sticky='w', pady=2)
        self.tone_var = tk.StringVar(value="존댓말")
        ttk.Combobox(left, textvariable=self.tone_var, values=["존댓말", "반말"], state='readonly', width=20).grid(row=row, column=1, sticky='ew', pady=2)
        row += 1

        # 키워드
        ttk.Label(left, text="키워드:").grid(row=row, column=0, sticky='w', pady=2)
        self.keyword_entry = ttk.Entry(left, width=25)
        self.keyword_entry.grid(row=row, column=1, sticky='ew', pady=2)
        row += 1

        # 연관 키워드
        ttk.Label(left, text="연관:").grid(row=row, column=0, sticky='w', pady=2)
        self.sub_kw_entry = ttk.Entry(left, width=25)
        self.sub_kw_entry.grid(row=row, column=1, sticky='ew', pady=2)
        row += 1

        # 해시태그
        ttk.Label(left, text="해시태그:").grid(row=row, column=0, sticky='w', pady=2)
        self.hash_entry = ttk.Entry(left, width=25)
        self.hash_entry.grid(row=row, column=1, sticky='ew', pady=2)
        row += 1

        # 글자수
        ttk.Label(left, text="글자수:").grid(row=row, column=0, sticky='w', pady=2)
        self.charcount_var = tk.StringVar(value="2000~2200")
        ttk.Combobox(left, textvariable=self.charcount_var,
                      values=["1500~1700", "2000~2200", "2500~2700", "3000~3300"],
                      width=20).grid(row=row, column=1, sticky='ew', pady=2)
        row += 1

        # 이미지 수
        ttk.Label(left, text="이미지수:").grid(row=row, column=0, sticky='w', pady=2)
        self.imgcount_var = tk.StringVar(value="15")
        ttk.Spinbox(left, from_=5, to=30, textvariable=self.imgcount_var, width=8).grid(row=row, column=1, sticky='w', pady=2)
        row += 1

        # 추가 지시
        ttk.Label(left, text="추가 지시:").grid(row=row, column=0, sticky='nw', pady=2)
        self.extra_text = tk.Text(left, width=25, height=3, font=('맑은 고딕', 9))
        self.extra_text.grid(row=row, column=1, sticky='ew', pady=2)
        row += 1

        left.columnconfigure(1, weight=1)

        # 우측: 이미지 프리뷰
        right = ttk.LabelFrame(top, text="이미지 프리뷰", padding=10)
        top.add(right, weight=2)

        # 이미지 그리드
        self.img_frame = ttk.Frame(right)
        self.img_frame.pack(fill='both', expand=True)

        # 이미지 버튼
        btn_frame = ttk.Frame(right)
        btn_frame.pack(fill='x', pady=5)
        ttk.Button(btn_frame, text="자동 선택", command=self._on_auto_select).pack(side='left', padx=2)
        ttk.Button(btn_frame, text="재선택", command=self._on_reselect).pack(side='left', padx=2)
        ttk.Button(btn_frame, text="수동 추가", command=self._on_manual_add).pack(side='left', padx=2)

        # 하단: 생성 버튼 + 결과
        bottom = ttk.Frame(tab)
        bottom.pack(fill='both', expand=True, padx=5)

        # 버튼 + 전광판
        ctrl = ttk.Frame(bottom)
        ctrl.pack(fill='x', pady=5)
        ttk.Button(ctrl, text="▶ 생성", style='Generate.TButton', command=self._on_generate).pack(side='left', padx=5)
        ttk.Button(ctrl, text="배치", command=self._on_batch).pack(side='left', padx=5)

        self.scoreboard = tk.Label(ctrl, text="글자수: 0  키워드: 0  이미지: 0",
                                   bg=self.THEME["inspect_bg"], fg=self.THEME["inspect_chars"],
                                   font=('맑은 고딕', 10, 'bold'), padx=10, pady=3)
        self.scoreboard.pack(side='left', padx=10)

        # 결과 텍스트
        self.result_text = scrolledtext.ScrolledText(bottom, wrap='word', height=15,
                                                      font=('맑은 고딕', 10),
                                                      bg=self.THEME["text_bg"], fg=self.THEME["text_fg"])
        self.result_text.pack(fill='both', expand=True, pady=5)

        # 저장 버튼
        save_frame = ttk.Frame(bottom)
        save_frame.pack(fill='x', pady=3)
        ttk.Button(save_frame, text="Word 저장", command=self._on_save_docx).pack(side='left', padx=2)
        ttk.Button(save_frame, text="텍스트 저장", command=self._on_save_txt).pack(side='left', padx=2)
        ttk.Button(save_frame, text="복사", command=self._on_copy).pack(side='left', padx=2)

    def _build_tab_substitute(self):
        """탭 2: 치환 모드"""
        tab = ttk.Frame(self.notebook)
        self.notebook.add(tab, text="  치환 모드  ")

        top = ttk.PanedWindow(tab, orient='horizontal')
        top.pack(fill='both', expand=True, padx=5, pady=5)

        # 좌측: 원본 입력
        left = ttk.LabelFrame(top, text="원본 원고", padding=10)
        top.add(left, weight=1)

        btn_f = ttk.Frame(left)
        btn_f.pack(fill='x', pady=3)
        ttk.Button(btn_f, text="DOCX 불러오기", command=self._on_load_docx).pack(side='left', padx=2)
        ttk.Button(btn_f, text="클립보드 붙여넣기", command=self._on_paste_text).pack(side='left', padx=2)

        self.sub_original = scrolledtext.ScrolledText(left, wrap='word', height=20,
                                                       font=('맑은 고딕', 9))
        self.sub_original.pack(fill='both', expand=True, pady=5)

        # 우측: 치환 설정
        right = ttk.LabelFrame(top, text="치환 설정", padding=10)
        top.add(right, weight=1)

        row = 0
        ttk.Label(right, text="대상 제품:").grid(row=row, column=0, sticky='w', pady=2)
        self.sub_product_var = tk.StringVar()
        self.sub_product_combo = ttk.Combobox(right, textvariable=self.sub_product_var, state='readonly', width=20)
        self.sub_product_combo.grid(row=row, column=1, sticky='ew', pady=2)
        row += 1

        ttk.Label(right, text="대상 키워드:").grid(row=row, column=0, sticky='w', pady=2)
        self.sub_keyword_entry = ttk.Entry(right, width=25)
        self.sub_keyword_entry.grid(row=row, column=1, sticky='ew', pady=2)
        row += 1

        ttk.Label(right, text="해시태그:").grid(row=row, column=0, sticky='w', pady=2)
        self.sub_hash_entry = ttk.Entry(right, width=25)
        self.sub_hash_entry.grid(row=row, column=1, sticky='ew', pady=2)
        row += 1

        right.columnconfigure(1, weight=1)

        # 이미지 매칭 (간략 리스트)
        ttk.Label(right, text="이미지 매칭:").grid(row=row, column=0, sticky='nw', pady=5)
        self.sub_match_text = tk.Text(right, width=30, height=10, font=('맑은 고딕', 8), state='disabled')
        self.sub_match_text.grid(row=row, column=1, sticky='nsew', pady=5)
        row += 1

        ttk.Button(right, text="자동 매칭", command=self._on_sub_auto_match).grid(row=row, column=0, columnspan=2, pady=5)
        row += 1
        right.rowconfigure(row - 2, weight=1)

        # 하단: 치환 생성
        bottom = ttk.Frame(tab)
        bottom.pack(fill='both', expand=True, padx=5)

        ttk.Button(bottom, text="▶ 치환 생성", style='Generate.TButton',
                   command=self._on_substitute).pack(pady=5)

        self.sub_result = scrolledtext.ScrolledText(bottom, wrap='word', height=15,
                                                     font=('맑은 고딕', 10),
                                                     bg=self.THEME["text_bg"])
        self.sub_result.pack(fill='both', expand=True, pady=5)

        save_f = ttk.Frame(bottom)
        save_f.pack(fill='x', pady=3)
        ttk.Button(save_f, text="Word 저장", command=self._on_sub_save_docx).pack(side='left', padx=2)
        ttk.Button(save_f, text="복사", command=self._on_sub_copy).pack(side='left', padx=2)

    def _build_tab_settings(self):
        """탭 3: 설정"""
        tab = ttk.Frame(self.notebook)
        self.notebook.add(tab, text="  설정  ")

        f = ttk.LabelFrame(tab, text="연결 설정", padding=15)
        f.pack(fill='x', padx=10, pady=10)

        ttk.Label(f, text="API Key:").grid(row=0, column=0, sticky='w', pady=3)
        self.api_key_entry = ttk.Entry(f, width=50, show='*')
        self.api_key_entry.grid(row=0, column=1, sticky='ew', pady=3)

        ttk.Label(f, text="Sheet ID:").grid(row=1, column=0, sticky='w', pady=3)
        self.sheet_id_entry = ttk.Entry(f, width=50)
        self.sheet_id_entry.grid(row=1, column=1, sticky='ew', pady=3)

        ttk.Label(f, text="이미지 Drive 폴더 ID:").grid(row=2, column=0, sticky='w', pady=3)
        self.img_folder_entry = ttk.Entry(f, width=50)
        self.img_folder_entry.grid(row=2, column=1, sticky='ew', pady=3)

        f.columnconfigure(1, weight=1)

        btn_f = ttk.Frame(tab)
        btn_f.pack(fill='x', padx=10, pady=5)
        ttk.Button(btn_f, text="저장 + 연결", command=self._on_save_settings).pack(side='left', padx=5)
        ttk.Button(btn_f, text="이미지 동기화", command=self._on_sync_images).pack(side='left', padx=5)

        self.status_label = ttk.Label(tab, text="", font=('맑은 고딕', 9))
        self.status_label.pack(padx=10, pady=5)

    def _build_tab_history(self):
        """탭 4: 이력"""
        tab = ttk.Frame(self.notebook)
        self.notebook.add(tab, text="  이력  ")

        self.history_text = scrolledtext.ScrolledText(tab, wrap='word', height=30,
                                                       font=('맑은 고딕', 9), state='disabled')
        self.history_text.pack(fill='both', expand=True, padx=10, pady=10)
        ttk.Button(tab, text="새로고침", command=self._refresh_history).pack(pady=5)

    # ── 초기화 ──

    def _init_load(self):
        api_key = lc.load_api_key(self.paths["api_key_file"])
        sheet_id = lc.load_sheet_id(self.paths["sheet_config_file"])
        if api_key:
            self.api_key_entry.insert(0, api_key)
        if sheet_id:
            self.sheet_id_entry.insert(0, sheet_id)

        if sheet_id:
            self._connect_and_load(sheet_id)

    def _connect_and_load(self, sheet_id):
        """시트 연결 + 데이터 로딩"""
        def _run():
            cred = self.paths["cred_file"]
            sp, err = lc.connect_sheet(sheet_id, cred)
            if err:
                self.root.after(0, lambda: self._set_status(f"시트 연결 실패: {err}"))
                return
            self.spreadsheet = sp
            self.sheet_data = lc.load_all_from_sheet(sp)

            # 이미지 메타데이터 로딩
            count = self.metadata_store.load_from_sheet(sp)

            # Drive 연결
            drv, err2 = lc.connect_drive(cred)
            if not err2:
                self.drive_service = drv

            self.root.after(0, lambda: self._on_data_loaded(count))

        threading.Thread(target=_run, daemon=True).start()
        self._set_status("데이터 로딩 중...")

    def _on_data_loaded(self, img_count):
        # 드롭다운 업데이트
        products = sorted(self.sheet_data["products"].keys())
        self.product_combo['values'] = products
        self.sub_product_combo['values'] = products
        if products:
            self.product_var.set(products[0])
            self.sub_product_var.set(products[0])

        types = sorted(self.sheet_data["prompts"].keys())
        self.type_combo['values'] = types
        if types:
            self.type_var.set(types[0])

        styles = sorted(self.sheet_data["styles"].keys())
        self.style_combo['values'] = styles
        if styles:
            self.style_var.set(styles[0])

        self._set_status(f"로딩 완료 — 제품 {len(products)}개, 유형 {len(types)}개, 이미지 {img_count}개")

    def _set_status(self, text):
        self.status_label.config(text=text)

    # ── 이미지 선택 ──

    def _on_auto_select(self):
        product = self.product_var.get()
        keyword = self.keyword_entry.get()
        prompt_type = self.type_var.get()
        img_count = int(self.imgcount_var.get())
        api_key = self.api_key_entry.get()

        if not product:
            messagebox.showwarning("알림", "제품을 선택하세요.")
            return

        def _run():
            slots = self.image_selector.auto_select(
                product, keyword, prompt_type, img_count,
                api_key=api_key if api_key else None,
            )
            self.root.after(0, self._refresh_image_grid)

        threading.Thread(target=_run, daemon=True).start()
        self._set_status("이미지 자동 선택 중...")

    def _on_reselect(self):
        product = self.product_var.get()
        keyword = self.keyword_entry.get()
        prompt_type = self.type_var.get()
        api_key = self.api_key_entry.get()
        self.image_selector.re_select_unlocked(
            product, keyword, prompt_type, api_key=api_key if api_key else None
        )
        self._refresh_image_grid()

    def _on_manual_add(self):
        """수동 이미지 추가 다이얼로그"""
        dialog = tk.Toplevel(self.root)
        dialog.title("이미지 브라우징")
        dialog.geometry("600x500")

        # 검색
        search_f = ttk.Frame(dialog)
        search_f.pack(fill='x', padx=10, pady=5)
        ttk.Label(search_f, text="검색:").pack(side='left')
        search_var = tk.StringVar()
        search_entry = ttk.Entry(search_f, textvariable=search_var, width=30)
        search_entry.pack(side='left', padx=5)

        # 필터
        product_filter = tk.StringVar(value="전체")
        ttk.Label(search_f, text="제품:").pack(side='left', padx=(10, 0))
        pf = ttk.Combobox(search_f, textvariable=product_filter,
                           values=["전체"] + self.metadata_store.get_products(),
                           state='readonly', width=12)
        pf.pack(side='left', padx=2)

        # 결과 리스트
        result_list = tk.Listbox(dialog, font=('맑은 고딕', 9), height=20)
        result_list.pack(fill='both', expand=True, padx=10, pady=5)

        all_images = self.metadata_store.all

        def _search(*args):
            result_list.delete(0, 'end')
            query = search_var.get().strip()
            pf_val = product_filter.get()
            if query:
                results = self.metadata_store.search(query, product=pf_val if pf_val != "전체" else None)
            else:
                results = self.metadata_store.filter(product=pf_val if pf_val != "전체" else None)
            for entry in results[:50]:
                result_list.insert('end', f"{entry['filename']} | {entry['scene']} | {entry['mood']}")

        search_var.trace_add('write', _search)
        product_filter.trace_add('write', _search)
        _search()

        def _select():
            sel = result_list.curselection()
            if not sel:
                return
            idx = sel[0]
            query = search_var.get().strip()
            pf_val = product_filter.get()
            if query:
                results = self.metadata_store.search(query, product=pf_val if pf_val != "전체" else None)
            else:
                results = self.metadata_store.filter(product=pf_val if pf_val != "전체" else None)
            if idx < len(results):
                self.image_selector.add_slot(entry=results[idx])
                self.imgcount_var.set(str(len(self.image_selector.slots)))
                self._refresh_image_grid()
            dialog.destroy()

        ttk.Button(dialog, text="선택", command=_select).pack(pady=5)

    def _refresh_image_grid(self):
        """이미지 프리뷰 그리드 업데이트"""
        for widget in self.img_frame.winfo_children():
            widget.destroy()

        cols = 5
        for i, slot in enumerate(self.image_selector.slots):
            row = i // cols
            col = i % cols

            frame = ttk.Frame(self.img_frame, relief='solid', borderwidth=1)
            frame.grid(row=row, column=col, padx=3, pady=3)

            # 번호 + 역할 배지
            badge = str(slot.index)
            if slot.role == "hooking":
                badge = f"H"
            elif slot.role == "product_cut":
                badge = f"P"

            lbl = ttk.Label(frame, text=badge, font=('맑은 고딕', 8, 'bold'))
            lbl.pack()

            # 썸네일 (있으면)
            if slot.entry:
                fid = slot.entry['drive_file_id']
                tk_img = self.thumb_cache.get_tk_image(fid, size=(70, 70))
                if tk_img:
                    self._tk_images[fid] = tk_img
                    img_lbl = ttk.Label(frame, image=tk_img)
                    img_lbl.pack()
                else:
                    ttk.Label(frame, text="🖼", font=('맑은 고딕', 20)).pack()

                # 장면 설명 (툴팁 대용)
                scene = slot.entry.get('scene', slot.entry.get('filename', ''))[:12]
                ttk.Label(frame, text=scene, font=('맑은 고딕', 7)).pack()
            else:
                ttk.Label(frame, text="(빈칸)", font=('맑은 고딕', 8)).pack(pady=10)

        self._set_status(f"이미지 {len(self.image_selector.slots)}장 선택됨")

    # ── 원고 생성 ──

    def _on_generate(self, is_batch=False):
        if self.is_generating:
            return

        api_key = self.api_key_entry.get()
        if not api_key:
            messagebox.showwarning("알림", "API Key를 설정하세요.")
            return

        # 배치 모드에서는 이미지 자동 선택 → 완료 후 생성 진행
        if is_batch:
            product = self.product_var.get()
            keyword = self.keyword_entry.get()
            prompt_type = self.type_var.get()
            img_count = int(self.imgcount_var.get())

            self.is_generating = True
            batch_info = f" ({self.batch_current}/{self.batch_count})"
            self._set_status(f"이미지 자동 선택 중...{batch_info}")

            def _run():
                self.image_selector.auto_select(
                    product, keyword, prompt_type, img_count,
                    api_key=api_key,
                )
                self.root.after(0, lambda: self._batch_after_image_select(api_key))

            threading.Thread(target=_run, daemon=True).start()
            return

        # 단건 모드: 이미지가 선택되어 있어야 함
        if not self.image_selector.slots:
            messagebox.showwarning("알림", "이미지를 먼저 선택하세요.")
            return

        self.is_generating = True
        self._set_status("원고 생성 중...")

        # 페르소나 → 제목 → 원고 순서
        self._generate_persona(api_key)

    def _batch_after_image_select(self, api_key):
        """배치 모드: 이미지 선택 완료 후 생성 진행"""
        self._refresh_image_grid()
        if not self.image_selector.slots:
            self._set_status("이미지 선택 실패, 다음 키워드로 넘어갑니다.")
            self.is_generating = False
            if self.batch_current < self.batch_count:
                self.root.after(500, self._batch_next)
            return
        batch_info = f" ({self.batch_current}/{self.batch_count})"
        self._set_status(f"페르소나 생성 중...{batch_info}")
        self._generate_persona(api_key, is_batch=True)

    def _generate_persona(self, api_key, is_batch=False):
        """Step 1: 페르소나 생성"""
        prompt = lc.build_persona_prompt(
            self.sheet_data, self.product_var.get(), self.type_var.get(),
            self.style_var.get(), self.tone_var.get(), self.keyword_entry.get()
        )

        def on_done(text):
            self.root.after(0, lambda: self._on_persona_done(text, api_key, is_batch))

        def on_err(err):
            self.root.after(0, lambda: self._on_gen_error(err))

        lc.call_claude_api(api_key, prompt, on_done, on_err, max_tokens=2048)

    def _on_persona_done(self, text, api_key, is_batch=False):
        """페르소나 결과 → 첫 번째 선택 → 제목 생성"""
        # 간단히 페르소나 A를 자동 선택
        persona = text
        m = re.search(r'\[페르소나 A\](.*?)(?=\[페르소나 [BC]\]|\Z)', text, re.DOTALL)
        if m:
            persona = m.group(1).strip()
        self._selected_persona = persona
        self._generate_title(api_key, persona, is_batch)

    def _generate_title(self, api_key, persona, is_batch=False):
        """Step 2: 제목 생성"""
        prompt = lc.build_title_prompt(
            self.product_var.get(), self.type_var.get(),
            self.keyword_entry.get(), self.sub_kw_entry.get(), persona
        )

        def on_done(text):
            self.root.after(0, lambda: self._on_title_done(text, api_key, persona, is_batch))

        def on_err(err):
            self.root.after(0, lambda: self._on_gen_error(err))

        lc.call_claude_api(api_key, prompt, on_done, on_err, max_tokens=1024)

    def _on_title_done(self, text, api_key, persona, is_batch=False):
        """제목 결과 → 제목 A 자동 선택 → 원고 생성"""
        title = text
        m = re.search(r'\[제목 A\]\s*\n(.+)', text)
        if m:
            title = m.group(1).strip()
        self._selected_title = title
        self._generate_manuscript(api_key, persona, title, is_batch)

    def _generate_manuscript(self, api_key, persona, title, is_batch=False):
        """Step 3: 원고 생성"""
        refs = lc.load_refs_for_product(self.paths["references"], self.product_var.get())

        prompt, sample_fname = build_prompt_with_images(
            self.sheet_data,
            self.product_var.get(),
            self.type_var.get(),
            self.style_var.get(),
            self.tone_var.get(),
            "13", "가운데정렬", "1",
            self.keyword_entry.get(),
            self.sub_kw_entry.get(),
            refs,
            self.extra_text.get("1.0", "end").strip(),
            True,
            self.image_selector,
            persona_text=persona,
            title_text=title,
            hashtags=self.hash_entry.get(),
            char_count=self.charcount_var.get(),
            samples_dir=self.paths["samples"],
        )

        def on_done(text):
            self.root.after(0, lambda: self._on_manuscript_done(text, is_batch))

        def on_err(err):
            self.root.after(0, lambda: self._on_gen_error(err, is_batch))

        lc.call_claude_api(api_key, prompt, on_done, on_err, max_tokens=8192)

    def _on_manuscript_done(self, text, is_batch=False):
        self.is_generating = False
        # 후처리: 이미지 번호 정규화 (00→0)
        text = re.sub(r'^0(\d)$', r'\1', text, flags=re.MULTILINE)

        self.result_text.delete("1.0", "end")
        self.result_text.insert("1.0", text)

        # 스코어보드
        char_count = len(text.replace('\n', '').replace(' ', ''))
        kw = self.keyword_entry.get()
        kw_count = text.count(kw) if kw else 0
        img_count = len(re.findall(r'^\d{1,2}$', text, re.MULTILINE))
        self.scoreboard.config(text=f"글자수: {char_count:,}  키워드: {kw_count}  이미지: {img_count}")

        self._set_status("원고 생성 완료!")

        # 자동 저장
        self._auto_save(text)

        # 배치 모드: 다음 키워드로 진행
        if is_batch and self.batch_current < self.batch_count:
            self.root.after(500, self._batch_next)
        elif is_batch:
            # 마지막 키워드 완료
            total = self.batch_count
            self.batch_count = 0
            self.batch_current = 0
            self.batch_keywords = []
            self._set_status(f"연속 생성 완료! (총 {total}개)")
            messagebox.showinfo("연속 생성 완료", f"모든 원고가 생성되었습니다! (총 {total}개)")

    def _on_gen_error(self, err, is_batch=False):
        self.is_generating = False
        self._set_status(f"오류: {err}")
        if is_batch:
            # 배치 모드에서는 오류 시 다음 키워드로 넘어감
            if self.batch_current < self.batch_count:
                self.root.after(500, self._batch_next)
            else:
                self.batch_count = 0
                self.batch_current = 0
                self.batch_keywords = []
        else:
            messagebox.showerror("생성 오류", str(err))

    def _auto_save(self, text):
        """자동 저장 (텍스트)"""
        keyword = self.keyword_entry.get().replace(' ', '') or "원고"
        product_code = lc._get_product_code(self.product_var.get(), self.sheet_data)
        date_str = __import__('datetime').datetime.now().strftime("%y%m%d")
        fname = f"{date_str}{keyword}_{self.type_var.get()}_{product_code}.txt"
        fpath = os.path.join(self.paths["output"], fname)
        with open(fpath, 'w', encoding='utf-8') as f:
            f.write(text)

    # ── 배치 모드 (연속 생성) ──

    def _on_batch(self):
        if self.is_generating:
            messagebox.showinfo("진행 중", "생성이 진행 중입니다.")
            return

        dialog = tk.Toplevel(self.root)
        dialog.title("연속 생성")
        dialog.geometry("600x450")
        dialog.transient(self.root)
        dialog.grab_set()

        ttk.Label(dialog, text="키워드별 연속 생성",
                  font=('맑은 고딕', 11, 'bold')).pack(pady=(15, 5))
        ttk.Label(dialog, text="한 줄에 하나씩 입력하세요. | 구분자로 연관 키워드, 해시태그를 함께 입력할 수 있습니다.",
                  font=('맑은 고딕', 9), foreground='#666').pack(pady=(0, 3))
        ttk.Label(dialog, text="형식: 메인키워드 | 연관키워드1,연관키워드2 | #해시1,#해시2   (연관/해시 생략 가능)",
                  font=('맑은 고딕', 9), foreground='#888').pack(pady=(0, 10))

        kw_frame = ttk.LabelFrame(dialog, text="키워드 목록", padding=8)
        kw_frame.pack(fill=tk.BOTH, expand=True, padx=15, pady=(0, 10))

        kw_text = tk.Text(kw_frame, height=10, font=('맑은 고딕', 10), wrap=tk.WORD)
        kw_text.pack(fill=tk.BOTH, expand=True)

        # 현재 키워드를 첫 줄에 미리 입력
        current_kw = self.keyword_entry.get().strip()
        if current_kw:
            parts = [current_kw]
            current_sub = self.sub_kw_entry.get().strip()
            current_hash = self.hash_entry.get().strip()
            if current_sub or current_hash:
                parts.append(current_sub)
            if current_hash:
                parts.append(current_hash)
            kw_text.insert('1.0', " | ".join(parts) + "\n")

        ttk.Label(dialog, text="각 키워드마다 이미지 자동 선택 → 페르소나/제목 자동 선택 → 원고 생성이 순서대로 실행됩니다.",
                  font=('맑은 고딕', 8), foreground='#666', wraplength=550).pack(padx=15, pady=(0, 10))

        def on_start():
            lines = [l.strip() for l in kw_text.get('1.0', tk.END).strip().split('\n') if l.strip()]
            if not lines:
                messagebox.showwarning("키워드", "키워드를 하나 이상 입력해주세요.", parent=dialog)
                return
            parsed = []
            for line in lines:
                segments = [s.strip() for s in line.split('|')]
                main_kw = segments[0] if len(segments) >= 1 else ""
                sub_kw = segments[1] if len(segments) >= 2 else ""
                hashtags = segments[2] if len(segments) >= 3 else ""
                if main_kw:
                    parsed.append((main_kw, sub_kw, hashtags))
            if not parsed:
                messagebox.showwarning("키워드", "메인 키워드를 하나 이상 입력해주세요.", parent=dialog)
                return
            dialog.destroy()
            self.batch_keywords = parsed
            self.batch_count = len(parsed)
            self.batch_current = 0
            self._batch_next()

        btn_frame = ttk.Frame(dialog)
        btn_frame.pack(fill=tk.X, padx=15, pady=(0, 15))
        ttk.Button(btn_frame, text="시작", style='Generate.TButton', command=on_start).pack(side=tk.LEFT, padx=(0, 10))
        ttk.Button(btn_frame, text="취소", command=dialog.destroy).pack(side=tk.LEFT)

    def _batch_next(self):
        if self.batch_current >= self.batch_count:
            total = self.batch_count
            self.batch_count = 0
            self.batch_current = 0
            self.batch_keywords = []
            self._set_status(f"연속 생성 완료! (총 {total}개)")
            messagebox.showinfo("연속 생성 완료", f"모든 원고가 생성되었습니다! (총 {total}개)")
            return
        # 현재 배치의 키워드를 입력란에 세팅
        main_kw, sub_kw, hashtags = self.batch_keywords[self.batch_current]
        self.keyword_entry.delete(0, tk.END)
        self.keyword_entry.insert(0, main_kw)
        self.sub_kw_entry.delete(0, tk.END)
        self.sub_kw_entry.insert(0, sub_kw)
        self.hash_entry.delete(0, tk.END)
        self.hash_entry.insert(0, hashtags)
        self.batch_current += 1
        self._set_status(f"연속 생성 {self.batch_current}/{self.batch_count} — 키워드: {main_kw}")
        # 이미지 자동 선택 → 생성까지 자동 진행
        self._on_generate(is_batch=True)

    # ── 저장 ──

    def _on_save_docx(self):
        text = self.result_text.get("1.0", "end").strip()
        if not text:
            return
        fpath = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word", "*.docx")],
            initialdir=self.paths["output"]
        )
        if fpath:
            save_as_docx(text, fpath)
            self._set_status(f"Word 저장: {os.path.basename(fpath)}")

    def _on_save_txt(self):
        text = self.result_text.get("1.0", "end").strip()
        if not text:
            return
        fpath = filedialog.asksaveasfilename(
            defaultextension=".txt",
            filetypes=[("Text", "*.txt")],
            initialdir=self.paths["output"]
        )
        if fpath:
            with open(fpath, 'w', encoding='utf-8') as f:
                f.write(text)
            self._set_status(f"텍스트 저장: {os.path.basename(fpath)}")

    def _on_copy(self):
        text = self.result_text.get("1.0", "end").strip()
        if text:
            self.root.clipboard_clear()
            self.root.clipboard_append(text)
            self._set_status("클립보드에 복사됨")

    # ── 치환 모드 ──

    def _on_load_docx(self):
        fpath = filedialog.askopenfilename(filetypes=[("Word", "*.docx"), ("Text", "*.txt")])
        if fpath:
            content = lc.read_file_content(fpath)
            self.sub_original.delete("1.0", "end")
            self.sub_original.insert("1.0", content)

    def _on_paste_text(self):
        try:
            text = self.root.clipboard_get()
            self.sub_original.delete("1.0", "end")
            self.sub_original.insert("1.0", text)
        except Exception:
            pass

    def _on_sub_auto_match(self):
        """치환 모드: 자동 이미지 매칭"""
        original = self.sub_original.get("1.0", "end").strip()
        if not original:
            messagebox.showwarning("알림", "원본 원고를 입력하세요.")
            return

        api_key = self.api_key_entry.get()
        target_product = self.sub_product_var.get()
        if not api_key or not target_product:
            messagebox.showwarning("알림", "API Key와 대상 제품을 확인하세요.")
            return

        def _run():
            segments, img_count = sub.parse_original_manuscript(original)
            contexts = sub.get_context_around_images(segments)
            scenes = sub.infer_scenes(api_key, contexts)
            slots = sub.match_images_for_substitution(
                self.metadata_store, scenes, target_product, img_count
            )
            self.root.after(0, lambda: self._on_sub_match_done(slots))

        threading.Thread(target=_run, daemon=True).start()
        self._set_status("이미지 매칭 중...")

    def _on_sub_match_done(self, slots):
        self._sub_slots = slots
        self.sub_match_text.config(state='normal')
        self.sub_match_text.delete("1.0", "end")
        for slot in slots:
            self.sub_match_text.insert("end", f"img {slot.index}: {slot.description}\n")
        self.sub_match_text.config(state='disabled')
        self._set_status(f"이미지 매칭 완료 ({len(slots)}장)")

    def _on_substitute(self):
        """치환 생성"""
        original = self.sub_original.get("1.0", "end").strip()
        if not original:
            return
        if not hasattr(self, '_sub_slots') or not self._sub_slots:
            messagebox.showwarning("알림", "먼저 자동 매칭을 실행하세요.")
            return

        api_key = self.api_key_entry.get()
        target_product = self.sub_product_var.get()
        product_guide = self.sheet_data["products"].get(target_product, "")
        product_link = self.sheet_data.get("product_links", {}).get(target_product, "")
        product_code = self.sheet_data.get("product_codes", {}).get(target_product, "")

        prompt = sub.build_substitution_prompt(
            original, target_product, product_guide,
            self._sub_slots,
            target_keyword=self.sub_keyword_entry.get(),
            hashtags=self.sub_hash_entry.get(),
            product_link=product_link,
            product_code=product_code,
        )

        def on_done(text):
            self.root.after(0, lambda: self._on_sub_done(text))

        def on_err(err):
            self.root.after(0, lambda: self._on_gen_error(err))

        self.is_generating = True
        self._set_status("치환 생성 중...")
        lc.call_claude_api(api_key, prompt, on_done, on_err, max_tokens=8192)

    def _on_sub_done(self, text):
        self.is_generating = False
        text = re.sub(r'^0(\d)$', r'\1', text, flags=re.MULTILINE)
        self.sub_result.delete("1.0", "end")
        self.sub_result.insert("1.0", text)
        self._set_status("치환 완료!")

    def _on_sub_save_docx(self):
        text = self.sub_result.get("1.0", "end").strip()
        if not text:
            return
        fpath = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word", "*.docx")])
        if fpath:
            save_as_docx(text, fpath)

    def _on_sub_copy(self):
        text = self.sub_result.get("1.0", "end").strip()
        if text:
            self.root.clipboard_clear()
            self.root.clipboard_append(text)

    # ── 설정 ──

    def _on_save_settings(self):
        api_key = self.api_key_entry.get().strip()
        sheet_id = self.sheet_id_entry.get().strip()
        if api_key:
            lc.save_api_key(self.paths["api_key_file"], api_key)
        if sheet_id:
            lc.save_sheet_id(self.paths["sheet_config_file"], sheet_id)
            self._connect_and_load(sheet_id)

    def _on_sync_images(self):
        """이미지 메타데이터 새로고침"""
        if self.spreadsheet:
            count = self.metadata_store.load_from_sheet(self.spreadsheet)
            self._set_status(f"이미지 메타데이터 새로고침: {count}개")

    # ── 이력 ──

    def _refresh_history(self):
        log = lc.load_generation_log(self.paths["log_file"])
        self.history_text.config(state='normal')
        self.history_text.delete("1.0", "end")
        for entry in reversed(log[-50:]):
            line = f"[{entry.get('timestamp', '')}] {entry.get('product', '')} / {entry.get('prompt_type', '')} / {entry.get('title', '')[:30]} ({entry.get('char_count', 0)}자)\n"
            self.history_text.insert("end", line)
        self.history_text.config(state='disabled')


def main():
    root = tk.Tk()
    app = ImageManuscriptApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()
