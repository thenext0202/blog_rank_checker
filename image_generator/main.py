"""
이미지 제너레이터 v2.0

완성된 원고를 기반으로 이미지를 자동 제안/선택/변형하는 독립 프로그램.

흐름:
  1. 원고 입력 (파일 불러오기 / 직접 입력)
  2. 제품 선택 (이미지 시트에서 로딩)
  3. 임베딩 기반 하이브리드 검색으로 이미지 자동 제안
     - 1차: 제품 태그 필터링
     - 2차: 원고 맥락 vs 이미지 메타데이터 코사인 유사도
  4. 각 이미지: 수락 / 거절(대체 5장 + 프롬프트 변형)
  5. 최종 출력: 폴더에 원고 텍스트 + 이미지 파일(1.png, 2.png...)

의존:
  lib_common.py, image_metadata.py, image_selector.py, image_transformer.py
"""

import os
import sys
import re
import json
import threading
import tkinter as tk
from tkinter import ttk, messagebox, filedialog, scrolledtext

from lib_common import (
    base_dir, connect_sheet, connect_drive, drive_download_bytes,
    load_image_metadata_from_sheet, get_embeddings,
    CRED_FILE,
)
from image_metadata import ImageMetadataStore, ThumbnailCache
from image_transformer import (
    extract_image_contexts, transform_image, generate_image,
    build_transform_prompt,
)

# ── 설정 파일 경로 ──
CONFIG_DIR = base_dir()
IMAGE_SHEET_CONFIG = os.path.join(CONFIG_DIR, ".image_sheet_id")
API_KEY_FILE = os.path.join(CONFIG_DIR, ".api_key")
GEMINI_KEY_FILE = os.path.join(CONFIG_DIR, ".gemini_key")
IMAGE_CACHE_DIR = os.path.join(CONFIG_DIR, "image_cache")
OUTPUT_DIR = os.path.join(CONFIG_DIR, "output")
os.makedirs(IMAGE_CACHE_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)

VERSION = "2.0"


# ════════════════════════════════════════════════════
#  유틸리티
# ════════════════════════════════════════════════════

def load_key(filepath):
    if os.path.exists(filepath):
        with open(filepath, 'r', encoding='utf-8') as f:
            return f.read().strip()
    return ""


def save_key(filepath, key):
    with open(filepath, 'w', encoding='utf-8') as f:
        f.write(key.strip())


def extract_sheet_id(raw):
    """URL이든 순수 ID든 시트 ID만 추출"""
    m = re.search(r'/spreadsheets/d/([a-zA-Z0-9_-]+)', raw)
    return m.group(1) if m else raw.strip()


def extract_image_numbers(text):
    """원고에서 이미지 번호 추출 (0번 제외, 1번부터)"""
    image_num_re = re.compile(r'^\d{1,2}$')
    numbers = []
    for line in text.split('\n'):
        stripped = line.strip()
        if image_num_re.match(stripped):
            num = int(stripped)
            if num >= 1:
                numbers.append(num)
    return sorted(set(numbers))


# ════════════════════════════════════════════════════
#  메인 GUI
# ════════════════════════════════════════════════════

class ImageGeneratorApp:

    THEME = {
        "bg": "#FAFAFA", "fg": "#1A1A2E", "accent": "#0F3460",
        "accent2": "#E94560", "card_bg": "#FFFFFF", "text_bg": "#FFFFFF",
        "text_fg": "#1A1A2E", "btn_bg": "#0F3460", "btn_fg": "#FFFFFF",
        "ok": "#2E7D32", "warn": "#C62828", "info": "#1565C0",
    }

    def __init__(self):
        self.root = tk.Tk()
        self.root.title(f"이미지 제너레이터 v{VERSION}")
        self.root.geometry("1100x800")
        self.root.configure(bg=self.THEME["bg"])

        # 상태
        self.image_spreadsheet = None
        self.drive_service = None
        self.metadata_store = ImageMetadataStore()
        self.thumb_cache = ThumbnailCache(IMAGE_CACHE_DIR)
        self.tk_images = {}  # 참조 유지용

        # 이미지 슬롯: {img_num: {entry, scene_rec, mood_rec, reason, status, bytes}}
        self.image_slots = {}
        self.manuscript_text = ""
        self.is_processing = False

        self._build_ui()
        self._auto_connect()

    # ────────────────────────────────────────
    #  UI 구성
    # ────────────────────────────────────────

    def _build_ui(self):
        style = ttk.Style()
        style.theme_use('clam')

        # ── 상단: 설정 바 ──
        settings_bar = ttk.Frame(self.root)
        settings_bar.pack(fill=tk.X, padx=10, pady=(10, 5))

        ttk.Label(settings_bar, text="Claude:").pack(side=tk.LEFT, padx=(0, 3))
        self.api_key_var = tk.StringVar(value=load_key(API_KEY_FILE))
        ttk.Entry(settings_bar, textvariable=self.api_key_var, width=25, show='*').pack(side=tk.LEFT, padx=(0, 10))

        ttk.Label(settings_bar, text="Gemini:").pack(side=tk.LEFT, padx=(0, 3))
        self.gemini_key_var = tk.StringVar(value=load_key(GEMINI_KEY_FILE))
        ttk.Entry(settings_bar, textvariable=self.gemini_key_var, width=25, show='*').pack(side=tk.LEFT, padx=(0, 10))

        ttk.Label(settings_bar, text="이미지시트:").pack(side=tk.LEFT, padx=(0, 3))
        self.sheet_id_var = tk.StringVar(value=load_key(IMAGE_SHEET_CONFIG))
        ttk.Entry(settings_bar, textvariable=self.sheet_id_var, width=30).pack(side=tk.LEFT, padx=(0, 5))
        ttk.Button(settings_bar, text="연결", command=self._on_connect_sheet).pack(side=tk.LEFT, padx=(0, 10))

        ttk.Button(settings_bar, text="키 저장", command=self._save_keys).pack(side=tk.LEFT, padx=(0, 5))
        self.embed_btn = ttk.Button(settings_bar, text="임베딩 생성", command=self._on_build_embeddings)
        self.embed_btn.pack(side=tk.LEFT, padx=(0, 10))

        self.sheet_status_var = tk.StringVar(value="미연결")
        ttk.Label(settings_bar, textvariable=self.sheet_status_var,
                  font=('맑은 고딕', 9)).pack(side=tk.RIGHT)

        # ── 좌우 분할 ──
        paned = ttk.PanedWindow(self.root, orient=tk.HORIZONTAL)
        paned.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)

        # ── 왼쪽: 원고 입력 ──
        left = ttk.Frame(paned)
        paned.add(left, weight=1)

        left_top = ttk.Frame(left)
        left_top.pack(fill=tk.X, pady=(0, 5))
        ttk.Label(left_top, text="원고 입력",
                  font=('맑은 고딕', 11, 'bold')).pack(side=tk.LEFT)
        ttk.Button(left_top, text="파일 불러오기",
                   command=self._on_load_file).pack(side=tk.RIGHT, padx=(5, 0))

        self.manuscript_input = scrolledtext.ScrolledText(
            left, wrap=tk.WORD, font=('맑은 고딕', 10), height=20,
            bg=self.THEME["text_bg"], fg=self.THEME["text_fg"]
        )
        self.manuscript_input.pack(fill=tk.BOTH, expand=True)

        # 제품 선택 + 버튼
        action_bar = ttk.Frame(left)
        action_bar.pack(fill=tk.X, pady=(5, 0))

        ttk.Label(action_bar, text="제품:").pack(side=tk.LEFT, padx=(0, 3))
        self.product_var = tk.StringVar()
        self.product_combo = ttk.Combobox(
            action_bar, textvariable=self.product_var,
            state='readonly', width=15, values=[])
        self.product_combo.pack(side=tk.LEFT, padx=(0, 10))

        self.analyze_btn = ttk.Button(
            action_bar, text="이미지 자동 제안", command=self._on_analyze)
        self.analyze_btn.pack(side=tk.LEFT, padx=(0, 10))

        self.save_btn = ttk.Button(
            action_bar, text="최종 저장", command=self._on_save_all)
        self.save_btn.pack(side=tk.LEFT, padx=(0, 10))

        self.status_var = tk.StringVar(value="원고를 입력하고 이미지 시트를 연결하세요")
        ttk.Label(action_bar, textvariable=self.status_var,
                  font=('맑은 고딕', 9)).pack(side=tk.RIGHT)

        # ── 오른쪽: 이미지 카드 리스트 ──
        right = ttk.Frame(paned)
        paned.add(right, weight=1)

        ttk.Label(right, text="이미지 목록",
                  font=('맑은 고딕', 11, 'bold')).pack(anchor='w', pady=(0, 5))

        self._canvas = tk.Canvas(right, bg=self.THEME["bg"], highlightthickness=0)
        self._scrollbar = ttk.Scrollbar(
            right, orient="vertical", command=self._canvas.yview)
        self._cards_frame = ttk.Frame(self._canvas)

        self._cards_frame.bind(
            "<Configure>",
            lambda e: self._canvas.configure(scrollregion=self._canvas.bbox("all")))
        self._canvas_window = self._canvas.create_window(
            (0, 0), window=self._cards_frame, anchor="nw")
        self._canvas.configure(yscrollcommand=self._scrollbar.set)

        self._canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        self._scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        self._canvas.bind(
            "<Configure>",
            lambda e: self._canvas.itemconfig(self._canvas_window, width=e.width))
        self._canvas.bind_all(
            "<MouseWheel>",
            lambda e: self._canvas.yview_scroll(int(-1 * (e.delta / 120)), "units"))

    # ────────────────────────────────────────
    #  설정 / 연결
    # ────────────────────────────────────────

    def _save_keys(self):
        save_key(API_KEY_FILE, self.api_key_var.get())
        save_key(GEMINI_KEY_FILE, self.gemini_key_var.get())
        save_key(IMAGE_SHEET_CONFIG, self.sheet_id_var.get())
        self.status_var.set("키 저장 완료")

    def _on_build_embeddings(self):
        """임베딩 생성 버튼 — 시트 메타데이터를 임베딩으로 변환 후 파일 저장"""
        gemini_key = self.gemini_key_var.get().strip()
        if not gemini_key:
            messagebox.showwarning("API", "Gemini API Key를 입력하세요.")
            return
        if not self.metadata_store.all:
            messagebox.showwarning("시트", "이미지 시트를 먼저 연결하세요.")
            return

        self.embed_btn.config(state='disabled')
        self.status_var.set("임베딩 생성 중... (최초 1회만 필요)")

        def _run():
            try:
                def _progress(current, total):
                    self.root.after(0, lambda: self.status_var.set(
                        f"임베딩 생성 중... {current}/{total}"))

                count = self.metadata_store.compute_and_save_embeddings(
                    gemini_key, on_progress=_progress)

                self.root.after(0, lambda: (
                    self.status_var.set(f"임베딩 생성 완료! ({count}개)"),
                    self.sheet_status_var.set(
                        self.sheet_status_var.get().replace(
                            "임베딩 없음 (생성 필요)", "임베딩 OK")),
                ))
            except Exception as e:
                self.root.after(0, lambda err=e: self.status_var.set(
                    f"임베딩 오류: {err}"))
            finally:
                self.root.after(0, lambda: self.embed_btn.config(state='normal'))

        threading.Thread(target=_run, daemon=True).start()

    def _auto_connect(self):
        raw = self.sheet_id_var.get().strip()
        if raw:
            sheet_id = extract_sheet_id(raw)
            if sheet_id != raw:
                self.sheet_id_var.set(sheet_id)
                save_key(IMAGE_SHEET_CONFIG, sheet_id)
            threading.Thread(
                target=self._connect_sheet_bg, args=(sheet_id,), daemon=True
            ).start()

    def _on_connect_sheet(self):
        raw = self.sheet_id_var.get().strip()
        if not raw:
            messagebox.showwarning("시트", "이미지 시트 ID를 입력하세요.")
            return
        sheet_id = extract_sheet_id(raw)
        self.sheet_id_var.set(sheet_id)  # UI에도 정리된 ID 표시
        save_key(IMAGE_SHEET_CONFIG, sheet_id)
        self.sheet_status_var.set("연결 중...")
        threading.Thread(
            target=self._connect_sheet_bg, args=(sheet_id,), daemon=True
        ).start()

    def _connect_sheet_bg(self, sheet_id):
        spreadsheet, err = connect_sheet(sheet_id)
        if err:
            self.root.after(0, lambda: self.sheet_status_var.set(f"오류: {err}"))
            return

        self.image_spreadsheet = spreadsheet
        count = self.metadata_store.load_from_sheet(spreadsheet)

        drive, drive_err = connect_drive()
        if not drive_err:
            self.drive_service = drive

        products = self.metadata_store.get_products()

        # 임베딩 파일이 있으면 자동 로딩
        emb_loaded = self.metadata_store.load_embeddings()
        emb_msg = " | 임베딩 OK" if emb_loaded else " | 임베딩 없음 (생성 필요)"

        def _update():
            self.product_combo['values'] = products
            if products:
                self.product_combo.current(0)
            self.sheet_status_var.set(
                f"연결됨 (이미지 {count}개, 제품 {len(products)}개{emb_msg})")

        self.root.after(0, _update)

    # ────────────────────────────────────────
    #  원고 입력
    # ────────────────────────────────────────

    def _on_load_file(self):
        path = filedialog.askopenfilename(
            title="원고 파일 선택",
            filetypes=[
                ("지원 파일", "*.txt;*.docx"),
                ("텍스트 파일", "*.txt"),
                ("Word 파일", "*.docx"),
                ("모든 파일", "*.*"),
            ]
        )
        if not path:
            return
        try:
            if path.lower().endswith('.docx'):
                text = self._read_docx(path)
            else:
                with open(path, 'r', encoding='utf-8') as f:
                    text = f.read()
            self.manuscript_input.delete('1.0', tk.END)
            self.manuscript_input.insert('1.0', text)
            nums = extract_image_numbers(text)
            if nums:
                self.status_var.set(
                    f"파일 로딩 완료 — 이미지 위치 {len(nums)}개 "
                    f"({min(nums)}~{max(nums)}번)")
            else:
                self.status_var.set("파일 로딩 완료 — 이미지 위치 없음")
        except Exception as e:
            messagebox.showerror("파일 오류", str(e))

    def _read_docx(self, path):
        """DOCX 파일에서 텍스트 추출"""
        from docx import Document
        doc = Document(path)
        lines = []
        for para in doc.paragraphs:
            lines.append(para.text)
        return '\n'.join(lines)

    # ────────────────────────────────────────
    #  이미지 자동 제안
    # ────────────────────────────────────────

    def _on_analyze(self):
        if self.is_processing:
            return

        text = self.manuscript_input.get('1.0', tk.END).strip()
        if not text:
            messagebox.showwarning("원고", "원고를 먼저 입력하세요.")
            return
        gemini_key = self.gemini_key_var.get().strip()
        if not gemini_key:
            messagebox.showwarning("API", "Gemini API Key를 입력하세요.")
            return
        product = self.product_var.get()
        if not product:
            messagebox.showwarning("제품", "제품을 선택하세요.")
            return
        if not self.metadata_store.all:
            messagebox.showwarning("시트", "이미지 시트를 먼저 연결하세요.")
            return
        if not self.metadata_store.has_embeddings:
            messagebox.showwarning(
                "임베딩",
                "임베딩 데이터가 없습니다.\n"
                "먼저 '임베딩 생성' 버튼을 클릭하세요.")
            return

        image_nums = extract_image_numbers(text)
        if not image_nums:
            messagebox.showwarning(
                "원고",
                "원고에서 이미지 번호를 찾을 수 없습니다.\n"
                "(줄에 숫자만 있는 형태: 1, 2, 3...)")
            return

        self.manuscript_text = text
        self.is_processing = True
        self.analyze_btn.config(state='disabled')
        self.image_slots = {}
        self.status_var.set(f"이미지 {len(image_nums)}개 위치 분석 중...")

        threading.Thread(
            target=self._analyze_bg,
            args=(gemini_key, text, product, image_nums),
            daemon=True,
        ).start()

    def _analyze_bg(self, gemini_key, text, product, image_nums):
        """임베딩 기반 하이브리드 검색: 제품 필터 → 코사인 유사도 매칭"""
        try:
            # 원고에서 각 이미지 위치의 앞뒤 맥락 추출
            contexts = extract_image_contexts(text)
            ctx_map = {c['index']: c for c in contexts}

            # 각 이미지 위치의 맥락 텍스트 준비
            query_texts = []
            for num in image_nums:
                ctx = ctx_map.get(num, {})
                before = ctx.get('before', '')
                after = ctx.get('after', '')
                q = f"{before} {after}".strip()
                if not q:
                    q = product  # 빈 맥락이면 제품명으로 매칭
                query_texts.append(q)

            # 맥락 텍스트들을 한 번에 임베딩 (Gemini 1회 호출)
            self.root.after(0, lambda: self.status_var.set(
                f"원고 맥락 {len(query_texts)}개 임베딩 중..."))
            query_embeddings = get_embeddings(gemini_key, query_texts)

            # 각 이미지 위치별로 가장 유사한 이미지 매칭
            used_ids = set()
            slots = {}

            for i, num in enumerate(image_nums):
                # 임베딩 유사도로 상위 1개 선택 (이미 사용된 이미지 제외)
                results = self.metadata_store.search_by_embedding(
                    query_embeddings[i],
                    product=product,
                    exclude_ids=used_ids,
                    top_k=1,
                )

                match = results[0] if results else None

                if match:
                    used_ids.add(match['drive_file_id'])

                # 맥락 텍스트를 reason으로 활용
                ctx = ctx_map.get(num, {})
                before = ctx.get('before', '')[:60]
                after = ctx.get('after', '')[:60]
                reason = f"맥락: {before}...→...{after}" if before or after else ""

                slots[num] = {
                    "entry": match,
                    "scene_rec": query_texts[i][:100],
                    "mood_rec": "",
                    "reason": reason,
                    "status": "pending",
                    "bytes": None,
                    "query_embedding": query_embeddings[i],
                }

            self.image_slots = slots

            # 썸네일 미리 다운로드
            if self.drive_service:
                file_ids = [
                    s["entry"]["drive_file_id"]
                    for s in slots.values()
                    if s["entry"] and s["entry"].get("drive_file_id")
                ]
                if file_ids:
                    self.thumb_cache.download_batch(self.drive_service, file_ids)

            self.root.after(0, self._render_cards)

        except Exception as e:
            self.root.after(0, lambda err=e: self.status_var.set(f"분석 오류: {err}"))
        finally:
            self.is_processing = False
            self.root.after(0, lambda: self.analyze_btn.config(state='normal'))

    # ────────────────────────────────────────
    #  카드 렌더링
    # ────────────────────────────────────────

    def _render_cards(self):
        for w in self._cards_frame.winfo_children():
            w.destroy()
        self.tk_images.clear()

        accepted = sum(
            1 for s in self.image_slots.values()
            if s["status"] in ("accepted", "transformed"))
        total = len(self.image_slots)
        self.status_var.set(f"이미지 {total}개 제안됨 — 확정 {accepted}/{total}")

        for num in sorted(self.image_slots.keys()):
            self._build_card(num)

    def _build_card(self, img_num):
        slot = self.image_slots[img_num]
        entry = slot["entry"]
        status = slot["status"]

        card = tk.Frame(
            self._cards_frame, bg=self.THEME["card_bg"],
            relief=tk.RIDGE, bd=1, padx=10, pady=8)
        card.pack(fill=tk.X, padx=5, pady=(0, 5))

        # ── 상단: 번호 + 상태 ──
        top = tk.Frame(card, bg=self.THEME["card_bg"])
        top.pack(fill=tk.X)

        status_colors = {
            "pending": "#FF8F00",
            "accepted": self.THEME["ok"],
            "transformed": self.THEME["info"],
        }
        status_labels = {
            "pending": "검토 중",
            "accepted": "수락",
            "transformed": "변형됨",
        }

        tk.Label(
            top, text=f"이미지 {img_num}",
            font=('맑은 고딕', 11, 'bold'),
            bg=self.THEME["card_bg"], fg=self.THEME["fg"],
        ).pack(side=tk.LEFT)

        tk.Label(
            top, text=f" [{status_labels.get(status, status)}]",
            font=('맑은 고딕', 10, 'bold'),
            bg=self.THEME["card_bg"],
            fg=status_colors.get(status, "#666"),
        ).pack(side=tk.LEFT, padx=(5, 0))

        # ── 중단: 썸네일 + 정보 ──
        mid = tk.Frame(card, bg=self.THEME["card_bg"])
        mid.pack(fill=tk.X, pady=(5, 0))

        thumb_label = tk.Label(
            mid, text="[로딩]", bg=self.THEME["card_bg"], width=10, height=5)
        thumb_label.pack(side=tk.LEFT, padx=(0, 10))

        if entry and entry.get('drive_file_id'):
            tk_img = self.thumb_cache.get_tk_image(
                entry['drive_file_id'], size=(100, 100))
            if tk_img:
                self.tk_images[f"thumb_{img_num}"] = tk_img
                thumb_label.config(image=tk_img, text="", width=100, height=100)
            else:
                thumb_label.config(text="[썸네일 없음]")

        info = tk.Frame(mid, bg=self.THEME["card_bg"])
        info.pack(side=tk.LEFT, fill=tk.X, expand=True)

        if entry:
            scene = entry.get('scene', entry.get('filename', ''))
            mood = entry.get('mood', '')
            category = entry.get('category', '')
            tk.Label(
                info, text=f"장면: {scene}", font=('맑은 고딕', 10),
                bg=self.THEME["card_bg"], fg=self.THEME["fg"],
                wraplength=350, justify='left',
            ).pack(anchor='w')
            tk.Label(
                info, text=f"분위기: {mood}  |  카테고리: {category}",
                font=('맑은 고딕', 9),
                bg=self.THEME["card_bg"], fg='#666',
            ).pack(anchor='w')
        else:
            tk.Label(
                info, text="매칭된 이미지 없음", font=('맑은 고딕', 10),
                bg=self.THEME["card_bg"], fg=self.THEME["warn"],
            ).pack(anchor='w')

        reason = slot.get("reason", "")
        if reason:
            tk.Label(
                info, text=f"추천 이유: {reason}",
                font=('맑은 고딕', 9, 'italic'),
                bg=self.THEME["card_bg"], fg=self.THEME["info"],
                wraplength=350, justify='left',
            ).pack(anchor='w', pady=(3, 0))

        # ── 버튼 ──
        btn_frame = tk.Frame(card, bg=self.THEME["card_bg"])
        btn_frame.pack(fill=tk.X, pady=(5, 0))

        if status == "pending":
            ttk.Button(
                btn_frame, text="수락",
                command=lambda n=img_num: self._accept_image(n),
            ).pack(side=tk.LEFT, padx=(0, 5))
            ttk.Button(
                btn_frame, text="다른 이미지 보기",
                command=lambda n=img_num: self._show_alternatives(n),
            ).pack(side=tk.LEFT, padx=(0, 5))
        elif status in ("accepted", "transformed"):
            ttk.Button(
                btn_frame, text="다시 선택",
                command=lambda n=img_num: self._reset_image(n),
            ).pack(side=tk.LEFT, padx=(0, 5))

    # ────────────────────────────────────────
    #  수락 / 리셋
    # ────────────────────────────────────────

    def _accept_image(self, img_num):
        slot = self.image_slots[img_num]
        if not slot["entry"]:
            messagebox.showwarning("이미지", "매칭된 이미지가 없습니다.")
            return

        if (self.drive_service
                and slot["entry"].get("drive_file_id")
                and not slot["bytes"]):
            self.status_var.set(f"이미지 {img_num} 다운로드 중...")

            def _dl():
                try:
                    img_bytes = drive_download_bytes(
                        self.drive_service, slot["entry"]["drive_file_id"])
                    slot["bytes"] = img_bytes
                except Exception:
                    pass
                slot["status"] = "accepted"
                self.root.after(0, self._render_cards)

            threading.Thread(target=_dl, daemon=True).start()
        else:
            slot["status"] = "accepted"
            self._render_cards()

    def _reset_image(self, img_num):
        slot = self.image_slots[img_num]
        slot["status"] = "pending"
        slot["bytes"] = None
        self._render_cards()

    # ────────────────────────────────────────
    #  대체 이미지 5장 + 프롬프트 변형
    # ────────────────────────────────────────

    def _show_alternatives(self, img_num):
        slot = self.image_slots[img_num]
        product = self.product_var.get()

        # 중복 방지: 다른 슬롯에서 사용 중인 이미지 제외
        used_ids = set()
        for n, s in self.image_slots.items():
            if s["entry"] and s["entry"].get("drive_file_id") and n != img_num:
                used_ids.add(s["entry"]["drive_file_id"])
        if slot["entry"] and slot["entry"].get("drive_file_id"):
            used_ids.add(slot["entry"]["drive_file_id"])

        # 대체 이미지 검색 (임베딩 기반)
        query_emb = slot.get("query_embedding")
        alternatives = []
        if query_emb and self.metadata_store.has_embeddings:
            alternatives = self.metadata_store.search_by_embedding(
                query_emb, product=product, exclude_ids=used_ids, top_k=5)

        # 임베딩 결과가 부족하면 태그 필터로 보충
        alt_ids = {a['drive_file_id'] for a in alternatives}
        if len(alternatives) < 5:
            extra = self.metadata_store.filter(
                product=product, exclude_ids=used_ids | alt_ids)
            alternatives.extend(extra)
            alt_ids.update(e['drive_file_id'] for e in extra)
        if len(alternatives) < 5:
            extra = self.metadata_store.filter(
                product="공통", exclude_ids=used_ids | alt_ids)
            alternatives.extend(extra)
        alternatives = alternatives[:5]

        # 썸네일 다운로드
        if self.drive_service:
            ids = [a['drive_file_id'] for a in alternatives
                   if a.get('drive_file_id')]
            self.thumb_cache.download_batch(self.drive_service, ids)

        # ── 다이얼로그 ──
        dialog = tk.Toplevel(self.root)
        dialog.title(f"이미지 {img_num} — 대체 이미지 선택")
        dialog.geometry("700x650")
        dialog.transient(self.root)
        dialog.grab_set()
        dialog._tk_images = {}

        tk.Label(
            dialog, text=f"이미지 {img_num} — 대체 이미지",
            font=('맑은 고딕', 12, 'bold'),
        ).pack(padx=15, pady=(15, 5), anchor='w')

        if slot.get("scene_rec"):
            tk.Label(
                dialog, text=f"추천 장면: {slot['scene_rec']}",
                font=('맑은 고딕', 9, 'italic'), fg='#666',
            ).pack(padx=15, anchor='w')

        # 대체 이미지 리스트
        alt_frame = ttk.Frame(dialog)
        alt_frame.pack(fill=tk.BOTH, expand=True, padx=15, pady=10)

        for i, alt in enumerate(alternatives):
            row = tk.Frame(
                alt_frame, bg=self.THEME["card_bg"],
                relief=tk.GROOVE, bd=1, padx=8, pady=5)
            row.pack(fill=tk.X, pady=(0, 4))

            thumb_lbl = tk.Label(row, text="[로딩]", bg=self.THEME["card_bg"])
            thumb_lbl.pack(side=tk.LEFT, padx=(0, 10))

            if alt.get('drive_file_id'):
                tk_img = self.thumb_cache.get_tk_image(
                    alt['drive_file_id'], size=(70, 70))
                if tk_img:
                    dialog._tk_images[i] = tk_img
                    thumb_lbl.config(image=tk_img, text="", width=70, height=70)

            info_f = tk.Frame(row, bg=self.THEME["card_bg"])
            info_f.pack(side=tk.LEFT, fill=tk.X, expand=True)

            scene = alt.get('scene', alt.get('filename', ''))
            mood = alt.get('mood', '')
            tk.Label(
                info_f, text=scene, font=('맑은 고딕', 10),
                bg=self.THEME["card_bg"], wraplength=350, justify='left',
            ).pack(anchor='w')
            tk.Label(
                info_f, text=f"분위기: {mood}", font=('맑은 고딕', 9),
                bg=self.THEME["card_bg"], fg='#666',
            ).pack(anchor='w')

            ttk.Button(
                row, text="선택",
                command=lambda a=alt, d=dialog: self._select_alternative(
                    img_num, a, d),
            ).pack(side=tk.RIGHT, padx=(10, 0))

        if not alternatives:
            tk.Label(
                alt_frame, text="대체 이미지가 없습니다.",
                font=('맑은 고딕', 10), fg=self.THEME["warn"],
            ).pack(pady=20)

        # ── 프롬프트 입력 (AI 변형) ──
        ttk.Separator(dialog, orient='horizontal').pack(
            fill=tk.X, padx=15, pady=(5, 10))

        tk.Label(
            dialog, text="직접 프롬프트 입력 (현재 이미지를 AI로 변형)",
            font=('맑은 고딕', 10, 'bold'),
        ).pack(padx=15, anchor='w')

        prompt_text = tk.Text(
            dialog, height=3, font=('맑은 고딕', 10), wrap=tk.WORD)
        prompt_text.pack(fill=tk.X, padx=15, pady=(5, 5))
        prompt_text.insert(
            '1.0', "예: 30대 여성이 거울을 보며 미소짓는 모습으로 변경")

        prompt_btn_frame = ttk.Frame(dialog)
        prompt_btn_frame.pack(fill=tk.X, padx=15, pady=(0, 15))

        transform_status = tk.StringVar(value="")
        ttk.Label(
            prompt_btn_frame, textvariable=transform_status,
            font=('맑은 고딕', 9),
        ).pack(side=tk.LEFT)

        def _on_transform():
            user_prompt = prompt_text.get('1.0', tk.END).strip()
            if not user_prompt or user_prompt.startswith("예:"):
                messagebox.showwarning("프롬프트", "변형 프롬프트를 입력하세요.")
                return

            gemini_key = self.gemini_key_var.get().strip()
            if not gemini_key:
                messagebox.showwarning("API", "Gemini API Key를 입력하세요.")
                return

            transform_status.set("이미지 변형 중...")

            def _run():
                try:
                    contexts = extract_image_contexts(self.manuscript_text)
                    ctx = next(
                        (c for c in contexts if c['index'] == img_num), {})

                    full_prompt = build_transform_prompt(
                        context_before=ctx.get('before', ''),
                        context_after=ctx.get('after', ''),
                        image_metadata=slot.get("entry"),
                        user_instruction=user_prompt,
                    )

                    ref_bytes = None
                    if (self.drive_service and slot["entry"]
                            and slot["entry"].get("drive_file_id")):
                        try:
                            ref_bytes = drive_download_bytes(
                                self.drive_service,
                                slot["entry"]["drive_file_id"])
                        except Exception:
                            pass

                    if ref_bytes:
                        img_bytes = transform_image(
                            gemini_key, ref_bytes, full_prompt)
                    else:
                        img_bytes = generate_image(gemini_key, full_prompt)

                    slot["bytes"] = img_bytes
                    slot["status"] = "transformed"

                    self.root.after(0, lambda: (
                        transform_status.set("변형 완료!"),
                        self._render_cards(),
                    ))

                except Exception as e:
                    self.root.after(0, lambda err=e: transform_status.set(
                        f"오류: {err}"))

            threading.Thread(target=_run, daemon=True).start()

        ttk.Button(
            prompt_btn_frame, text="변형 실행", command=_on_transform,
        ).pack(side=tk.RIGHT)

    def _select_alternative(self, img_num, alt_entry, dialog):
        slot = self.image_slots[img_num]
        slot["entry"] = alt_entry
        slot["bytes"] = None
        slot["status"] = "pending"
        self._accept_image(img_num)
        dialog.destroy()

    # ────────────────────────────────────────
    #  최종 저장
    # ────────────────────────────────────────

    def _on_save_all(self):
        if not self.image_slots:
            messagebox.showwarning("저장", "먼저 이미지 자동 제안을 실행하세요.")
            return

        pending = [
            n for n, s in self.image_slots.items()
            if s["status"] == "pending"
        ]
        if pending:
            ans = messagebox.askyesno(
                "미확정 이미지",
                f"아직 확정되지 않은 이미지가 {len(pending)}개 있습니다.\n"
                f"({', '.join(str(n) for n in sorted(pending))}번)\n\n"
                "미확정 이미지는 제안된 이미지로 자동 저장됩니다.\n"
                "계속하시겠습니까?")
            if not ans:
                return

        folder = filedialog.askdirectory(
            title="저장 폴더 선택", initialdir=OUTPUT_DIR)
        if not folder:
            return

        self.status_var.set("저장 중...")
        threading.Thread(
            target=self._save_all_bg, args=(folder,), daemon=True
        ).start()

    def _save_all_bg(self, folder):
        try:
            saved_count = 0
            errors = []

            for num in sorted(self.image_slots.keys()):
                slot = self.image_slots[num]

                # 이미 변형/다운로드된 이미지가 있으면 그대로 사용
                if slot["bytes"]:
                    filepath = os.path.join(folder, f"{num}.png")
                    with open(filepath, 'wb') as f:
                        f.write(slot["bytes"])
                    saved_count += 1
                    continue

                # Drive에서 다운로드
                if (self.drive_service and slot["entry"]
                        and slot["entry"].get("drive_file_id")):
                    try:
                        img_bytes = drive_download_bytes(
                            self.drive_service,
                            slot["entry"]["drive_file_id"])

                        # 확장자 판별
                        if img_bytes[:8].startswith(b'\x89PNG'):
                            ext = "png"
                        elif img_bytes[:2] == b'\xff\xd8':
                            ext = "jpg"
                        elif img_bytes[:4] == b'RIFF':
                            ext = "webp"
                        else:
                            ext = "png"

                        filepath = os.path.join(folder, f"{num}.{ext}")
                        with open(filepath, 'wb') as f:
                            f.write(img_bytes)
                        saved_count += 1
                    except Exception as e:
                        errors.append(f"이미지 {num}: {e}")
                else:
                    errors.append(f"이미지 {num}: 이미지 없음")

            # 원고 Word 저장
            text = self.manuscript_input.get('1.0', tk.END).strip()
            if text:
                from docx import Document
                doc = Document()
                for line in text.split('\n'):
                    doc.add_paragraph(line)
                docx_path = os.path.join(folder, "원고.docx")
                doc.save(docx_path)

            msg = f"저장 완료: 이미지 {saved_count}개 + 원고.docx"
            if errors:
                msg += f"\n\n오류 {len(errors)}건:\n" + "\n".join(errors)

            self.root.after(0, lambda: (
                self.status_var.set(f"저장 완료 ({saved_count}개)"),
                messagebox.showinfo("저장 완료", msg),
            ))

        except Exception as e:
            self.root.after(0, lambda err=e: (
                self.status_var.set(f"저장 오류: {err}"),
                messagebox.showerror("저장 오류", str(err)),
            ))

    # ────────────────────────────────────────
    #  실행
    # ────────────────────────────────────────

    def run(self):
        self.root.mainloop()


if __name__ == "__main__":
    app = ImageGeneratorApp()
    app.run()
